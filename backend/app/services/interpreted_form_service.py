"""
Interpreted Form Service for general purpose form filling
Handles data upload to Pinecone, template analysis, and form generation
Enhanced with deterministic document filling to solve field confusion issues
"""

import os
import json
import uuid
import re
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import tempfile
import zipfile
from datetime import datetime
"""
Interpreted Form Service for AI-powered document filling

This service handles the complete workflow of intelligent form filling:
1. Upload person data to Pinecone vector database
2. Analyze templates to identify fillable fields
3. Search person data using contextual queries
4. Fill forms with appropriate data

IMPORTANT: Context length limits implemented to prevent Google 500 errors:
- _extract_value_from_text: Limits context to 2000 chars with smart truncation
- _generate_document_content: Limits field data to 1500 chars for API safety
"""

import asyncio
import logging
import shutil
from docx import Document
from docx.shared import Inches

from .gemini_service import gemini_service
from .pinecone_service import pinecone_service
from .llm_service import LLMService, DocumentProcessor
from .deterministic_document_filler import DeterministicDocumentFiller
from .template_mapping_service import TemplateMappingService

logger = logging.getLogger(__name__)

class InterpretedFormService:
    """Service for handling interpreted form filling workflow"""
    
    def __init__(self):
        self.gemini_service = gemini_service
        self.pinecone_service = pinecone_service
        self.doc_processor = DocumentProcessor()
        self.output_dir = Path("filled_templates")
        self.output_dir.mkdir(exist_ok=True)
        
        # Enhanced services for deterministic document filling
        self.deterministic_filler = DeterministicDocumentFiller()
        self.mapping_service = TemplateMappingService()
        
        # Create separate storage for person data
        self.person_storage_dir = Path("local_vector_storage/persons")
        self.person_storage_dir.mkdir(exist_ok=True, parents=True)
        
        logger.info("InterpretedFormService initialized with deterministic filling capabilities")
    
    def _get_person_storage_file(self, person_id: str) -> Path:
        """Get storage file path for person data"""
        return self.person_storage_dir / f"person_{person_id}_data.json"
    
    async def cleanup_person_data(self, person_id: str = None, max_age_hours: int = 24) -> Dict[str, Any]:
        """
        Clean up old person data chunks. Called on page refresh or periodically.
        
        Args:
            person_id: Specific person to clean up, or None for all
            max_age_hours: Remove data older than this many hours
        """
        try:
            import time
            current_time = time.time()
            cleanup_threshold = current_time - (max_age_hours * 3600)  # Convert hours to seconds
            
            cleaned_files = []
            total_cleaned = 0
            
            if person_id:
                # Clean specific person
                person_file = self._get_person_storage_file(person_id)
                if person_file.exists():
                    person_file.unlink()
                    cleaned_files.append(person_id)
                    total_cleaned = 1
                    logger.info(f"🧹 Cleaned person data for: {person_id}")
            else:
                # Clean all old person data
                if self.person_storage_dir.exists():
                    for person_file in self.person_storage_dir.glob("person_*_data.json"):
                        try:
                            # Check file age
                            file_age = person_file.stat().st_mtime
                            if file_age < cleanup_threshold:
                                person_id_from_file = person_file.stem.replace('person_', '').replace('_data', '')
                                person_file.unlink()
                                cleaned_files.append(person_id_from_file)
                                total_cleaned += 1
                                logger.info(f"🧹 Cleaned old person data: {person_id_from_file}")
                        except Exception as e:
                            logger.warning(f"⚠️ Failed to clean {person_file}: {e}")
            
            return {
                'success': True,
                'cleaned_persons': cleaned_files,
                'total_cleaned': total_cleaned,
                'message': f'Cleaned {total_cleaned} person data files'
            }
            
        except Exception as e:
            logger.error(f"❌ Failed to cleanup person data: {e}")
            return {
                'success': False,
                'error': str(e),
                'message': 'Failed to cleanup person data'
            }
    
    async def cleanup_all_person_data(self) -> Dict[str, Any]:
        """
        Clean up ALL person data chunks (for page refresh/reset)
        """
        try:
            cleaned_files = []
            total_cleaned = 0
            
            if self.person_storage_dir.exists():
                for person_file in self.person_storage_dir.glob("person_*_data.json"):
                    try:
                        person_id_from_file = person_file.stem.replace('person_', '').replace('_data', '')
                        person_file.unlink()
                        cleaned_files.append(person_id_from_file)
                        total_cleaned += 1
                        logger.info(f"🧹 Cleaned person data: {person_id_from_file}")
                    except Exception as e:
                        logger.warning(f"⚠️ Failed to clean {person_file}: {e}")
            
            logger.info(f"🧹 Total cleanup: {total_cleaned} person data files removed")
            
            return {
                'success': True,
                'cleaned_persons': cleaned_files,
                'total_cleaned': total_cleaned,
                'message': f'Cleaned all {total_cleaned} person data files'
            }
            
        except Exception as e:
            logger.error(f"❌ Failed to cleanup all person data: {e}")
            return {
                'success': False,
                'error': str(e),
                'message': 'Failed to cleanup all person data'
            }
    
    async def _store_person_vectors(self, person_id: str, vectors: List[Dict[str, Any]]) -> bool:
        """Store person vectors in separate storage"""
        try:
            storage_file = self._get_person_storage_file(person_id)
            
            # Load existing data if any
            existing_vectors = []
            if storage_file.exists():
                with open(storage_file, 'r', encoding='utf-8') as f:
                    existing_vectors = json.load(f)
            
            # Add new vectors
            existing_vectors.extend(vectors)
            
            # Save updated data
            with open(storage_file, 'w', encoding='utf-8') as f:
                json.dump(existing_vectors, f, indent=2)
            
            logger.info(f"✅ Stored {len(vectors)} vectors for person {person_id}")
            return True
            
        except Exception as e:
            logger.error(f"❌ Failed to store person vectors: {e}")
            return False
    
    async def _search_person_vectors(self, person_id: str, query_embedding: List[float], top_k: int = 5) -> List[Dict[str, Any]]:
        """Search vectors for a specific person"""
        try:
            storage_file = self._get_person_storage_file(person_id)
            
            if not storage_file.exists():
                logger.warning(f"❌ No data found for person {person_id}")
                return []
            
            # Load person data
            with open(storage_file, 'r', encoding='utf-8') as f:
                person_vectors = json.load(f)
            
            # Calculate similarity scores
            results = []
            for vector in person_vectors:
                if 'embedding' in vector and 'values' in vector:
                    embedding = vector.get('values', vector.get('embedding', []))
                    if embedding:
                        # Calculate cosine similarity
                        similarity = self._calculate_cosine_similarity(query_embedding, embedding)
                        results.append({
                            'text': vector.get('text', ''),
                            'metadata': vector.get('metadata', {}),
                            'score': similarity
                        })
            
            # Sort by similarity and return top_k
            results.sort(key=lambda x: x['score'], reverse=True)
            return results[:top_k]
            
        except Exception as e:
            logger.error(f"❌ Failed to search person vectors: {e}")
            return []
    
    def _calculate_cosine_similarity(self, vec1: List[float], vec2: List[float]) -> float:
        """Calculate cosine similarity between two vectors"""
        try:
            import numpy as np
            
            # Convert to numpy arrays
            a = np.array(vec1)
            b = np.array(vec2)
            
            # Calculate cosine similarity
            dot_product = np.dot(a, b)
            norms = np.linalg.norm(a) * np.linalg.norm(b)
            
            if norms == 0:
                return 0.0
            
            return dot_product / norms
            
        except Exception as e:
            logger.error(f"❌ Failed to calculate similarity: {e}")
            return 0.0
        
    async def upload_person_data(
        self, 
        person_data_file: bytes, 
        filename: str, 
        device_id: str,
        person_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Upload person data to Pinecone database
        Supports JSON, CSV, TXT, PDF, DOCX formats
        """
        try:
            # Generate unique person ID if not provided
            if not person_id:
                person_id = str(uuid.uuid4())
            
            logger.info(f"🔄 Starting person data upload - device_id: {device_id}, person_id: {person_id}, filename: {filename}")
            
            # Create temporary file for processing
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as tmp_file:
                tmp_file.write(person_data_file)
                tmp_file_path = tmp_file.name
            
            try:
                # Extract content based on file type
                file_ext = Path(filename).suffix.lower()
                logger.info(f"📁 Processing {file_ext} file: {filename}")
                
                if file_ext == '.json':
                    with open(tmp_file_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    # Convert JSON to text chunks for embedding
                    chunks = self._json_to_chunks(data, person_id)
                
                elif file_ext == '.csv':
                    import csv
                    chunks = []
                    with open(tmp_file_path, 'r', encoding='utf-8') as f:
                        reader = csv.DictReader(f)
                        rows = list(reader)
                        
                        # Process each row for better chunking
                        for row_idx, row in enumerate(rows):
                            row_chunks = self._process_csv_row(row, row_idx, person_id, device_id, filename)
                            chunks.extend(row_chunks)
                        
                        # Create summary chunks for the entire CSV
                        if rows:
                            summary_chunks = self._create_csv_summary_chunks(rows, person_id, device_id, filename)
                            chunks.extend(summary_chunks)
                
                elif file_ext in ['.txt', '.pdf', '.docx']:
                    # Extract text and create chunks
                    text_content = self.doc_processor.extract_text(tmp_file_path)
                    chunks = self._text_to_chunks(text_content, person_id, device_id, filename)
                
                else:
                    raise ValueError(f"Unsupported file format: {file_ext}")
                
                logger.info(f"📊 Created {len(chunks)} chunks from {filename}")
                
                # Generate embeddings and store in Pinecone
                vectors = []
                for i, chunk in enumerate(chunks):
                    try:
                        logger.info(f"🔤 Generating embedding for chunk {i+1}/{len(chunks)}: '{chunk['text'][:100]}...'")
                        embedding = await self.gemini_service.get_embedding(chunk['text'])
                        vector = {
                            'id': f"person_{person_id}_{uuid.uuid4()}",
                            'values': embedding,  # For compatibility
                            'embedding': embedding,  # For local search
                            'metadata': chunk['metadata'],
                            'text': chunk['text']  # Store text for search results
                        }
                        vectors.append(vector)
                        logger.info(f"✅ Generated embedding {i+1}/{len(chunks)}")
                    except Exception as e:
                        logger.warning(f"Failed to generate embedding for chunk {i}: {e}")
                        continue
                
                logger.info(f"🎯 Generated {len(vectors)} vectors for upload to person storage")
                
                # Store in separate person storage instead of device storage
                success = await self._store_person_vectors(person_id, vectors)
                
                if success:
                    logger.info(f"✅ Successfully stored {len(vectors)} person data vectors for person {person_id}")
                    return {
                        'success': True,
                        'person_id': person_id,
                        'chunks_processed': len(vectors),
                        'storage_type': 'person_storage',
                        'message': f'Successfully uploaded person data with {len(vectors)} data points'
                    }
                else:
                    raise Exception("Failed to store person vectors")
                    
            finally:
                # Clean up temporary file
                os.unlink(tmp_file_path)
                
        except Exception as e:
            logger.error(f"❌ Failed to upload person data: {e}")
            return {
                'success': False,
                'error': str(e),
                'message': 'Failed to upload person data'
            }
    
    async def analyze_template(
        self, 
        template_file: bytes, 
        filename: str, 
        device_id: str
    ) -> Dict[str, Any]:
        """
        Analyze template to identify blank spaces and field requirements
        """
        try:
            # Create temporary file for processing
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as tmp_file:
                tmp_file.write(template_file)
                tmp_file_path = tmp_file.name
            
            try:
                # Extract text from template
                template_text = self.doc_processor.extract_text(tmp_file_path)
                
                # Use LLM to analyze template and identify blank fields
                analysis = await self._analyze_template_with_llm(template_text, filename)
                
                # Store template file for later use in form filling
                template_id = str(uuid.uuid4())
                stored_template_path = self.output_dir / f"template_{template_id}_{filename}"
                shutil.copy2(tmp_file_path, stored_template_path)
                
                # Add template path to analysis
                analysis['template_path'] = str(stored_template_path)
                analysis['template_id'] = template_id
                
                return {
                    'success': True,
                    'template_analysis': analysis,
                    'message': f'Successfully analyzed template: {filename}'
                }
                
            finally:
                # Clean up temporary file
                os.unlink(tmp_file_path)
                
        except Exception as e:
            logger.error(f"❌ Failed to analyze template: {e}")
            return {
                'success': False,
                'error': str(e),
                'message': 'Failed to analyze template'
            }
    
    async def fill_form_deterministic(
        self,
        template_content: str,
        input_data: Dict[str, Any],
        template_name: Optional[str] = None,
        output_filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Fill form using deterministic mapping approach to solve field confusion issues.
        This is the enhanced method that uses explicit field mapping.
        
        Args:
            template_content: The template content to analyze
            input_data: Dictionary containing the form data
            template_name: Optional template name (auto-detected if not provided)
            output_filename: Optional output filename
            
        Returns:
            Dictionary with success status, filled document path, and metadata
        """
        try:
            logger.info("🎯 Starting deterministic form filling...")
            
            # Auto-detect template type if not specified
            detected_template = None
            if not template_name:
                detected_template = self.mapping_service.detect_template_type(template_content)
                if detected_template:
                    template_name = detected_template
                    logger.info(f"📋 Auto-detected template type: {template_name}")
                else:
                    logger.warning("⚠️ Could not auto-detect template type")
            
            # Get available templates for reference
            available_templates = self.mapping_service.get_available_templates()
            logger.info(f"📚 Available templates: {[t['name'] for t in available_templates]}")
            
            # Try deterministic filling if template is recognized
            if template_name and template_name in [t['name'] for t in available_templates]:
                logger.info(f"✅ Using deterministic filling for: {template_name}")
                
                success, message, filled_doc = self.deterministic_filler.fill_document_deterministic(
                    template_content, input_data, template_name
                )
                
                if success and filled_doc:
                    # Generate output filename if not provided
                    if not output_filename:
                        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                        output_filename = f"filled_{template_name}_{timestamp}.docx"
                    
                    # Save the document
                    output_path = self.deterministic_filler.save_filled_document(
                        filled_doc, 
                        str(self.output_dir), 
                        output_filename
                    )
                    
                    # Get template preview for validation
                    template_preview = self.deterministic_filler.get_template_preview(template_name)
                    
                    return {
                        'success': True,
                        'method': 'deterministic_mapping',
                        'template_type': template_name,
                        'detected_template': detected_template,
                        'filled_document_path': output_path,
                        'filename': output_filename,
                        'message': message,
                        'template_preview': template_preview,
                        'validation_passed': True,
                        'fields_filled': len(input_data),
                        'available_templates': available_templates
                    }
                else:
                    logger.error(f"❌ Deterministic filling failed: {message}")
                    return {
                        'success': False,
                        'method': 'deterministic_mapping',
                        'template_type': template_name,
                        'error': message,
                        'available_templates': available_templates
                    }
            
            else:
                # Fallback to original method for unrecognized templates
                logger.info("📝 Template not recognized for deterministic filling, using fallback approach")
                
                # You can add fallback logic here or return guidance
                return {
                    'success': False,
                    'method': 'deterministic_mapping',
                    'template_type': template_name or 'unknown',
                    'detected_template': detected_template,
                    'error': f'Template type not supported for deterministic filling. Supported types: {[t["name"] for t in available_templates]}',
                    'available_templates': available_templates,
                    'suggestion': 'Please specify one of the supported template types or use the legacy filling method'
                }
                
        except Exception as e:
            logger.error(f"❌ Error in deterministic form filling: {str(e)}", exc_info=True)
            return {
                'success': False,
                'method': 'deterministic_mapping',
                'error': str(e),
                'message': 'Deterministic form filling failed'
            }
    
    async def generate_filled_forms(
        self,
        person_id: str,
        templates: List[Dict[str, Any]],  # List of template analyses
        device_id: str,
        batch_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Generate filled forms using person data from Pinecone and template analyses
        Supports up to 5 templates at once
        """
        try:
            if len(templates) > 5:
                raise ValueError("Maximum 5 templates allowed per batch")
            
            if not batch_id:
                batch_id = str(uuid.uuid4())
            
            filled_documents = []
            download_links = []
            
            for template_idx, template in enumerate(templates):
                try:
                    # Fill individual template
                    filled_doc = await self._fill_single_template(
                        person_id, 
                        template, 
                        device_id, 
                        batch_id, 
                        template_idx
                    )
                    
                    if filled_doc['success']:
                        filled_documents.append(filled_doc)
                        download_links.append(filled_doc['download_url'])
                    
                except Exception as e:
                    logger.error(f"❌ Failed to fill template {template_idx}: {e}")
                    continue
            
            # Create batch download if multiple documents
            if len(filled_documents) > 1:
                batch_download_url = await self._create_batch_download(
                    filled_documents, 
                    batch_id
                )
                download_links.append(batch_download_url)
            
            return {
                'success': True,
                'batch_id': batch_id,
                'filled_documents': filled_documents,
                'download_links': download_links,
                'total_documents': len(filled_documents),
                'message': f'Successfully generated {len(filled_documents)} filled documents'
            }
            
        except Exception as e:
            logger.error(f"❌ Failed to generate filled forms: {e}")
            return {
                'success': False,
                'error': str(e),
                'message': 'Failed to generate filled forms'
            }
    
    async def _analyze_template_with_llm(
        self, 
        template_text: str, 
        filename: str
    ) -> Dict[str, Any]:
        """
        Use LLM to analyze template and identify fields to fill with enhanced contextual understanding
        """
        prompt = f"""
You are an expert legal document analyzer specializing in affidavits and wills. Analyze this template and identify ALL fillable fields with PRECISE LEGAL CONTEXTUAL UNDERSTANDING.

Template: {filename}
Content: {template_text[:2500]}

CRITICAL REQUIREMENTS:
1. Return ONLY valid JSON - no extra text, explanations, or markdown
2. This is a LEGAL AFFIDAVIT - understand legal roles and terminology
3. Identify ALL blank spaces with dots: ............., _____, [FIELD], etc.
4. Different legal roles (testator, witness A, witness B, etc.) need different field IDs

LEGAL CONTEXT PATTERNS TO RECOGNIZE:
- "executed by Shri ............" = testator name
- "son of ............" (after testator context) = testator father name
- "resident of ............" (after testator context) = testator address  
- "A, aged about ............" = witness A age
- "B, aged about ............" = witness B age
- "son of Shri ............" (after witness A) = witness A father name
- "resident of ............" (after witness A) = witness A address
- "executed his will on ............" = will execution date
- "Sub-Registrar ............" = registrar office/name
- "Verified at ............" = verification location

Return this EXACT JSON structure:
{{
  "template_info": {{
    "filename": "{filename}",
    "estimated_fields": 0,
    "template_type": "legal_affidavit"
  }},
  "identified_fields": [
  ]
}}

For each field, add to identified_fields array:
{{
  "field_id": "role_type_number",
  "field_type": "name|date|address|location|age|text",
  "semantic_role": "testator|witness_a|witness_b|father_of_testator|father_of_witness_a|father_of_witness_b|sub_registrar|general",
  "context": "surrounding text with legal terms",
  "description": "specific legal role description",
  "required": true,
  "contextual_keywords": ["legal", "role", "keywords"],
  "placeholder_patterns": ["dots pattern"]
}}

LEGAL ROLE EXAMPLES FOR AFFIDAVITS:
- "executed by Shri ............" → field_id: "testator_name_1", semantic_role: "testator", field_type: "name"
- "son of ............" (testator context) → field_id: "testator_father_name_1", semantic_role: "father_of_testator", field_type: "name"
- "resident of ............" (testator context) → field_id: "testator_address_1", semantic_role: "testator", field_type: "address"
- "A, aged about ............" → field_id: "witness_a_age_1", semantic_role: "witness_a", field_type: "age"
- "B, aged about ............" → field_id: "witness_b_age_1", semantic_role: "witness_b", field_type: "age"

IMPORTANT: 
- Distinguish between testator vs witness A vs witness B details
- Make descriptions SHORT and clear. No quotes in descriptions.
- Each context should get unique field_id based on legal role

Return ONLY the JSON object."""
        
        try:
            response = await self.gemini_service.generate_response(
                prompt=prompt,
                max_tokens=3000,
                temperature=0.1
            )
            
            # Multiple attempts to parse JSON
            for attempt in range(3):
                try:
                    # Clean the response
                    response_text = response.strip()
                    
                    # Remove markdown formatting
                    if response_text.startswith('```json'):
                        response_text = response_text[7:]
                    elif response_text.startswith('```'):
                        response_text = response_text[3:]
                    if response_text.endswith('```'):
                        response_text = response_text[:-3]
                    response_text = response_text.strip()
                    
                    # Apply different levels of fixing based on attempt
                    if attempt == 0:
                        # First attempt: minimal cleaning
                        cleaned_text = response_text
                    elif attempt == 1:
                        # Second attempt: basic JSON fixing
                        cleaned_text = self._fix_json_formatting(response_text)
                    else:
                        # Third attempt: aggressive fixing
                        cleaned_text = self._aggressive_json_fix(response_text)
                    
                    # Try to parse
                    analysis = json.loads(cleaned_text)
                    
                    # Validate structure
                    if self._validate_analysis_structure(analysis):
                        logger.info(f"✅ JSON parsed successfully on attempt {attempt + 1}")
                        return analysis
                    else:
                        logger.warning(f"⚠️ Invalid structure on attempt {attempt + 1}")
                        
                except json.JSONDecodeError as e:
                    logger.warning(f"❌ JSON parsing attempt {attempt + 1} failed: {e}")
                    if attempt == 2:  # Last attempt
                        logger.error(f"Response text: {response_text[:500]}...")
                
            # If all JSON parsing attempts fail, use fallback
            logger.warning("🔄 Using fallback template analysis")
            return await self._fallback_template_analysis(template_text, filename)
            
        except Exception as e:
            logger.error(f"❌ LLM template analysis failed: {e}")
            return await self._fallback_template_analysis(template_text, filename)

    def _fix_json_formatting(self, json_text: str) -> str:
        """
        Comprehensive JSON formatting fix for LLM responses
        """
        try:
            # Remove any markdown or extra formatting
            json_text = json_text.strip()
            if json_text.startswith('```json'):
                json_text = json_text[7:]
            if json_text.startswith('```'):
                json_text = json_text[3:]
            if json_text.endswith('```'):
                json_text = json_text[:-3]
            json_text = json_text.strip()
            
            # Fix common issues
            # 1. Remove any text before the first {
            first_brace = json_text.find('{')
            if first_brace > 0:
                json_text = json_text[first_brace:]
            
            # 2. Remove any text after the last }
            last_brace = json_text.rfind('}')
            if last_brace >= 0:
                json_text = json_text[:last_brace + 1]
            
            # 3. Fix unterminated strings
            lines = json_text.split('\n')
            fixed_lines = []
            in_string = False
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                # Count unescaped quotes
                quote_count = 0
                i = 0
                while i < len(line):
                    if line[i] == '"' and (i == 0 or line[i-1] != '\\'):
                        quote_count += 1
                    i += 1
                
                # If odd number of quotes, likely unterminated
                if quote_count % 2 == 1:
                    # Find the last quote and close it
                    if line.endswith(',') or line.endswith('{') or line.endswith('['):
                        line = line[:-1] + '"' + line[-1]
                    else:
                        line = line + '"'
                
                fixed_lines.append(line)
            
            json_text = '\n'.join(fixed_lines)
            
            # 4. Remove trailing commas before closing braces/brackets
            json_text = re.sub(r',(\s*[}\]])', r'\1', json_text)
            
            # 5. Ensure proper structure
            if not json_text.startswith('{'):
                json_text = '{' + json_text
            if not json_text.endswith('}'):
                json_text = json_text + '}'
            
            return json_text
            
        except Exception as e:
            logger.warning(f"JSON fix failed: {e}")
            return json_text

    def _aggressive_json_fix(self, json_text: str) -> str:
        """
        Aggressive JSON fixing for stubborn cases
        """
        try:
            # Remove everything before first { and after last }
            start = json_text.find('{')
            end = json_text.rfind('}')
            if start >= 0 and end >= 0:
                json_text = json_text[start:end+1]
            
            # Replace problematic characters
            json_text = json_text.replace('\n', ' ')
            json_text = json_text.replace('\t', ' ')
            json_text = re.sub(r'\s+', ' ', json_text)
            
            # Remove trailing commas
            json_text = re.sub(r',(\s*[}\]])', r'\1', json_text)
            
            # Ensure proper array/object endings
            json_text = re.sub(r',\s*}', '}', json_text)
            json_text = re.sub(r',\s*]', ']', json_text)
            
            return json_text
            
        except Exception as e:
            logger.warning(f"Aggressive JSON fix failed: {e}")
            return json_text

    def _validate_analysis_structure(self, analysis: dict) -> bool:
        """
        Validate the analysis structure
        """
        try:
            # Check required keys
            if not isinstance(analysis, dict):
                return False
            
            if 'template_info' not in analysis or 'identified_fields' not in analysis:
                return False
            
            template_info = analysis['template_info']
            if not isinstance(template_info, dict):
                return False
            
            identified_fields = analysis['identified_fields']
            if not isinstance(identified_fields, list):
                return False
            
            # Validate each field
            for field in identified_fields:
                if not isinstance(field, dict):
                    return False
                required_keys = ['field_id', 'field_type', 'semantic_role']
                if not all(key in field for key in required_keys):
                    return False
            
            return True
            
        except Exception as e:
            logger.warning(f"Validation failed: {e}")
            return False

    async def _fallback_template_analysis(
        self, 
        template_text: str, 
        filename: str
    ) -> Dict[str, Any]:
        """
        Fallback template analysis using pattern matching when LLM fails
        """
        import re
        
        # Look for common placeholder patterns
        patterns = {
            'underscores': r'_{3,}',  # Three or more underscores
            'brackets': r'\[([^\]]+)\]',  # Text in brackets
            'braces': r'\{([^}]+)\}',  # Text in braces
            'parentheses': r'\(([^)]*blank[^)]*)\)',  # Parentheses with "blank"
        }
        
        identified_fields = []
        field_counter = 0
        
        for pattern_name, pattern in patterns.items():
            matches = re.finditer(pattern, template_text, re.IGNORECASE)
            for match in matches:
                field_counter += 1
                field_id = f"field_{field_counter}"
                
                # Extract context around the match
                start = max(0, match.start() - 50)
                end = min(len(template_text), match.end() + 50)
                context = template_text[start:end].strip()
                
                # Try to guess field type and semantic role based on context
                field_type, semantic_role = self._guess_field_type(context)
                
                # Create contextual field_id
                contextual_field_id = f"{semantic_role}_{field_type}_{field_counter}" if semantic_role != "general" else f"{field_type}_{field_counter}"
                
                identified_fields.append({
                    "field_id": contextual_field_id,
                    "field_type": field_type,
                    "semantic_role": semantic_role,
                    "context": context,
                    "description": f"{semantic_role.title()} {field_type}" if semantic_role != "general" else f"Field requiring {field_type}",
                    "required": True,
                    "contextual_keywords": [semantic_role, field_type],
                    "placeholder_patterns": [match.group(0)]
                })
        
        return {
            "template_info": {
                "filename": filename,
                "estimated_fields": len(identified_fields),
                "template_type": "general_form"
            },
            "identified_fields": identified_fields
        }

    def _guess_field_type(self, context: str) -> tuple[str, str]:
        """
        Guess field type and semantic role based on context
        Returns (field_type, semantic_role)
        """
        context_lower = context.lower()
        
        # Determine semantic role
        semantic_role = "general"
        if any(word in context_lower for word in ['testator', 'deceased', 'deponent']):
            semantic_role = "testator"
        elif any(word in context_lower for word in ['witness', 'attestor', 'attesting']):
            semantic_role = "witness"
        elif any(word in context_lower for word in ['applicant', 'petitioner', 'plaintiff']):
            semantic_role = "applicant"
        elif any(word in context_lower for word in ['father', 'father\'s', 'paternal']):
            semantic_role = "father"
        elif any(word in context_lower for word in ['mother', 'mother\'s', 'maternal']):
            semantic_role = "mother"
        elif any(word in context_lower for word in ['spouse', 'husband', 'wife']):
            semantic_role = "spouse"
        elif any(word in context_lower for word in ['guardian', 'custodian']):
            semantic_role = "guardian"
        elif any(word in context_lower for word in ['beneficiary', 'heir', 'inheritor']):
            semantic_role = "beneficiary"
        
        # Determine field type
        field_type = "text"
        if any(word in context_lower for word in ['age', 'years old', 'aged', 'years']):
            field_type = 'age'
        elif any(word in context_lower for word in ['name', 'full name', 'first name', 'last name']):
            field_type = 'name'
        elif any(word in context_lower for word in ['date', 'day', 'month', 'year', 'birth', 'dob']) and 'age' not in context_lower:
            field_type = 'date'
        elif any(word in context_lower for word in ['address', 'street', 'city', 'state', 'residence']):
            field_type = 'address'
        elif any(word in context_lower for word in ['phone', 'telephone', 'mobile', 'contact']):
            field_type = 'phone'
        elif any(word in context_lower for word in ['email', 'e-mail', 'mail']):
            field_type = 'email'
        elif any(word in context_lower for word in ['signature', 'sign', 'signed']):
            field_type = 'signature'
        elif any(word in context_lower for word in ['number', 'amount', 'value']):
            field_type = 'number'
        
        return field_type, semantic_role

    def _get_suggestions_for_type(self, field_type: str) -> List[str]:
        """
        Get search suggestions based on field type
        """
        suggestions_map = {
            'name': ['full_name', 'first_name', 'last_name', 'legal_name', 'name'],
            'date': ['date', 'birth_date', 'current_date', 'today'],
            'address': ['address', 'home_address', 'mailing_address', 'street_address'],
            'phone': ['phone', 'mobile', 'telephone', 'contact_number'],
            'email': ['email', 'email_address', 'contact_email'],
            'signature': ['signature', 'full_name', 'name'],
            'number': ['age', 'number', 'amount', 'quantity'],
            'text': ['information', 'details', 'description']
        }
        return suggestions_map.get(field_type, ['information', 'details'])

    async def _fill_single_template(
        self,
        person_id: str,
        template: Dict[str, Any],
        device_id: str,
        batch_id: str,
        template_idx: int
    ) -> Dict[str, Any]:
        """
        Fill a single template using person data from Pinecone
        """
        try:
            template_analysis = template['template_analysis']
            template_path = template_analysis.get('template_path')
            filled_fields = {}
            
            # For each identified field, search for relevant data in Pinecone
            for field in template_analysis.get('identified_fields', []):
                field_id = field['field_id']
                field_type = field['field_type']
                semantic_role = field.get('semantic_role', 'general')
                contextual_keywords = field.get('contextual_keywords', [])
                description = field.get('description', '')
                
                # Backward compatibility: if no contextual_keywords, use suggestions
                if not contextual_keywords:
                    suggestions = field.get('suggestions', [])
                    contextual_keywords = suggestions
                
                # Generate contextual search queries
                search_queries = self._generate_contextual_search_queries(
                    field_type, 
                    semantic_role, 
                    contextual_keywords, 
                    description
                )
                
                # Use contextual search if semantic_role is available, otherwise fall back
                if semantic_role != 'general' and hasattr(self, '_search_person_data_contextual'):
                    field_value = await self._search_person_data_contextual(
                        person_id, 
                        search_queries, 
                        device_id,
                        field_type,
                        semantic_role
                    )
                else:
                    # Fallback to original search for compatibility
                    field_value = await self._search_person_data(
                        person_id, 
                        search_queries, 
                        device_id,
                        field_type
                    )
                
                filled_fields[field_id] = {
                    'value': field_value,
                    'field_info': field,
                    'found': bool(field_value)
                }
            
            # Generate filled document based on file type
            filename = f"filled_{batch_id}_{template_idx}_{template_analysis['template_info']['filename']}"
            file_path = self.output_dir / filename
            
            # Initialize skip tracking
            skip_info = {'skipped_fields': [], 'filled_fields_count': 0}
            
            if template_path and Path(template_path).suffix.lower() == '.docx':
                # Handle Word documents properly
                skip_info = await self._fill_word_document(
                    template_path,
                    file_path,
                    template_analysis,
                    filled_fields
                )
            else:
                # Handle other document types or create new document
                skip_info = await self._create_filled_document(
                    file_path,
                    template_analysis,
                    filled_fields,
                    template_idx
                )
            
            # Calculate statistics including skipped fields
            filled_count = sum(1 for f in filled_fields.values() if f['found'])
            total_identified_fields = len(filled_fields)
            skipped_fields = skip_info.get('skipped_fields', [])
            
            return {
                'success': True,
                'template_index': template_idx,
                'filename': filename,
                'file_path': str(file_path),
                'download_url': f"/api/download/interpreted/{filename}",
                'filled_fields': filled_fields,
                'fields_found': filled_count,
                'total_fields': total_identified_fields,
                'fields_skipped': len(skipped_fields),
                'skipped_field_summary': [
                    f"{skip['field_type']} field ({skip['field_id']}): {skip['reason']}"
                    for skip in skipped_fields
                ],
                'fill_rate': f"{filled_count}/{total_identified_fields}" if total_identified_fields > 0 else "0/0",
                'message': f"Successfully filled {filled_count} out of {total_identified_fields} fields" + 
                          (f" (skipped {len(skipped_fields)} fields with missing data)" if skipped_fields else "")
            }
            
        except Exception as e:
            logger.error(f"❌ Failed to fill single template: {e}")
            return {
                'success': False,
                'error': str(e),
                'template_index': template_idx
            }
    
    def _generate_search_queries(
        self, 
        field_type: str, 
        suggestions: List[str], 
        description: str
    ) -> List[str]:
        """
        Generate search queries for finding relevant person data
        """
        queries = []
        
        # Add suggestion-based queries
        for suggestion in suggestions:
            clean_suggestion = suggestion.replace('_', ' ').strip()
            if clean_suggestion:
                queries.append(clean_suggestion)
        
        # Add type-based queries with better coverage
        type_queries = {
            'name': ['name', 'full name', 'first name', 'last name', 'person name', 'applicant name'],
            'address': ['address', 'street address', 'home address', 'residence', 'location', 'plot', 'area'],
            'phone': ['phone number', 'telephone', 'mobile', 'contact number', 'phone'],
            'email': ['email address', 'email', 'e-mail', 'electronic mail'],
            'date': ['date', 'date of birth', 'birth date', 'birthday', 'DOB', 'birth'],
            'signature': ['signature', 'sign', 'signed by'],
            'age': ['age', 'years old'],
            'father': ['father name', 'father', 'parent'],
            'today': ['today date', 'current date', 'date'],
            'text': ['information', 'details', 'data']
        }
        
        # Get queries for this field type
        if field_type.lower() in type_queries:
            queries.extend(type_queries[field_type.lower()])
        
        # Also try to match common field type variations
        for field_variant in ['name', 'address', 'phone', 'email', 'date']:
            if field_variant in field_type.lower():
                queries.extend(type_queries.get(field_variant, []))
        
        # Add description-based queries
        if description:
            desc_words = description.lower().split()
            for word in desc_words:
                if len(word) > 3 and word not in ['the', 'and', 'for', 'with', 'this', 'that']:
                    queries.append(word)
        
        # Remove duplicates and empty queries
        unique_queries = []
        for q in queries:
            if q and q not in unique_queries:
                unique_queries.append(q)
        
        return unique_queries[:10]  # Limit to first 10 queries

    async def _search_person_data(
        self,
        person_id: str,
        search_queries: List[str],
        device_id: str,  # Keep for compatibility but not used for person search
        field_type: str
    ) -> Optional[str]:
        """
        Search for person data using the separate person storage with enhanced matching
        """
        try:
            logger.info(f"🔍 Enhanced search for {field_type} with queries: {search_queries}")
            best_match = None
            best_score = 0.0
            best_metadata = None
            
            for query in search_queries:
                try:
                    logger.info(f"🔍 Processing query: '{query}'")
                    
                    # Generate embedding for query
                    query_embedding = await self.gemini_service.get_embedding(query)
                    logger.info(f"✅ Generated embedding for query: '{query}'")
                    
                    # Search in person storage
                    results = await self._search_person_vectors(
                        person_id,
                        query_embedding,
                        top_k=10  # Get more results for better analysis
                    )
                    
                    logger.info(f"📊 Person storage returned {len(results)} results for query '{query}'")
                    
                    # Process results with enhanced scoring
                    for i, match in enumerate(results):
                        # Calculate enhanced score based on multiple factors
                        enhanced_score = self._calculate_enhanced_match_score(
                            match, query, field_type
                        )
                        
                        logger.info(f"📋 Result {i+1}: base_score={match['score']:.3f}, enhanced_score={enhanced_score:.3f}")
                        
                        if enhanced_score > best_score and enhanced_score > 0.15:  # Higher threshold
                            best_score = enhanced_score
                            best_metadata = match['metadata']
                            
                            # Extract value from metadata or text
                            if 'field_value' in match['metadata']:
                                best_match = match['metadata']['field_value']
                                logger.info(f"✅ Found enhanced field_value: '{best_match}'")
                            else:
                                # Extract value from text using improved extraction
                                best_match = await self._extract_value_from_text(
                                    match['text'], 
                                    field_type, 
                                    query
                                )
                                logger.info(f"🤖 Enhanced extracted value: '{best_match}' for query '{query}'")
                    
                except Exception as e:
                    logger.warning(f"❌ Enhanced query '{query}' failed: {e}")
                    continue
            
            # Post-process the value with validation
            if best_match and best_metadata:
                final_value = self._validate_and_format_field_value(
                    best_match, field_type, best_metadata
                )
                logger.info(f"✅ Final enhanced value for {field_type}: '{final_value}' (score: {best_score:.3f})")
                return final_value
            else:
                logger.warning(f"❌ No enhanced data found for {field_type} with person_id: {person_id}")
            
            return None
            
        except Exception as e:
            logger.error(f"❌ Failed enhanced person data search: {e}")
            return None
    
    def _calculate_enhanced_match_score(
        self, 
        match: Dict[str, Any], 
        query: str, 
        field_type: str
    ) -> float:
        """
        Calculate enhanced match score based on multiple factors
        """
        base_score = match.get('score', 0.0)
        metadata = match.get('metadata', {})
        text = match.get('text', '').lower()
        query_lower = query.lower()
        
        # Start with base embedding similarity
        enhanced_score = base_score
        
        # Boost for exact field type matches
        if metadata.get('field_type') == field_type:
            enhanced_score += 0.2
            logger.debug(f"🎯 Field type match bonus: +0.2")
        
        # Boost for chunk type quality
        chunk_type = metadata.get('chunk_type', '')
        chunk_type_boosts = {
            'structured_data': 0.15,
            'csv_data': 0.1,
            'summary': 0.05,
            'complete_summary': 0.1,
            'csv_field_summary': 0.08
        }
        if chunk_type in chunk_type_boosts:
            enhanced_score += chunk_type_boosts[chunk_type]
            logger.debug(f"🔥 Chunk type '{chunk_type}' bonus: +{chunk_type_boosts[chunk_type]}")
        
        # Boost for semantic role matches
        semantic_role = metadata.get('semantic_role', 'general')
        if semantic_role != 'general':
            enhanced_score += 0.1
            logger.debug(f"👤 Semantic role '{semantic_role}' bonus: +0.1")
        
        # Boost for direct query matches in text
        if query_lower in text:
            enhanced_score += 0.15
            logger.debug(f"🎪 Direct query match bonus: +0.15")
        
        # Boost for field name relevance
        field_name = metadata.get('field_name', '').lower()
        field_relevance_score = self._calculate_field_relevance(field_name, field_type)
        enhanced_score += field_relevance_score
        
        # Penalty for overly generic chunks
        if len(text) < 10 or 'general' in text:
            enhanced_score -= 0.05
            logger.debug(f"⚠️ Generic content penalty: -0.05")
        
        return min(enhanced_score, 1.0)  # Cap at 1.0
    
    def _calculate_field_relevance(self, field_name: str, target_field_type: str) -> float:
        """
        Calculate how relevant a field name is to the target field type
        """
        field_keywords = {
            'name': ['name', 'full', 'first', 'last', 'person'],
            'address': ['address', 'residence', 'location', 'street', 'area', 'plot'],
            'phone': ['phone', 'mobile', 'contact', 'telephone'],
            'email': ['email', 'mail'],
            'age': ['age', 'years', 'old'],
            'date': ['date', 'birth', 'dob', 'born'],
            'father': ['father', 'parent', 'guardian'],
            'occupation': ['occupation', 'profession', 'job', 'work']
        }
        
        if target_field_type not in field_keywords:
            return 0.0
        
        target_keywords = field_keywords[target_field_type]
        relevance_score = 0.0
        
        for keyword in target_keywords:
            if keyword in field_name:
                relevance_score += 0.05
        
        return min(relevance_score, 0.2)  # Cap at 0.2
    
    def _validate_and_format_field_value(
        self, 
        value: str, 
        field_type: str, 
        metadata: Dict[str, Any]
    ) -> str:
        """
        Validate and format field value with enhanced logic
        """
        if not value or not value.strip():
            return ""
        
        value = value.strip()
        
        # Additional validation based on field type
        if field_type == 'name':
            # Ensure it looks like a name (letters and spaces)
            import re
            if re.match(r'^[A-Za-z\s\.]+$', value) and len(value) > 1:
                return self._format_field_value(value, field_type)
            else:
                logger.warning(f"⚠️ Invalid name format: '{value}'")
                return ""
        
        elif field_type == 'phone':
            # Validate phone number format
            import re
            # Remove common formatting
            clean_phone = re.sub(r'[^\d\+]', '', value)
            if len(clean_phone) >= 10:
                return self._format_field_value(value, field_type)
            else:
                logger.warning(f"⚠️ Invalid phone format: '{value}'")
                return ""
        
        elif field_type == 'email':
            # Basic email validation
            import re
            if re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', value):
                return value.lower()
            else:
                logger.warning(f"⚠️ Invalid email format: '{value}'")
                return ""
        
        elif field_type == 'age':
            # Validate age is a reasonable number
            try:
                age_val = int(value)
                if 0 <= age_val <= 150:
                    return str(age_val)
                else:
                    logger.warning(f"⚠️ Invalid age value: '{value}'")
                    return ""
            except ValueError:
                logger.warning(f"⚠️ Non-numeric age: '{value}'")
                return ""
        
        # For other field types, use standard formatting
        return self._format_field_value(value, field_type)

    def _generate_contextual_search_queries(
        self, 
        field_type: str, 
        semantic_role: str,
        contextual_keywords: List[str], 
        description: str
    ) -> List[str]:
        """
        Generate contextual search queries that understand semantic roles with precise field type matching
        """
        queries = []
        
        # ENHANCED Role-specific queries with field type precision
        role_specific_queries = {
            'testator': {
                'name': ['testator name', 'deceased name', 'person name', 'full name', 'applicant name'],
                'age': ['testator age', 'person age', 'age', 'years old', 'aged'],
                'date': ['testator date of birth', 'birth date', 'date of birth', 'DOB'],
                'address': ['testator address', 'person address', 'residence', 'home address'],
            },
            'witness_a': {
                'name': ['witness A name', 'first witness name', 'witness name'],
                'age': ['witness A age', 'first witness age', 'witness age', 'aged about'],
                'date': ['witness A birth date', 'witness date of birth'],
                'address': ['witness A address', 'witness residence', 'witness home'],
            },
            'witness_b': {
                'name': ['witness B name', 'second witness name', 'witness name'],
                'age': ['witness B age', 'second witness age', 'witness age', 'aged about'],
                'date': ['witness B birth date', 'witness date of birth'],
                'address': ['witness B address', 'witness residence', 'witness home'],
            },
            'father_of_testator': {
                'name': ['testator father name', 'father name', 'parent name', 'son of', 'daughter of'],
                'age': ['father age', 'parent age'],
                'address': ['father address', 'parent address'],
            },
            'father_of_witness_a': {
                'name': ['witness A father name', 'witness father name', 'son of', 'daughter of'],
                'age': ['witness father age'],
            },
            'father_of_witness_b': {
                'name': ['witness B father name', 'witness father name', 'son of', 'daughter of'],
                'age': ['witness father age'],
            },
            'sub_registrar': {
                'name': ['sub registrar name', 'registrar name', 'sub-registrar'],
                'location': ['sub registrar office', 'registration office', 'registrar location'],
            },
            'general': {
                'name': ['name', 'full name', 'person name'],
                'age': ['age', 'years old', 'aged', 'years'],
                'date': ['date', 'birth date', 'date of birth'],
                'address': ['address', 'residence', 'home'],
                'location': ['location', 'place', 'at'],
            }
        }
        
        # Get role-specific queries with fallback to general
        if semantic_role in role_specific_queries and field_type in role_specific_queries[semantic_role]:
            queries.extend(role_specific_queries[semantic_role][field_type])
        elif field_type in role_specific_queries['general']:
            # Fallback to general queries for this field type
            queries.extend(role_specific_queries['general'][field_type])
        
        # Enhanced field type specific queries
        enhanced_field_queries = {
            'age': ['age', 'years old', 'aged', 'years', 'age of person', 'how old'],
            'name': ['name', 'full name', 'person name', 'individual name', 'called'],
            'address': ['address', 'residence', 'home', 'lives at', 'resident of', 'located at'],
            'date': ['date', 'birth date', 'date of birth', 'born on', 'DOB'],
            'phone': ['phone', 'mobile', 'contact', 'telephone', 'number'],
            'email': ['email', 'mail', 'electronic mail'],
            'location': ['location', 'place', 'at', 'where', 'venue'],
            'text': ['information', 'details', 'data']
        }
        
        # Add enhanced field-specific queries
        if field_type in enhanced_field_queries:
            for query in enhanced_field_queries[field_type]:
                if query not in queries:
                    queries.append(query)
        
        # Add contextual keywords (cleaned)
        for keyword in contextual_keywords:
            if keyword and keyword.strip():
                clean_keyword = keyword.replace('_', ' ').strip()
                if clean_keyword not in queries and len(clean_keyword) > 1:
                    queries.append(clean_keyword)
                
                # Combine keyword with field type for better precision
                if field_type != 'text':
                    combined = f"{clean_keyword} {field_type}".strip()
                    if combined not in queries:
                        queries.append(combined)
        
        # Add description-based queries (filtered)
        if description:
            desc_words = description.lower().split()
            relevant_words = []
            # Filter out common stop words
            stop_words = {'the', 'and', 'for', 'with', 'this', 'that', 'field', 'information', 'data'}
            for word in desc_words:
                if len(word) > 2 and word not in stop_words and word not in queries:
                    relevant_words.append(word)
            
            # Add the most relevant description words
            queries.extend(relevant_words[:3])  # Limit to top 3 relevant words
        
        # Remove duplicates while preserving order
        unique_queries = []
        for q in queries:
            if q and q.strip() and q not in unique_queries:
                unique_queries.append(q.strip())
        
        # Prioritize more specific queries first
        prioritized_queries = []
        
        # Add role+field specific queries first
        for q in unique_queries:
            if semantic_role != 'general' and (semantic_role.replace('_', ' ') in q or field_type in q):
                prioritized_queries.append(q)
        
        # Add field-specific queries
        for q in unique_queries:
            if q not in prioritized_queries and field_type in q:
                prioritized_queries.append(q)
        
        # Add remaining queries
        for q in unique_queries:
            if q not in prioritized_queries:
                prioritized_queries.append(q)
        
        return prioritized_queries[:8]  # Limit for performance but allow more for better coverage

    async def _search_person_data_contextual(
        self,
        person_id: str,
        search_queries: List[str],
        device_id: str,
        field_type: str,
        semantic_role: str
    ) -> Optional[str]:
        """
        Contextual search that considers semantic roles for better accuracy
        """
        try:
            logger.info(f"🎯 Contextual search for {semantic_role} {field_type} with queries: {search_queries}")
            best_match = None
            best_score = 0.0
            
            for query in search_queries:
                try:
                    logger.info(f"🔍 Processing contextual query: '{query}' for {semantic_role}")
                    
                    # Generate embedding for query
                    query_embedding = await self.gemini_service.get_embedding(query)
                    
                    # Search in person storage
                    results = await self._search_person_vectors(
                        person_id,
                        query_embedding,
                        top_k=5
                    )
                    
                    logger.info(f"📊 Found {len(results)} results for contextual query '{query}'")
                    
                    # Process results with semantic role consideration
                    for i, match in enumerate(results):
                        context_score = match['score']
                        
                        # Boost score if the result text contains role-specific keywords
                        text_content = match.get('text', '').lower()
                        role_boost = 0.0
                        
                        if semantic_role != 'general':
                            role_keywords = {
                                'testator': ['testator', 'deceased', 'deponent'],
                                'witness': ['witness', 'attesting', 'attestor'],
                                'applicant': ['applicant', 'petitioner'],
                                'father': ['father', 'paternal', 'parent']
                            }
                            
                            if semantic_role in role_keywords:
                                for keyword in role_keywords[semantic_role]:
                                    if keyword in text_content:
                                        role_boost = 0.1
                                        break
                        
                        adjusted_score = context_score + role_boost
                        
                        logger.info(f"📋 Result {i+1}: score={context_score:.3f}, adjusted={adjusted_score:.3f}, role={semantic_role}")
                        
                        if adjusted_score > best_score and adjusted_score > 0.1:
                            best_score = adjusted_score
                            
                            # Extract value using improved methods
                            if 'field_value' in match['metadata']:
                                best_match = match['metadata']['field_value']
                                logger.info(f"✅ Found contextual field_value: '{best_match}'")
                            else:
                                best_match = await self._extract_value_from_text(
                                    match['text'], 
                                    field_type, 
                                    query
                                )
                                logger.info(f"🎯 Contextually extracted: '{best_match}' for {semantic_role} {field_type}")
                    
                except Exception as e:
                    logger.warning(f"❌ Contextual query '{query}' failed: {e}")
                    continue
            
            # Format the final value
            if best_match:
                final_value = self._format_field_value(best_match, field_type)
                logger.info(f"✅ Final contextual value for {semantic_role} {field_type}: '{final_value}' (score: {best_score:.3f})")
                return final_value
            else:
                logger.warning(f"❌ No contextual data found for {semantic_role} {field_type} with person_id: {person_id}")
            
            return None
            
        except Exception as e:
            logger.error(f"❌ Failed contextual search: {e}")
            return None
    
    async def _extract_value_from_text(
        self, 
        text: str, 
        field_type: str, 
        query: str
    ) -> Optional[str]:
        """
        Use LLM to extract specific value from text
        """
        # Try regex-based extraction first for common patterns
        regex_result = self._extract_with_regex(text, field_type, query)
        if regex_result:
            logger.info(f"🎯 Regex extracted for {field_type}: '{regex_result}'")
            return regex_result
        
        # INTELLIGENT FORM FIX: Truncate text to prevent Google 500 errors
        # Limit context to prevent API token limits while preserving relevant content
        MAX_CONTEXT_LENGTH = 2000  # Safe limit for Google Gemini API
        
        truncated_text = text
        if len(text) > MAX_CONTEXT_LENGTH:
            # Smart truncation: Try to keep the most relevant part around the query
            query_lower = query.lower()
            text_lower = text.lower()
            
            # Find query position in text
            query_pos = text_lower.find(query_lower)
            if query_pos >= 0:
                # Extract context around the query
                start = max(0, query_pos - MAX_CONTEXT_LENGTH // 2)
                end = min(len(text), start + MAX_CONTEXT_LENGTH)
                truncated_text = text[start:end]
                logger.info(f"🔧 Context truncated around query '{query}': {len(text)} → {len(truncated_text)} chars")
            else:
                # Fallback: Take first part of text
                truncated_text = text[:MAX_CONTEXT_LENGTH]
                logger.info(f"🔧 Context truncated (query not found): {len(text)} → {len(truncated_text)} chars")
        
        # Fallback to LLM extraction with truncated text
        prompt = f"""
Extract ONLY the {field_type} information from this text.

Query: "{query}"
Text: {truncated_text}

Return ONLY the exact value, no explanation. Examples:
- For name query: "Arun Yadav" (not "Name: Arun Yadav")
- For address query: "Plot No 45, Industrial Area Phase II, Panchkula, Haryana"
- For phone query: "+91-9876543210"
- For email query: "arun@example.com"
- For date query: "25/08/1990"

Extract only the {field_type} value:"""

        try:
            response = await self.gemini_service.generate_response(
                prompt=prompt,
                max_tokens=100,
                temperature=0.0
            )
            
            extracted_value = response.strip()
            # Clean up common LLM artifacts
            extracted_value = extracted_value.replace('"', '').replace("'", '').strip()
            
            if extracted_value and extracted_value.lower() not in ["not found", "none", "n/a", "not available"]:
                logger.info(f"🤖 LLM extracted for {field_type}: '{extracted_value}'")
                return extracted_value
            
        except Exception as e:
            logger.error(f"❌ LLM extraction failed: {e}")
        
        return None

    def _extract_with_regex(self, text: str, field_type: str, query: str) -> Optional[str]:
        """
        Extract values using regex patterns for common field types
        """
        import re
        
        try:
            # Name extraction
            if field_type.lower() in ['name', 'full_name', 'first_name', 'last_name']:
                # Look for "Name: value" pattern (stop at Father or next field)
                name_match = re.search(r'Name:\s*([A-Za-z\s]+?)(?:\s+Father|\s+Address|$|\n)', text, re.IGNORECASE)
                if name_match:
                    return name_match.group(1).strip()
                
                # Look for "Full Name: value" pattern
                name_match = re.search(r'Full\s*Name:\s*([A-Za-z\s]+?)(?:\s+|$|\n)', text, re.IGNORECASE)
                if name_match:
                    return name_match.group(1).strip()
            
            # Address extraction - more precise
            elif field_type.lower() in ['address', 'street_address', 'home_address']:
                # Look for "Address: value" pattern (stop at Age)
                addr_match = re.search(r'Address:\s*(.*?)(?:\s+Age)', text, re.IGNORECASE)
                if addr_match:
                    address = addr_match.group(1).strip()
                    return address
                
                # Fallback pattern for plot/area style addresses
                plot_match = re.search(r'Plot\s+No\s+\d+[^A]+?(?=\s+Age|\s+Phone|\s+Email|$)', text, re.IGNORECASE)
                if plot_match:
                    return plot_match.group().strip()
            
            # Phone extraction
            elif field_type.lower() in ['phone', 'telephone', 'mobile', 'contact']:
                # Look for phone patterns
                phone_match = re.search(r'(?:Phone|Mobile|Contact):\s*([+]?[\d\s\-\(\)]+)', text, re.IGNORECASE)
                if phone_match:
                    return phone_match.group(1).strip()
                
                # Look for standalone phone numbers
                phone_match = re.search(r'[+]?[\d]{10,}', text)
                if phone_match:
                    return phone_match.group().strip()
            
            # Date extraction - more precise
            elif field_type.lower() in ['date', 'birth_date', 'dob', 'date_of_birth']:
                # Look for "Date of Birth: value" pattern
                date_match = re.search(r'Date\s*of\s*Birth:\s*(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{4})', text, re.IGNORECASE)
                if date_match:
                    return date_match.group(1).strip()
                
                # Look for standalone dates in DD/MM/YYYY format
                date_match = re.search(r'\b(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{4})\b', text)
                if date_match:
                    return date_match.group(1).strip()
            
            # Age extraction
            elif field_type.lower() in ['age']:
                age_match = re.search(r'Age:\s*(\d+)', text, re.IGNORECASE)
                if age_match:
                    return age_match.group(1).strip()
            
            # Father name extraction
            elif field_type.lower() in ['father', 'father_name', 'parent']:
                father_match = re.search(r'Father\s+Name:\s*([A-Za-z\s]+?)(?:\s+Address|\s+Age|$|\n)', text, re.IGNORECASE)
                if father_match:
                    return father_match.group(1).strip()
            
            # Email extraction
            elif field_type.lower() in ['email', 'e-mail', 'email_address']:
                email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
                if email_match:
                    return email_match.group().strip()
            
        except Exception as e:
            logger.warning(f"❌ Regex extraction failed for {field_type}: {e}")
        
        return None
    
    def _format_field_value(self, value: str, field_type: str) -> str:
        """
        Format field value according to field type
        """
        if not value:
            return ""
        
        value = value.strip()
        
        if field_type == 'date':
            # Try to standardize date format
            try:
                from datetime import datetime
                # Try common date formats
                for fmt in ['%Y-%m-%d', '%m/%d/%Y', '%d/%m/%Y', '%Y%m%d']:
                    try:
                        date_obj = datetime.strptime(value, fmt)
                        return date_obj.strftime('%Y-%m-%d')
                    except ValueError:
                        continue
            except:
                pass
        
        elif field_type == 'phone':
            # Clean phone number
            import re
            digits = re.sub(r'\D', '', value)
            if len(digits) == 10:
                return f"({digits[:3]}) {digits[3:6]}-{digits[6:]}"
            elif len(digits) == 11 and digits[0] == '1':
                return f"({digits[1:4]}) {digits[4:7]}-{digits[7:]}"
        
        elif field_type == 'name':
            # Capitalize names properly
            return value.title()
        
        return value
    
    async def _generate_document_content(
        self,
        template_analysis: Dict[str, Any],
        filled_fields: Dict[str, Any],
        template_idx: int
    ) -> str:
        """
        Generate the final filled document content
        """
        # INTELLIGENT FORM FIX: Limit field data to prevent Google 500 errors
        MAX_FIELD_DATA_LENGTH = 1500  # Safe limit for field data
        
        # Prepare simplified field data to prevent API token limits
        field_values = {k: v['value'] for k, v in filled_fields.items()}
        field_details = {}
        
        # Truncate field details if too large
        field_details_str = json.dumps({k: v['field_info'] for k, v in filled_fields.items()}, indent=2)
        if len(field_details_str) > MAX_FIELD_DATA_LENGTH:
            # Use simplified field details
            for k, v in filled_fields.items():
                field_info = v['field_info']
                field_details[k] = {
                    'field_type': field_info.get('field_type', 'text'),
                    'required': field_info.get('required', False)
                }
            logger.info(f"🔧 Field details truncated for document generation: {len(field_details_str)} → simplified")
        else:
            field_details = {k: v['field_info'] for k, v in filled_fields.items()}
        
        prompt = f"""
You are a professional document formatter. Create a complete, well-formatted document using the template analysis and filled field data provided.

Template Information:
{json.dumps(template_analysis['template_info'], indent=2)}

Identified Fields and Values:
{json.dumps(field_values, indent=2)}

Field Details:
{json.dumps(field_details, indent=2)}

Instructions:
1. Create a professional, complete document
2. Replace all identified blank spaces with the appropriate filled values
3. Maintain proper formatting and structure
4. Include all necessary legal language and clauses
5. Ensure the document flows naturally and reads professionally
6. If a field value is missing, use "[FIELD_NAME_TO_BE_FILLED]" as placeholder

Generate the complete filled document:
"""
        
        try:
            response = await self.gemini_service.generate_response(
                prompt=prompt,
                max_tokens=3000,
                temperature=0.2
            )
            
            # Add header with metadata
            header = f"""
Generated Filled Document #{template_idx + 1}
Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Template: {template_analysis['template_info']['filename']}
Fields filled: {sum(1 for f in filled_fields.values() if f['found'])}/{len(filled_fields)}

{'-' * 80}

"""
            
            return header + response
            
        except Exception as e:
            logger.error(f"❌ Failed to generate document content: {e}")
            # Return fallback content
            return f"""
Generated Filled Document #{template_idx + 1}
Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Template: {template_analysis['template_info']['filename']}

ERROR: Could not generate document content due to: {str(e)}

Field Values:
{json.dumps({k: v['value'] for k, v in filled_fields.items()}, indent=2)}
"""
    
    async def _create_batch_download(
        self,
        filled_documents: List[Dict[str, Any]],
        batch_id: str
    ) -> str:
        """
        Create a zip file containing all filled documents
        """
        try:
            zip_filename = f"batch_{batch_id}_filled_documents.zip"
            zip_path = self.output_dir / zip_filename
            
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for doc in filled_documents:
                    if doc['success']:
                        zipf.write(doc['file_path'], doc['filename'])
            
            return f"/api/download/interpreted/{zip_filename}"
            
        except Exception as e:
            logger.error(f"❌ Failed to create batch download: {e}")
            return ""
    
    def _json_to_chunks(self, data: Dict[str, Any], person_id: str) -> List[Dict[str, Any]]:
        """
        Convert JSON data to high-quality text chunks for embedding with enhanced context
        """
        chunks = []
        
        def process_value(key: str, value: Any, prefix: str = "", parent_context: str = ""):
            full_key = f"{prefix}.{key}" if prefix else key
            current_context = f"{parent_context} {key}" if parent_context else key
            
            if isinstance(value, dict):
                for sub_key, sub_value in value.items():
                    process_value(sub_key, sub_value, full_key, current_context)
            elif isinstance(value, list):
                for i, item in enumerate(value):
                    process_value(f"{key}[{i}]", item, prefix, current_context)
            else:
                # Create enhanced chunks with better semantic understanding
                value_str = str(value).strip()
                if not value_str or len(value_str) < 2:  # Skip empty or very short values
                    return
                
                # Determine field type and semantic role
                field_type, semantic_role = self._classify_json_field(key, value_str, current_context)
                
                # Create contextual text representations
                contextual_texts = self._create_contextual_representations(key, value_str, field_type, semantic_role)
                
                # Create multiple chunks for better searchability
                for i, text_repr in enumerate(contextual_texts):
                    chunk = {
                        'text': text_repr,
                        'metadata': {
                            'person_id': person_id,
                            'field_name': full_key,
                            'field_value': value_str,
                            'field_type': field_type,
                            'semantic_role': semantic_role,
                            'chunk_type': 'personal_data',
                            'chunk_variant': i,
                            'source_context': current_context
                        }
                    }
                    chunks.append(chunk)
        
        for key, value in data.items():
            process_value(key, value)
        
        # Add comprehensive summary chunks
        chunks.extend(self._create_summary_chunks(data, person_id))
        
        logger.info(f"📊 Created {len(chunks)} high-quality JSON chunks for person {person_id}")
        return chunks
    
    def _classify_json_field(self, key: str, value: str, context: str) -> tuple[str, str]:
        """
        Enhanced field classification for legal documents and affidavits
        """
        key_lower = key.lower()
        value_lower = value.lower()
        context_lower = context.lower()
        
        # Determine field type with enhanced legal document awareness
        field_type = "text"
        
        # Name field detection
        if any(term in key_lower for term in ['name', 'full_name', 'first_name', 'last_name', 'father_name', 'mother_name', 'spouse_name']):
            field_type = 'name'
        # Address field detection 
        elif any(term in key_lower for term in ['address', 'street', 'residence', 'location', 'area', 'plot', 'house', 'flat', 'sector']):
            field_type = 'address'
        # Phone field detection
        elif any(term in key_lower for term in ['phone', 'mobile', 'telephone', 'contact']):
            field_type = 'phone'
        # Email field detection
        elif any(term in key_lower for term in ['email', 'mail']):
            field_type = 'email'
        # Age field detection
        elif any(term in key_lower for term in ['age', 'years']) and value.strip().isdigit():
            field_type = 'age'
        # Date field detection (enhanced for legal documents)
        elif any(term in key_lower for term in ['date', 'birth', 'dob', 'execution_date', 'verification_date']) or self._is_date_pattern(value):
            field_type = 'date'
        # Location field detection (for legal venues)
        elif any(term in key_lower for term in ['location', 'office', 'registrar', 'court', 'venue']):
            field_type = 'location'
        # Occupation field detection
        elif any(term in key_lower for term in ['occupation', 'profession', 'job']):
            field_type = 'occupation'
        # Document type field
        elif any(term in key_lower for term in ['document_type', 'type']):
            field_type = 'document_type'
        
        # Determine semantic role with enhanced legal context awareness
        semantic_role = "general"
        
        # Check for testator context (main person in will)
        if any(term in context_lower for term in ['testator', 'testator_info']):
            semantic_role = 'testator'
        elif any(term in key_lower for term in ['testator']):
            semantic_role = 'testator'
        
        # Check for witness context
        elif any(term in context_lower for term in ['witness_a', 'witness a']):
            semantic_role = 'witness_a'
        elif any(term in context_lower for term in ['witness_b', 'witness b']):
            semantic_role = 'witness_b'
        elif any(term in key_lower for term in ['witness_a']):
            semantic_role = 'witness_a'
        elif any(term in key_lower for term in ['witness_b']):
            semantic_role = 'witness_b'
        elif any(term in key_lower for term in ['witness', 'attestor']):
            semantic_role = 'witness'
        
        # Check for family relations
        elif any(term in key_lower for term in ['father_name', 'father']):
            if 'testator' in context_lower:
                semantic_role = 'father_of_testator'
            elif 'witness_a' in context_lower:
                semantic_role = 'father_of_witness_a'
            elif 'witness_b' in context_lower:
                semantic_role = 'father_of_witness_b'
            else:
                semantic_role = 'father'
        elif any(term in key_lower for term in ['mother']):
            semantic_role = 'mother'
        elif any(term in key_lower for term in ['spouse', 'wife', 'husband']):
            semantic_role = 'spouse'
        
        # Check for legal document context
        elif any(term in context_lower for term in ['legal_context', 'legal']):
            if 'location' in key_lower or 'office' in key_lower or 'registrar' in key_lower:
                semantic_role = 'sub_registrar'
            elif 'date' in key_lower:
                semantic_role = 'legal_date'
            else:
                semantic_role = 'legal_context'
        
        # Check for registration office context
        elif any(term in key_lower for term in ['registrar', 'sub_registrar', 'office']):
            semantic_role = 'sub_registrar'
        elif any(term in key_lower for term in ['registration_location', 'verification_location']):
            semantic_role = 'sub_registrar'
        
        return field_type, semantic_role
    
    def _is_date_pattern(self, value: str) -> bool:
        """Check if value matches common date patterns"""
        import re
        date_patterns = [
            r'\d{1,2}[\/\-]\d{1,2}[\/\-]\d{4}',  # DD/MM/YYYY or MM/DD/YYYY
            r'\d{4}[\/\-]\d{1,2}[\/\-]\d{1,2}',  # YYYY/MM/DD
            r'\d{1,2}\s+(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}',  # DD Mon YYYY
        ]
        
        for pattern in date_patterns:
            if re.search(pattern, value, re.IGNORECASE):
                return True
        return False
    
    def _create_contextual_representations(self, key: str, value: str, field_type: str, semantic_role: str) -> List[str]:
        """
        Create multiple contextual text representations for better search
        """
        representations = []
        
        # Clean field name for display
        clean_key = key.replace('_', ' ').replace('.', ' ').title()
        
        # Base representation
        representations.append(f"{clean_key}: {value}")
        
        # Field type specific representations
        if field_type == 'name':
            representations.extend([
                f"Name is {value}",
                f"Full name: {value}",
                f"Person name: {value}",
                f"{value} is the name"
            ])
            
            if semantic_role == 'father':
                representations.extend([
                    f"Father name: {value}",
                    f"Father's name is {value}",
                    f"Son of {value}",
                    f"Daughter of {value}",
                    f"Parent name: {value}"
                ])
            elif semantic_role == 'testator':
                representations.extend([
                    f"Testator name: {value}",
                    f"Deceased name: {value}",
                    f"Will executed by {value}"
                ])
            elif semantic_role == 'witness':
                representations.extend([
                    f"Witness name: {value}",
                    f"Attesting witness: {value}"
                ])
        
        elif field_type == 'address':
            representations.extend([
                f"Address: {value}",
                f"Residence: {value}",
                f"Lives at {value}",
                f"Resident of {value}",
                f"Home address: {value}",
                f"Street address: {value}"
            ])
        
        elif field_type == 'phone':
            representations.extend([
                f"Phone number: {value}",
                f"Mobile number: {value}",
                f"Contact number: {value}",
                f"Telephone: {value}"
            ])
        
        elif field_type == 'email':
            representations.extend([
                f"Email address: {value}",
                f"Email: {value}",
                f"Electronic mail: {value}"
            ])
        
        elif field_type == 'age':
            representations.extend([
                f"Age: {value}",
                f"Age is {value}",
                f"{value} years old",
                f"Aged {value}",
                f"Aged about {value}"
            ])
        
        elif field_type == 'date':
            representations.extend([
                f"Date: {value}",
                f"Date of birth: {value}",
                f"Birth date: {value}",
                f"Born on {value}",
                f"DOB: {value}"
            ])
        
        elif field_type == 'occupation':
            representations.extend([
                f"Occupation: {value}",
                f"Profession: {value}",
                f"Job: {value}",
                f"Works as {value}"
            ])
        
        # Add semantic role context
        if semantic_role != 'general':
            role_context = f"{semantic_role.title()} {field_type}: {value}"
            if role_context not in representations:
                representations.append(role_context)
        
        return representations
    
    def _create_summary_chunks(self, data: Dict[str, Any], person_id: str) -> List[Dict[str, Any]]:
        """
        Create comprehensive summary chunks for better document filling
        """
        summary_chunks = []
        
        # Extract key information
        name_fields = []
        address_fields = []
        contact_fields = []
        family_fields = []
        personal_fields = []
        
        def extract_fields(obj, prefix=""):
            if isinstance(obj, dict):
                for key, value in obj.items():
                    full_key = f"{prefix}.{key}" if prefix else key
                    key_lower = key.lower()
                    
                    if any(term in key_lower for term in ['name']) and 'father' not in key_lower:
                        name_fields.append(f"{key}: {value}")
                    elif any(term in key_lower for term in ['address', 'residence', 'location']):
                        address_fields.append(f"{key}: {value}")
                    elif any(term in key_lower for term in ['phone', 'email', 'contact']):
                        contact_fields.append(f"{key}: {value}")
                    elif any(term in key_lower for term in ['father', 'mother', 'spouse', 'parent']):
                        family_fields.append(f"{key}: {value}")
                    elif any(term in key_lower for term in ['age', 'date', 'birth', 'occupation']):
                        personal_fields.append(f"{key}: {value}")
                    
                    # Recursively process nested objects
                    if isinstance(value, (dict, list)):
                        extract_fields(value, full_key)
            elif isinstance(obj, list):
                for i, item in enumerate(obj):
                    extract_fields(item, f"{prefix}[{i}]")
        
        extract_fields(data)
        
        # Create category-based summary chunks
        if name_fields:
            summary_chunks.append({
                'text': f"Person name information: {'; '.join(name_fields)}",
                'metadata': {
                    'person_id': person_id,
                    'chunk_type': 'summary',
                    'summary_category': 'names',
                    'field_count': len(name_fields)
                }
            })
        
        if address_fields:
            summary_chunks.append({
                'text': f"Address and location information: {'; '.join(address_fields)}",
                'metadata': {
                    'person_id': person_id,
                    'chunk_type': 'summary',
                    'summary_category': 'address',
                    'field_count': len(address_fields)
                }
            })
        
        if contact_fields:
            summary_chunks.append({
                'text': f"Contact information: {'; '.join(contact_fields)}",
                'metadata': {
                    'person_id': person_id,
                    'chunk_type': 'summary',
                    'summary_category': 'contact',
                    'field_count': len(contact_fields)
                }
            })
        
        if family_fields:
            summary_chunks.append({
                'text': f"Family information: {'; '.join(family_fields)}",
                'metadata': {
                    'person_id': person_id,
                    'chunk_type': 'summary',
                    'summary_category': 'family',
                    'field_count': len(family_fields)
                }
            })
        
        if personal_fields:
            summary_chunks.append({
                'text': f"Personal details: {'; '.join(personal_fields)}",
                'metadata': {
                    'person_id': person_id,
                    'chunk_type': 'summary',
                    'summary_category': 'personal',
                    'field_count': len(personal_fields)
                }
            })
        
        # Create a comprehensive overview chunk
        all_fields = name_fields + address_fields + contact_fields + family_fields + personal_fields
        if all_fields:
            summary_chunks.append({
                'text': f"Complete person profile: {'; '.join(all_fields)}",
                'metadata': {
                    'person_id': person_id,
                    'chunk_type': 'complete_summary',
                    'summary_category': 'overview',
                    'field_count': len(all_fields)
                }
            })
        
        return summary_chunks
    
    def _text_to_chunks(
        self, 
        text: str, 
        person_id: str, 
        device_id: str, 
        filename: str
    ) -> List[Dict[str, Any]]:
        """
        Convert text to high-quality chunks for embedding with intelligent parsing
        """
        chunks = []
        
        # First, try to parse structured data from text
        structured_chunks = self._parse_structured_text(text, person_id, device_id, filename)
        if structured_chunks:
            chunks.extend(structured_chunks)
            logger.info(f"📋 Extracted {len(structured_chunks)} structured chunks from text")
        
        # Then create semantic chunks from remaining text
        semantic_chunks = self._create_semantic_text_chunks(text, person_id, device_id, filename)
        chunks.extend(semantic_chunks)
        
        # Add contextual chunks for better search
        contextual_chunks = self._create_contextual_text_chunks(text, person_id, device_id, filename)
        chunks.extend(contextual_chunks)
        
        logger.info(f"📊 Created {len(chunks)} total high-quality text chunks for person {person_id}")
        return chunks
    
    def _parse_structured_text(
        self, 
        text: str, 
        person_id: str, 
        device_id: str, 
        filename: str
    ) -> List[Dict[str, Any]]:
        """
        Parse structured information from text using patterns
        """
        import re
        chunks = []
        
        # Define patterns for common personal data fields
        patterns = {
            'name': [
                r'(?:Name|Full\s*Name|Person\s*Name):\s*([A-Za-z\s]+?)(?:\n|$|[A-Z][a-z]+:)',
                # Enhanced pattern for Indian names - more context-aware
                r'(?:Shri|Mr\.?|Ms\.?|Mrs\.?)\s+([A-Za-z\s]+?)(?:\s+son\s+of|\s+daughter\s+of|\s+aged|\s+resident|\s+,)',
                r'I,\s+([A-Za-z\s]+?)(?:,|\s+son|\s+daughter|\s+aged)',
                # Pattern for "executed by Shri [Name]" - capture only actual names
                r'executed\s+by\s+(?:Shri|Mr\.?|Ms\.?|Mrs\.?)\s+([A-Za-z\s]+?)(?:\s+son\s+of|\s+on|\s+,|\s+\.)',
            ],
            'father_name': [
                r'(?:Father|Father\'s\s*Name|Parent):\s*([A-Za-z\s]+?)(?:\n|$|[A-Z][a-z]+:)',
                r'(?:son|daughter)\s+of\s+(?:Shri|Mr\.?|late)?\s*([A-Za-z\s]+?)(?:\n|,|$|\s+resident)',
            ],
            'address': [
                r'(?:Address|Residence|Home\s*Address):\s*([^:]+?)(?:\n[A-Z][a-z]+:|$)',
                r'(?:resident\s+of|residing\s+at|living\s+at)\s+([^:,]+?)(?:\n|,|$)',
                r'(Plot\s+No\.?\s*\d+[^:]+?)(?:\n|$)',
            ],
            'age': [
                r'(?:Age|Aged):\s*(\d+)',
                r'aged\s+about\s+(\d+)',
                r'(\d+)\s+years?\s+old',
            ],
            'date_of_birth': [
                r'(?:Date\s+of\s+Birth|DOB|Birth\s+Date):\s*(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{4})',
                r'born\s+on\s+(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{4})',
            ],
            'phone': [
                r'(?:Phone|Mobile|Contact|Telephone):\s*([\+\d\s\-\(\)]+)',
                r'(?:Ph\.?|Mob\.?):\s*([\+\d\s\-\(\)]+)',
            ],
            'email': [
                r'(?:Email|E-mail):\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
                r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
            ],
            'occupation': [
                r'(?:Occupation|Profession|Job|Work):\s*([^:]+?)(?:\n|$)',
                r'(?:working\s+as|employed\s+as)\s+([^:,]+?)(?:\n|,|$)',
            ]
        }
        
        # Extract information using patterns
        for field_type, field_patterns in patterns.items():
            for pattern in field_patterns:
                matches = re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE)
                for match in matches:
                    extracted_value = match.group(1).strip()
                    if len(extracted_value) > 1 and len(extracted_value) < 200:  # Reasonable length
                        
                        # Enhanced validation for name fields to filter out place names
                        if field_type == 'name' and self._is_likely_place_name(extracted_value):
                            continue  # Skip this match as it's likely a place name
                        
                        # Determine semantic role
                        semantic_role = self._determine_semantic_role_from_context(
                            match.group(0), text, field_type
                        )
                        
                        # Create multiple representations
                        representations = self._create_contextual_representations(
                            field_type, extracted_value, field_type, semantic_role
                        )
                        
                        for i, text_repr in enumerate(representations):
                            chunk = {
                                'text': text_repr,
                                'metadata': {
                                    'person_id': person_id,
                                    'device_id': device_id,
                                    'source_file': filename,
                                    'field_name': field_type,
                                    'field_value': extracted_value,
                                    'field_type': field_type,
                                    'semantic_role': semantic_role,
                                    'chunk_type': 'structured_data',
                                    'extraction_method': 'pattern_matching',
                                    'chunk_variant': i
                                }
                            }
                            chunks.append(chunk)
        
        return chunks
    
    def _is_likely_place_name(self, name: str) -> bool:
        """
        Check if the extracted name is likely a place name rather than a person name
        """
        name_lower = name.lower().strip()
        
        # Common Indian city/place names that might be mistaken for person names
        indian_place_names = {
            'panchkula', 'chandigarh', 'delhi', 'mumbai', 'bangalore', 'chennai', 'kolkata',
            'hyderabad', 'pune', 'ahmedabad', 'surat', 'jaipur', 'lucknow', 'kanpur',
            'nagpur', 'indore', 'thane', 'bhopal', 'visakhapatnam', 'pimpri', 'patna',
            'vadodara', 'ghaziabad', 'ludhiana', 'agra', 'nashik', 'faridabad', 'meerut',
            'rajkot', 'kalyan', 'vasai', 'varanasi', 'srinagar', 'aurangabad', 'dhanbad',
            'amritsar', 'navi mumbai', 'allahabad', 'ranchi', 'howrah', 'coimbatore',
            'jabalpur', 'gwalior', 'vijayawada', 'jodhpur', 'madurai', 'raipur',
            'kota', 'guwahati', 'chandigarh', 'solapur', 'hubli', 'bareilly', 'moradabad',
            'mysore', 'gurgaon', 'aligarh', 'jalandhar', 'tiruchirappalli', 'bhubaneswar'
        }
        
        # Check against known place names
        if name_lower in indian_place_names:
            return True
        
        # Check for place-like patterns
        place_patterns = [
            r'\bplot\s+no\.?\s*\d+',  # Plot No 45
            r'\bblock\s+[a-z]\b',     # Block A
            r'\bsector\s+\d+',        # Sector 15
            r'\bphase\s+\d+',         # Phase 2
            r'\bcolony\b',            # Colony
            r'\bnagar\b',             # Nagar
            r'\bchowk\b',             # Chowk
            r'\bpuram\b',             # Puram
            r'\babad\b',              # Abad (like Ahmedabad)
            r'\bpur\b$',              # Pur (like Raipur)
        ]
        
        import re
        for pattern in place_patterns:
            if re.search(pattern, name_lower):
                return True
        
        return False

    def _determine_semantic_role_from_context(self, match_text: str, full_text: str, field_type: str) -> str:
        """
        Determine semantic role based on context around the match
        """
        match_lower = match_text.lower()
        full_text_lower = full_text.lower()
        
        # Find position of match in full text
        match_pos = full_text_lower.find(match_lower)
        if match_pos == -1:
            return 'general'
        
        # Get context around the match (200 chars before and after)
        start = max(0, match_pos - 200)
        end = min(len(full_text), match_pos + len(match_text) + 200)
        context = full_text_lower[start:end]
        
        # Determine role based on context
        if any(term in context for term in ['father', 'parent', 'son of', 'daughter of']):
            return 'father'
        elif any(term in context for term in ['mother']):
            return 'mother'
        elif any(term in context for term in ['spouse', 'wife', 'husband']):
            return 'spouse'
        elif any(term in context for term in ['testator', 'deceased', 'will']):
            return 'testator'
        elif any(term in context for term in ['witness', 'attestor', 'attesting']):
            return 'witness'
        elif any(term in context for term in ['applicant', 'petitioner']):
            return 'applicant'
        
        return 'general'
    
    def _create_semantic_text_chunks(
        self, 
        text: str, 
        person_id: str, 
        device_id: str, 
        filename: str
    ) -> List[Dict[str, Any]]:
        """
        Create semantic chunks from text with better context preservation
        """
        chunks = []
        
        # Split into sentences first
        import re
        sentences = re.split(r'[.!?]+', text)
        
        current_chunk = ""
        chunk_size = 300  # characters - smaller for better precision
        overlap_size = 50  # characters for context overlap
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            # If adding this sentence would exceed chunk size
            if len(current_chunk) + len(sentence) > chunk_size and current_chunk:
                # Save current chunk
                chunk_text = current_chunk.strip()
                if len(chunk_text) > 20:  # Only save meaningful chunks
                    chunk = {
                        'text': chunk_text,
                        'metadata': {
                            'person_id': person_id,
                            'device_id': device_id,
                            'source_file': filename,
                            'chunk_type': 'semantic_text',
                            'chunk_size': len(chunk_text)
                        }
                    }
                    chunks.append(chunk)
                
                # Start new chunk with overlap from previous chunk
                if len(current_chunk) > overlap_size:
                    overlap_text = current_chunk[-overlap_size:]
                    current_chunk = overlap_text + " " + sentence
                else:
                    current_chunk = sentence
            else:
                current_chunk += " " + sentence if current_chunk else sentence
        
        # Add final chunk
        if current_chunk.strip() and len(current_chunk.strip()) > 20:
            chunk = {
                'text': current_chunk.strip(),
                'metadata': {
                    'person_id': person_id,
                    'device_id': device_id,
                    'source_file': filename,
                    'chunk_type': 'semantic_text',
                    'chunk_size': len(current_chunk.strip())
                }
            }
            chunks.append(chunk)
        
        return chunks
    
    def _create_contextual_text_chunks(
        self, 
        text: str, 
        person_id: str, 
        device_id: str, 
        filename: str
    ) -> List[Dict[str, Any]]:
        """
        Create contextual chunks for specific search scenarios
        """
        chunks = []
        
        # Create chunks for common legal document sections
        sections = {
            'personal_info': r'(I,\s+[^,]+,[^.]+\.)',
            'address_info': r'(resident\s+of[^.]+\.)',
            'family_info': r'(son\s+of|daughter\s+of|father[^.]+\.)',
            'declaration': r'(do\s+hereby[^.]+\.)',
            'verification': r'(verified\s+at[^.]+\.)',
        }
        
        import re
        for section_type, pattern in sections.items():
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                section_text = match.group(1).strip()
                if len(section_text) > 10:
                    chunk = {
                        'text': section_text,
                        'metadata': {
                            'person_id': person_id,
                            'device_id': device_id,
                            'source_file': filename,
                            'chunk_type': 'contextual_section',
                            'section_type': section_type
                        }
                    }
                    chunks.append(chunk)
        
        # Create question-answer style chunks for better search
        qa_chunks = [
            f"What is the person's name? Based on the document: {text[:200]}...",
            f"What is the person's address? Based on the document: {text[:200]}...",
            f"Who is the person's father? Based on the document: {text[:200]}...",
            f"What is the person's age? Based on the document: {text[:200]}...",
        ]
        
        for qa_text in qa_chunks:
            chunk = {
                'text': qa_text,
                'metadata': {
                    'person_id': person_id,
                    'device_id': device_id,
                    'source_file': filename,
                    'chunk_type': 'qa_style',
                }
            }
            chunks.append(chunk)
        
        return chunks
    
    def _process_csv_row(
        self, 
        row: Dict[str, str], 
        row_idx: int, 
        person_id: str, 
        device_id: str, 
        filename: str
    ) -> List[Dict[str, Any]]:
        """
        Process a single CSV row to create high-quality chunks
        """
        chunks = []
        
        for key, value in row.items():
            if value and value.strip():
                value = value.strip()
                
                # Classify the field
                field_type, semantic_role = self._classify_json_field(key, value, "")
                
                # Create contextual representations
                representations = self._create_contextual_representations(key, value, field_type, semantic_role)
                
                for i, text_repr in enumerate(representations):
                    chunk = {
                        'text': text_repr,
                        'metadata': {
                            'person_id': person_id,
                            'device_id': device_id,
                            'source_file': filename,
                            'field_name': key,
                            'field_value': value,
                            'field_type': field_type,
                            'semantic_role': semantic_role,
                            'row_index': row_idx,
                            'chunk_type': 'csv_data',
                            'chunk_variant': i
                        }
                    }
                    chunks.append(chunk)
        
        # Create a row summary chunk
        if row:
            valid_fields = {k: v for k, v in row.items() if v and v.strip()}
            if valid_fields:
                row_summary = "; ".join([f"{k}: {v}" for k, v in valid_fields.items()])
                chunk = {
                    'text': f"Person record {row_idx + 1}: {row_summary}",
                    'metadata': {
                        'person_id': person_id,
                        'device_id': device_id,
                        'source_file': filename,
                        'row_index': row_idx,
                        'chunk_type': 'csv_row_summary',
                        'field_count': len(valid_fields)
                    }
                }
                chunks.append(chunk)
        
        return chunks
    
    def _create_csv_summary_chunks(
        self, 
        rows: List[Dict[str, str]], 
        person_id: str, 
        device_id: str, 
        filename: str
    ) -> List[Dict[str, Any]]:
        """
        Create summary chunks from CSV data for better search
        """
        chunks = []
        
        if not rows:
            return chunks
        
        # Get all unique column names
        all_columns = set()
        for row in rows:
            all_columns.update(row.keys())
        
        # Create field-specific summaries
        for column in all_columns:
            values = [row.get(column, '').strip() for row in rows if row.get(column, '').strip()]
            if values:
                # Determine field type
                field_type, semantic_role = self._classify_json_field(column, values[0], "")
                
                # Create summary text
                if len(values) == 1:
                    summary_text = f"{column}: {values[0]}"
                else:
                    summary_text = f"{column} values: {', '.join(values[:5])}"  # First 5 values
                    if len(values) > 5:
                        summary_text += f" (and {len(values) - 5} more)"
                
                chunk = {
                    'text': summary_text,
                    'metadata': {
                        'person_id': person_id,
                        'device_id': device_id,
                        'source_file': filename,
                        'field_name': column,
                        'field_type': field_type,
                        'semantic_role': semantic_role,
                        'chunk_type': 'csv_field_summary',
                        'value_count': len(values)
                    }
                }
                chunks.append(chunk)
        
        # Create complete CSV summary
        all_data = []
        for row in rows:
            row_data = [f"{k}: {v}" for k, v in row.items() if v and v.strip()]
            if row_data:
                all_data.append("; ".join(row_data))
        
        if all_data:
            complete_summary = f"Complete CSV data: {' | '.join(all_data)}"
            chunk = {
                'text': complete_summary,
                'metadata': {
                    'person_id': person_id,
                    'device_id': device_id,
                    'source_file': filename,
                    'chunk_type': 'csv_complete_summary',
                    'row_count': len(rows),
                    'column_count': len(all_columns)
                }
            }
            chunks.append(chunk)
        
        return chunks

    async def _fill_word_document(
        self,
        template_path: str,
        output_path: Path,
        template_analysis: Dict[str, Any],
        filled_fields: Dict[str, Any]
    ):
        """
        Fill a Word document template by replacing placeholders with field-specific intelligent matching
        """
        try:
            # Copy the original template to output location
            shutil.copy2(template_path, output_path)
            
            # Open the document
            doc = Document(str(output_path))
            
            logger.info(f"🔄 Filling Word document with {len(filled_fields)} fields")
            
            # Extract all text to analyze patterns
            all_text = self._extract_all_text_from_doc(doc)
            
            # Position-based intelligent replacement mapping
            field_replacements = {}
            
            # Step 1: Find all placeholders in document order
            import re
            placeholder_matches = list(re.finditer(r'\.{5,}', all_text))
            
            logger.info(f"🔄 Found {len(placeholder_matches)} placeholders in document order")
            
            # Step 2: Create a priority-ordered list of fields to fill
            field_priority_list = []
            skipped_fields = []
            
            for field_id, field_data in filled_fields.items():
                field_info = field_data['field_info']
                field_value = field_data['value']
                
                # Enhanced field value validation - skip if data is not meaningful
                should_skip = self._should_skip_field(field_value, field_info, field_id)
                
                if should_skip:
                    skipped_fields.append({
                        'field_id': field_id,
                        'field_type': field_info.get('field_type', 'unknown'),
                        'semantic_role': field_info.get('semantic_role', 'general'),
                        'reason': should_skip
                    })
                    logger.info(f"⏭️ Skipping field {field_id} ({field_info.get('field_type', 'unknown')}) - {should_skip}")
                    continue
                
                priority = self._calculate_field_priority(field_info)
                field_priority_list.append({
                    'field_id': field_id,
                    'field_data': field_data,
                    'priority': priority,
                    'semantic_role': field_info.get('semantic_role', 'general'),
                    'field_type': field_info.get('field_type', 'text')
                })
            
            # Log skipping summary
            if skipped_fields:
                logger.info(f"📊 Skipped {len(skipped_fields)} fields due to missing/invalid data:")
                for skip_info in skipped_fields:
                    logger.info(f"  ❌ {skip_info['field_id']} ({skip_info['field_type']}/{skip_info['semantic_role']}): {skip_info['reason']}")
                    
            logger.info(f"✅ Processing {len(field_priority_list)} fields with valid data")
            
            # Sort by priority (highest first)
            field_priority_list.sort(key=lambda x: x['priority'], reverse=True)
            
            # ENHANCED: Sequential context-aware placeholder assignment
            logger.info(f"🔄 Using SEQUENTIAL context-aware assignment for {len(placeholder_matches)} placeholders")
            
            assigned_placeholders = set()  # Track assigned field_ids
            successful_assignments = 0
            position_to_field = {}  # Maps placeholder position to field data
            
            # Process placeholders in document order (left to right, top to bottom)
            for i, match in enumerate(placeholder_matches):
                # Get immediate context around this specific placeholder
                context_window = 50  # Smaller window for precise context
                start = max(0, match.start() - context_window)
                end = min(len(all_text), match.end() + context_window)
                immediate_context = all_text[start:end]
                
                # Determine what this placeholder should contain based on IMMEDIATE context
                field_value, field_id, confidence = self._determine_placeholder_content_sequential(
                    immediate_context, match, filled_fields, assigned_placeholders
                )
                
                if field_value and confidence > 0.7:  # High confidence threshold
                    position_to_field[i] = {
                        'value': field_value,
                        'field_id': field_id,
                        'pattern': match.group(),
                        'score': confidence * 100
                    }
                    assigned_placeholders.add(field_id)  # Mark field as used
                    successful_assignments += 1
                    logger.info(f"✅ Sequential assignment #{i+1}: '{match.group()[:15]}...' → '{field_value}' ({field_id}, confidence: {confidence:.2f})")
                else:
                    logger.info(f"ℹ️ Skipping placeholder #{i+1}: '{match.group()[:15]}...' (low confidence or no suitable data)")
            
            logger.info(f"🔄 Sequential assignment: {successful_assignments} placeholders assigned out of {len(placeholder_matches)} total")
            
            # Step 4: Apply replacements using sequential position-based approach
            replacement_count = 0
            
            # Process each placeholder position in document order
            for position_idx in sorted(position_to_field.keys()):
                field_data = position_to_field[position_idx]
                pattern = field_data['pattern']
                value = field_data['value']
                field_id = field_data['field_id']
                
                # Apply replacement to paragraphs with format preservation
                for paragraph in doc.paragraphs:
                    if self._replace_in_paragraph(paragraph, pattern, str(value)):
                        replacement_count += 1
                        logger.info(f"📝 Position {position_idx+1}: Replaced '{pattern}' with '{value}' (field: {field_id})")
                        break  # Move to next position after replacing
                
                # Apply replacement to tables with format preservation
                replaced_in_table = False
                for table in doc.tables:
                    if replaced_in_table:
                        break
                    for row in table.rows:
                        if replaced_in_table:
                            break
                        for cell in row.cells:
                            if replaced_in_table:
                                break
                            for paragraph in cell.paragraphs:
                                if self._replace_in_paragraph(paragraph, pattern, str(value)):
                                    replacement_count += 1
                                    logger.info(f"📝 Position {position_idx+1}: Replaced '{pattern}' with '{value}' in table (field: {field_id})")
                                    replaced_in_table = True
                                    break
            
            # Save the document
            doc.save(str(output_path))
            logger.info(f"✅ Document saved with {replacement_count} total replacements made")
            
            # Return skip information for the caller
            return {
                'skipped_fields': skipped_fields,
                'filled_fields_count': successful_assignments,
                'total_replacements': replacement_count
            }
            
        except Exception as e:
            logger.error(f"❌ Failed to fill Word document: {e}")
            # Fallback to creating a new document
            fallback_result = await self._create_filled_document(output_path, template_analysis, filled_fields, 0)
            return fallback_result

    def _replace_in_paragraph(self, paragraph, pattern: str, value: str) -> bool:
        """
        Replace text in paragraph while preserving formatting
        Returns True if replacement was made, False otherwise
        """
        from docx.shared import RGBColor
        
        full_text = paragraph.text
        if pattern not in full_text:
            return False
        
        # Find the position of the pattern
        pattern_start = full_text.find(pattern)
        if pattern_start == -1:
            return False
        
        pattern_end = pattern_start + len(pattern)
        
        # Split text into: before_pattern + pattern + after_pattern
        before_text = full_text[:pattern_start]
        after_text = full_text[pattern_end:]
        
        # Clear the paragraph but preserve paragraph formatting
        p_format = paragraph.paragraph_format
        paragraph.clear()
        
        # Rebuild paragraph with preserved formatting
        if before_text:
            run_before = paragraph.add_run(before_text)
            # Try to preserve some basic formatting
            try:
                run_before.font.size = None  # Use default
                run_before.font.name = None  # Use default
            except:
                pass
        
        # Add the replacement value with emphasis
        run_value = paragraph.add_run(str(value))
        try:
            run_value.bold = True  # Make filled values bold
            # You can add more formatting here if needed
        except:
            pass
        
        if after_text:
            run_after = paragraph.add_run(after_text)
            try:
                run_after.font.size = None  # Use default
                run_after.font.name = None  # Use default
            except:
                pass
        
        return True

    def _determine_placeholder_content_sequential(
        self, 
        immediate_context: str, 
        match_obj, 
        filled_fields: Dict[str, Any],
        assigned_placeholders: set
    ) -> tuple[str, str, float]:
        """
        Determine what content should fill a placeholder based on immediate context
        Returns: (field_value, field_id, confidence_score)
        """
        context = immediate_context.lower()
        
        # ENHANCED: Position-aware AND person-aware pattern matching
        # Determine which person context we're in based on document position
        placeholder_global_pos = 0  # We'll work with immediate context for now
        
        person_context = "testator"  # default
        
        # Check for person transitions in the immediate context and before
        text_before_placeholder = immediate_context.lower()
        
        if "and b," in text_before_placeholder or "and b " in text_before_placeholder:
            person_context = "witness_a"
        elif "and c," in text_before_placeholder or "and c " in text_before_placeholder:
            person_context = "witness_b"
        
        # ENHANCED position-aware pattern matching with person context
        placeholder_text = match_obj.group()
        placeholder_pos = immediate_context.lower().find(placeholder_text.lower())
        
        if placeholder_pos == -1:
            # Fallback: find dots pattern
            placeholder_pos = immediate_context.find('.....')
        
        # Define person-aware context patterns
        context_patterns = [
            # Age patterns - person-aware
            {
                'patterns': ['aged about', 'age about'],
                'field_types': {
                    'testator': ['testator_age'],
                    'witness_a': ['witness_a_age'], 
                    'witness_b': ['witness_b_age']
                },
                'confidence': 0.95,
                'position': 'before'
            },
            # Father's name patterns - person-aware
            {
                'patterns': ['son of shri', 'son of late', 'son of mr'],
                'field_types': {
                    'testator': ['testator_father', 'father_name'],
                    'witness_a': ['witness_a_father'],
                    'witness_b': ['witness_b_father']
                },
                'confidence': 0.90,
                'position': 'before'
            },
            # Address patterns - person-aware
            {
                'patterns': ['resident of', 'residing at'],
                'field_types': {
                    'testator': ['testator_address', 'address'],
                    'witness_a': ['witness_a_address'],
                    'witness_b': ['witness_b_address']
                },
                'confidence': 0.85,
                'position': 'before'
            },
            # Testator name patterns (only for testator)
            {
                'patterns': ['executed by shri', 'testator shri', 'will executed by'],
                'field_types': {
                    'testator': ['testator_name', 'name'],
                    'witness_a': [],
                    'witness_b': []
                },
                'confidence': 0.90,
                'position': 'before'
            },
            # Date patterns (general)
            {
                'patterns': ['executed on', 'on this', 'day of'],
                'field_types': {
                    'testator': ['execution_date', 'legal_date', 'date'],
                    'witness_a': ['execution_date', 'legal_date', 'date'],
                    'witness_b': ['execution_date', 'legal_date', 'date']
                },
                'confidence': 0.85,
                'position': 'before'
            }
        ]
        
        # Position-aware pattern matching with person context
        best_match = None
        best_score = 0
        
        if placeholder_pos >= 0:
            context_before = context[:placeholder_pos]
            context_after = context[placeholder_pos + len(placeholder_text):]
            
            for pattern_group in context_patterns:
                # Get field types for the current person context
                person_field_types = pattern_group['field_types'].get(person_context, [])
                if not person_field_types:
                    continue
                
                for pattern in pattern_group['patterns']:
                    found_pos = -1
                    
                    if pattern_group['position'] == 'before':
                        found_pos = context_before.rfind(pattern)
                    else:
                        found_pos = context_after.find(pattern)
                    
                    if found_pos >= 0:
                        # Calculate proximity score
                        if pattern_group['position'] == 'before':
                            proximity_score = max(0, 50 - (placeholder_pos - found_pos - len(pattern))) / 50
                        else:
                            proximity_score = max(0, 50 - found_pos) / 50
                        
                        total_score = pattern_group['confidence'] * (0.7 + 0.3 * proximity_score)
                        
                        if total_score > best_score:
                            best_score = total_score
                            best_match = {
                                'field_types': person_field_types,
                                'confidence': total_score,
                                'matched_pattern': pattern,
                                'person_context': person_context,
                                'position': pattern_group['position']
                            }
        
        if not best_match:
            return None, None, 0
        
        best_field_types = best_match['field_types']
        best_confidence = best_match['confidence']
        
        # Find matching field data that hasn't been assigned yet
        for field_type in best_field_types:
            for field_id, field_data in filled_fields.items():
                if field_id in assigned_placeholders:
                    continue
                    
                field_info = field_data['field_info']
                actual_field_type = field_info.get('field_type', '')
                semantic_role = field_info.get('semantic_role', '')
                
                # Check if this field matches what we're looking for
                if (field_type in field_id.lower() or 
                    field_type == actual_field_type or
                    field_type in semantic_role.lower()):
                    
                    field_value = field_data['value']
                    
                    # Additional validation
                    if field_type in ['age', 'testator_age', 'witness_a_age', 'witness_b_age']:
                        if str(field_value).strip().isdigit():
                            return str(field_value), field_id, best_confidence
                    elif field_type in ['date', 'execution_date', 'legal_date']:
                        if any(char in str(field_value) for char in ['/', '-', '19', '20']):
                            return str(field_value), field_id, best_confidence
                    else:
                        # For names, addresses, etc.
                        if len(str(field_value).strip()) > 2:
                            return str(field_value), field_id, best_confidence
        
        return None, None, 0

    def _find_field_specific_placeholders(
        self, 
        text: str, 
        field_type: str, 
        field_info: Dict[str, Any], 
        field_value: str
    ) -> List[str]:
        """
        Find placeholders that are specific to a particular field type
        """
        import re
        placeholders = []
        
        # Get field suggestions and context
        suggestions = field_info.get('suggestions', [])
        context = field_info.get('context', '').lower()
        description = field_info.get('description', '').lower()
        
        # Create keyword sets for different field types
        field_keywords = {
            'name': ['name', 'applicant', 'person', 'individual', 'deponent'],
            'address': ['address', 'residence', 'location', 'plot', 'area', 'street'],
            'date': ['date', 'birth', 'dob', 'born', 'day', 'month', 'year'],
            'phone': ['phone', 'mobile', 'telephone', 'contact', 'number'],
            'email': ['email', 'mail', 'electronic'],
            'age': ['age', 'years', 'old'],
            'father': ['father', 'parent', 'guardian'],
            'signature': ['signature', 'sign', 'signed']
        }
        
        # Get relevant keywords for this field type
        relevant_keywords = []
        if field_type in field_keywords:
            relevant_keywords.extend(field_keywords[field_type])
        
        # Add suggestion keywords
        for suggestion in suggestions:
            relevant_keywords.extend(suggestion.lower().split())
        
        # Remove duplicates
        relevant_keywords = list(set(relevant_keywords))
        
        logger.info(f"� Looking for {field_type} placeholders with keywords: {relevant_keywords}")
        
        # Find placeholders with field-specific patterns
        patterns = [
            r'_{3,}',  # Multiple underscores
            r'\[([^\]]*)\]',  # Square brackets
            r'\{([^}]*)\}',  # Curly braces
            r'\.{3,}',  # Multiple dots
            r'-{3,}',  # Multiple dashes
            r'\([^)]*\)',  # Parentheses
            r'\b35\b',  # Specific number 35 (common placeholder in your documents)
            r'\b\d{2}\b\.{3,}',  # Numbers followed by dots
            r'\d{2}\.{2,}',  # Numbers with multiple dots
            r'\b\d{1,2}\s*\.{2,}',  # Numbers with dots and optional spaces
            r'3535\.{2,}',  # Specific pattern like 3535....
            r'\b35\.{2,}'  # 35 followed by dots
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                placeholder_text = match.group(0)
                placeholder_context = text[max(0, match.start()-50):match.end()+50].lower()
                
                # Check if this placeholder is relevant to the field type
                is_relevant = False
                
                # Method 1: Check if any relevant keywords appear near the placeholder
                for keyword in relevant_keywords:
                    if keyword in placeholder_context:
                        is_relevant = True
                        logger.info(f"✅ Found relevant placeholder '{placeholder_text}' (keyword: '{keyword}' in context)")
                        break
                
                # Enhanced contextual matching for legal documents
                field_semantic_role = field_info.get('semantic_role', 'general')
                
                if not is_relevant:
                    # For testator names - look for "executed by Shri" pattern
                    if (field_type == 'name' and field_semantic_role == 'testator' and 
                        any(term in placeholder_context for term in ['executed by shri', 'testator shri'])):
                        is_relevant = True
                        logger.info(f"✅ Found testator name placeholder '{placeholder_text}' by legal context")
                    
                    # For testator father names - look for "son of" after testator context
                    elif (field_type == 'name' and field_semantic_role in ['father_of_testator', 'father'] and 
                          any(term in placeholder_context for term in ['shri', 'son of']) and
                          any(term in placeholder_context for term in ['executed', 'will', 'testator'])):
                        is_relevant = True
                        logger.info(f"✅ Found testator father name placeholder '{placeholder_text}' by legal context")
                    
                    # For testator address - look for "resident of" after testator context
                    elif (field_type == 'address' and field_semantic_role == 'testator' and 
                          any(term in placeholder_context for term in ['resident of']) and
                          any(term in placeholder_context for term in ['will', 'executed', 'testator'])):
                        is_relevant = True
                        logger.info(f"✅ Found testator address placeholder '{placeholder_text}' by legal context")
                    
                    # For witness A details - look for "A, aged about"
                    elif (field_semantic_role == 'witness_a' and 
                          any(term in placeholder_context for term in ['a, aged about', 'affidavit of a'])):
                        is_relevant = True
                        logger.info(f"✅ Found witness A placeholder '{placeholder_text}' by legal context")
                    
                    # For witness B details - look for "B, aged about"
                    elif (field_semantic_role == 'witness_b' and 
                          any(term in placeholder_context for term in ['b, aged about', 'and b, aged'])):
                        is_relevant = True
                        logger.info(f"✅ Found witness B placeholder '{placeholder_text}' by legal context")
                    
                    # For will execution dates - look for "executed his will on" or "will on"
                    elif (field_type == 'date' and 
                          any(term in placeholder_context for term in ['executed his will on', 'will on', 'executed on'])):
                        is_relevant = True
                        logger.info(f"✅ Found will execution date placeholder '{placeholder_text}' by legal context")
                    
                    # For sub-registrar - look for "Sub-Registrar"
                    elif (field_semantic_role == 'sub_registrar' and 
                          any(term in placeholder_context for term in ['sub-registrar', 'before the sub'])):
                        is_relevant = True
                        logger.info(f"✅ Found sub-registrar placeholder '{placeholder_text}' by legal context")
                    
                    # For verification location - look for "Verified at"
                    elif (field_type in ['location', 'text'] and 
                          any(term in placeholder_context for term in ['verified at'])):
                        is_relevant = True
                        logger.info(f"✅ Found verification location placeholder '{placeholder_text}' by legal context")
                
                # Method 3: General position matching
                if not is_relevant and field_type == 'name':
                    if re.search(r'\b(name|applicant|deponent|shri)\s*[:\-]?\s*' + re.escape(placeholder_text), placeholder_context, re.IGNORECASE):
                        is_relevant = True
                        logger.info(f"✅ Found name placeholder '{placeholder_text}' by general position")
                
                elif not is_relevant and field_type == 'date':
                    if re.search(r'\b(date|born|birth|day|on)\s*[:\-]?\s*' + re.escape(placeholder_text), placeholder_context, re.IGNORECASE):
                        is_relevant = True
                        logger.info(f"✅ Found date placeholder '{placeholder_text}' by general position")
                
                elif not is_relevant and field_type == 'address':
                    # Look for address-like positions
                    if re.search(r'\b(address|residence|plot)\s*[:\-]?\s*' + re.escape(placeholder_text), placeholder_context, re.IGNORECASE):
                        is_relevant = True
                        logger.info(f"✅ Found address placeholder '{placeholder_text}' by position")
                
                if is_relevant and len(placeholder_text) < 50:  # Reasonable length
                    placeholders.append(placeholder_text)
        
        # If no specific placeholders found, use a more intelligent approach
        if not placeholders:
            logger.warning(f"⚠️ No specific placeholders found for {field_type}, using general patterns")
            
            # For the document format you showed, try to find "35" patterns contextually
            if field_type in ['name', 'date', 'address', 'age']:
                # Look for "35" patterns near relevant keywords
                context_patterns = {
                    'name': [r'(Shri\s+35)', r'(by\s+Shri\s+35)', r'(deponents?\s+35)', r'(A,\s*aged\s*about\s*35)', r'(B,\s*aged\s*about\s*35)'],
                    'age': [r'(aged\s+about\s+35)', r'(aged\s+35)'],
                    'date': [r'(on\s+35)', r'(executed\s+.*\s+on\s+35)', r'(day\s+of\s+35)'],
                    'address': [r'(resident\s+of\s+35)', r'(of\s+35\.+)']
                }
                
                if field_type in context_patterns:
                    for pattern in context_patterns[field_type]:
                        matches = re.findall(pattern, text, re.IGNORECASE)
                        if matches:
                            # Extract just the "35" part
                            placeholders.extend(['35'])
                            logger.info(f"✅ Found contextual {field_type} placeholder: 35")
                            break
            
            # If still no placeholders, try general patterns but limit to first occurrence
            if not placeholders:
                general_patterns = re.findall(r'_{3,}', text)
                if general_patterns:
                    placeholders.extend(general_patterns[:1])  # Take only the first one to avoid duplicates
        
        return list(set(placeholders))  # Remove duplicates

    def _should_skip_field(self, field_value: Any, field_info: Dict[str, Any], field_id: str) -> Optional[str]:
        """
        Determine if a field should be skipped during document filling
        Returns skip reason if field should be skipped, None if field should be processed
        """
        # Check if value is None or empty
        if field_value is None:
            return "no data found in knowledge base"
        
        # Convert to string for validation
        value_str = str(field_value).strip()
        
        # Check for empty strings
        if not value_str:
            return "empty value"
        
        # Check for placeholder indicators (these indicate missing data)
        placeholder_indicators = [
            'TO_BE_FILLED',
            'NOT_FOUND', 
            'NOT_AVAILABLE',
            'MISSING',
            'N/A',
            'NULL',
            'NONE',
            'UNKNOWN'
        ]
        
        value_upper = value_str.upper()
        for indicator in placeholder_indicators:
            if indicator in value_upper:
                return f"contains placeholder indicator '{indicator}'"
        
        # Check for bracket placeholders like [FIELD_NAME_TO_BE_FILLED]
        if value_str.startswith('[') and value_str.endswith(']'):
            return "contains placeholder brackets"
        
        # Check for very short values that might not be meaningful
        if len(value_str) < 2:
            return "value too short (less than 2 characters)"
        
        # Check for values that are just numbers without context (might be IDs, not real data)
        if value_str.isdigit() and len(value_str) > 10:
            return "appears to be ID number, not meaningful data"
        
        # Field-type specific validations
        field_type = field_info.get('field_type', '').lower()
        
        if field_type == 'age':
            # Age should be numeric and reasonable
            if not value_str.isdigit():
                return "age field contains non-numeric value"
            age_val = int(value_str)
            if age_val < 1 or age_val > 120:
                return f"age value {age_val} is unrealistic"
        
        elif field_type == 'email':
            # Basic email validation
            if '@' not in value_str or '.' not in value_str:
                return "email field missing @ or . symbols"
        
        elif field_type == 'phone':
            # Phone should contain numbers
            if not any(c.isdigit() for c in value_str):
                return "phone field contains no digits"
        
        elif field_type == 'name':
            # Names should not be just numbers
            if value_str.isdigit():
                return "name field contains only numbers"
            # Names should have reasonable length
            if len(value_str) > 100:
                return "name value too long (over 100 characters)"
        
        # Check for generic "not found" patterns
        not_found_patterns = [
            'not found',
            'no data',
            'no information',
            'data not available',
            'information not available',
            'could not find',
            'unable to determine'
        ]
        
        value_lower = value_str.lower()
        for pattern in not_found_patterns:
            if pattern in value_lower:
                return f"contains '{pattern}' indicating no data"
        
        # If all checks pass, field should not be skipped
        return None

    def _calculate_field_priority(self, field_info: Dict[str, Any]) -> int:
        """
        Calculate priority for field assignment to resolve conflicts
        Higher priority fields get first choice of placeholders
        """
        priority = 0
        field_type = field_info.get('field_type', '').lower()
        semantic_role = field_info.get('semantic_role', 'general')
        field_id = field_info.get('field_id', '')
        
        # Priority based on semantic role importance
        role_priorities = {
            'testator': 100,  # Highest priority - main subject
            'father_of_testator': 90,
            'witness_a': 80,
            'father_of_witness_a': 70, 
            'witness_b': 60,
            'father_of_witness_b': 50,
            'sub_registrar': 40,
            'general': 30
        }
        
        priority += role_priorities.get(semantic_role, 20)
        
        # Additional priority based on field type importance
        type_priorities = {
            'name': 20,
            'address': 15,
            'date': 10,
            'age': 8,
            'location': 5,
            'text': 1
        }
        
        priority += type_priorities.get(field_type, 1)
        
        # Boost priority for certain key field patterns
        if 'testator_name' in field_id:
            priority += 50  # Testator name is most important
        elif 'execution_date' in field_id:
            priority += 30  # Will execution dates are important
        elif '_1' in field_id:  # First occurrence of field type gets priority
            priority += 10
        
        return priority

    def _score_placeholder_for_field(self, context: str, semantic_role: str, field_type: str, position: int) -> float:
        """
        ENHANCED: Precise pattern matching for legal affidavit documents
        """
        score = 0.0
        context_lower = context.lower()
        
        # ULTRA-PRECISE pattern matching for legal documents
        exact_patterns = {
            'age': [
                'aged about ............. years',
                'aged about .............years', 
                ', aged about ............. years,',
                'age about ............. years',
                'aged ............. years'
            ],
            'name': [
                'executed by shri .............', 
                'shri ............. son of',
                'testator shri .............', 
                'son of shri .............', 
                'by shri ............. resident',
                'will executed by shri .............'
            ],
            'address': [
                'resident of .............', 
                'resident of .......................',
                'address .............', 
                'residing at ..............'
            ],
            'date': [
                'executed on .............', 
                'will on .............', 
                'on .............', 
                'day of .............', 
                'verified at ............. on'
            ],
            'location': [
                'sub-registrar .............', 
                'before the sub-registrar .............', 
                'verified at .............', 
                'at ............. on this'
            ]
        }
        
        # Check for exact pattern matches (HIGHEST PRIORITY)
        if field_type in exact_patterns:
            for exact_pattern in exact_patterns[field_type]:
                if exact_pattern in context_lower:
                    score += 100  # Very high score for exact matches
                    logger.info(f"🎯 EXACT PATTERN MATCH: {field_type} found '{exact_pattern}'")
                    
        # Strong context indicators
        strong_indicators = {
            'age': ['aged about', 'aged', 'years'],
            'name': ['executed by', 'shri', 'son of', 'testator'],
            'address': ['resident of', 'address', 'residing'],
            'date': ['executed on', 'on', 'day of'],
            'location': ['sub-registrar', 'before the', 'verified at']
        }
        
        if field_type in strong_indicators:
            for indicator in strong_indicators[field_type]:
                if indicator in context_lower:
                    score += 30
                    logger.info(f"🎯 STRONG INDICATOR: {field_type} found '{indicator}'")
        
        # WITNESS-SPECIFIC patterns (very important for affidavits)
        witness_patterns = {
            ('witness_a', 'name'): ['affidavit of a', 'a,', 'and a,'],
            ('witness_a', 'age'): ['a, aged about', 'of a, aged'],
            ('witness_b', 'name'): ['and b,', 'b,', ', b '],
            ('witness_b', 'age'): ['b, aged about', 'and b, aged'],
            ('testator', 'name'): ['executed by shri', 'testator shri'],
            ('testator', 'address'): ['resident of', 'testator.*resident'],
            ('father_of_testator', 'name'): ['son of shri', 'son of'],
            ('sub_registrar', 'location'): ['sub-registrar', 'before the sub']
        }
        
        pattern_key = (semantic_role, field_type)
        if pattern_key in witness_patterns:
            for pattern in witness_patterns[pattern_key]:
                if pattern in context_lower:
                    score += 50  # High boost for role-specific patterns
                    logger.info(f"🏆 WITNESS PATTERN: {semantic_role}+{field_type} matched '{pattern}'")
        
        # NEGATIVE patterns - heavily penalize wrong assignments
        negative_rules = {
            'age': ['executed by', 'shri', 'son of', 'resident of', 'address'],
            'name': ['aged about', 'years', 'resident of', 'address'],
            'address': ['aged about', 'years', 'executed by', 'son of'],
            'date': ['aged about', 'shri', 'resident of', 'son of'],
            'location': ['aged about', 'years', 'son of']
        }
        
        if field_type in negative_rules:
            for neg_pattern in negative_rules[field_type]:
                if neg_pattern in context_lower:
                    score -= 40  # Heavy penalty for conflicting patterns
                    logger.info(f"❌ NEGATIVE PATTERN: {field_type} penalized for '{neg_pattern}'")
        
        # Position-based scoring for legal document structure
        position_preferences = {
            'sub_registrar': (0, 2),
            'testator': (2, 8), 
            'father_of_testator': (4, 10),
            'witness_a': (8, 15),
            'witness_b': (12, 20),
            'date': (15, 25),
            'location': (20, 30)
        }
        
        if semantic_role in position_preferences:
            start, end = position_preferences[semantic_role]
            if start <= position <= end:
                score += 15  # Position bonus
            else:
                penalty = min(abs(position - start), abs(position - end)) * 2
                score -= penalty
                
        logger.info(f"📊 FINAL SCORE: {field_type}+{semantic_role} at pos {position} = {score:.2f}")
        return max(score, 0.0)

    def _extract_all_text_from_doc(self, doc: Document) -> str:
        """
        Extract all text from a Word document
        """
        all_text = ""
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            all_text += paragraph.text + "\n"
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        all_text += paragraph.text + " "
        
        return all_text

    def _find_placeholders_in_text(self, text: str, field_type: str = '') -> List[str]:
        """
        Find potential placeholders in text using enhanced pattern matching
        """
        placeholders = []
        
        # Enhanced placeholder patterns for legal documents
        patterns = [
            r'\.{4,}',  # Multiple dots (4 or more) - common in legal docs
            r'_{3,}',  # Multiple underscores (3 or more)
            r'\s{4,}',  # Multiple spaces (4 or more) - blank lines
            r'\[([^\]]*)\]',  # Text in square brackets
            r'\{([^}]*)\}',  # Text in curly braces
            r'\(\s*([^)]*)\s*\)',  # Text in parentheses with optional spaces
            r'<<[^>]*>>',  # Text in double angle brackets
            r'\b[A-Z_]{3,}\b',  # ALL_CAPS words (likely placeholders)
            r'_+[A-Za-z0-9_]*_+',  # Words surrounded by underscores
            r'\*{2,}[^*]*\*{2,}',  # Text surrounded by multiple asterisks
            r'-{3,}',  # Multiple dashes (3 or more)
            r'={3,}',  # Multiple equals signs (3 or more)
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                if isinstance(match, tuple):
                    match = match[0] if match else ''
                
                # Clean up the match
                match = str(match).strip()
                
                # Filter criteria
                if (len(match) >= 3 and len(match) <= 100 and  # Reasonable length
                    not match.isspace() and  # Not just whitespace
                    not match.replace('.', '').replace('_', '').replace('-', '').replace('=', '') == ''):  # Not just symbols
                    placeholders.append(match)
        
        # Also look for the original dot patterns
        dot_patterns = re.findall(r'\.{4,}', text)
        for pattern in dot_patterns:
            if len(pattern) >= 4:
                placeholders.append(pattern)
        
        return list(set(placeholders))  # Remove duplicates

    async def _create_filled_document(
        self,
        output_path: Path,
        template_analysis: Dict[str, Any],
        filled_fields: Dict[str, Any],
        template_idx: int
    ):
        """
        Create a new filled document when template modification fails
        """
        try:
            # Check for skipped fields (same logic as Word document)
            skipped_fields = []
            valid_fields = {}
            
            for field_id, field_data in filled_fields.items():
                field_info = field_data['field_info']
                field_value = field_data['value']
                
                should_skip = self._should_skip_field(field_value, field_info, field_id)
                
                if should_skip:
                    skipped_fields.append({
                        'field_id': field_id,
                        'field_type': field_info.get('field_type', 'unknown'),
                        'semantic_role': field_info.get('semantic_role', 'general'),
                        'reason': should_skip
                    })
                    logger.info(f"⏭️ Skipping field {field_id} in new document - {should_skip}")
                else:
                    valid_fields[field_id] = field_data
            
            # Create a new Word document
            doc = Document()
            
            # Add title
            title = doc.add_heading(f"Filled Document #{template_idx + 1}", 0)
            
            # Add metadata including skip information
            metadata_para = doc.add_paragraph()
            metadata_para.add_run("Generated on: ").bold = True
            metadata_para.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
            metadata_para.add_run("\nTemplate: ").bold = True
            metadata_para.add_run(template_analysis['template_info']['filename'])
            metadata_para.add_run(f"\nFields filled: ").bold = True
            metadata_para.add_run(f"{len(valid_fields)}/{len(filled_fields)}")
            
            if skipped_fields:
                metadata_para.add_run(f"\nFields skipped: ").bold = True
                metadata_para.add_run(f"{len(skipped_fields)} (missing data)")
            
            # Add separator
            doc.add_paragraph("_" * 80)
            
            # Generate document content using LLM with valid fields only
            document_content = await self._generate_document_content(
                template_analysis,
                valid_fields,  # Use only valid fields
                template_idx
            )
            
            # Add the generated content
            doc.add_paragraph(document_content)
            
            # Save the document
            doc.save(str(output_path))
            
            # Return skip information
            return {
                'skipped_fields': skipped_fields,
                'filled_fields_count': len(valid_fields),
                'total_fields': len(filled_fields)
            }
            
        except Exception as e:
            logger.error(f"❌ Failed to create filled document: {e}")
            # Ultimate fallback - save as text file
            with open(output_path.with_suffix('.txt'), 'w', encoding='utf-8') as f:
                f.write(f"Error creating document: {e}\n\n")
                f.write(json.dumps(filled_fields, indent=2))
            
            # Return minimal info even on error
            return {
                'skipped_fields': [],
                'filled_fields_count': 0,
                'total_fields': len(filled_fields),
                'error': str(e)
            }

    async def debug_stored_data(self, device_id: str, person_id: str = None) -> Dict[str, Any]:
        """
        Debug method to check what data is stored for a device/person
        """
        try:
            result = {
                'device_data': {},
                'person_data': {},
                'success': True
            }
            
            # Check device storage (original functionality)
            storage_file = Path("local_vector_storage") / f"device_{device_id}_vectors.json"
            if storage_file.exists():
                with open(storage_file, 'r', encoding='utf-8') as f:
                    device_data = json.load(f)
                
                result['device_data'] = {
                    'total_vectors': len(device_data),
                    'first_5_samples': device_data[:5]
                }
                logger.info(f"📊 Found {len(device_data)} device vectors for {device_id}")
            
            # Check person storage (new functionality)
            if person_id:
                person_storage_file = self._get_person_storage_file(person_id)
                if person_storage_file.exists():
                    with open(person_storage_file, 'r', encoding='utf-8') as f:
                        person_data = json.load(f)
                    
                    result['person_data'] = {
                        'person_id': person_id,
                        'total_vectors': len(person_data),
                        'first_5_samples': person_data[:5]
                    }
                    logger.info(f"👤 Found {len(person_data)} person vectors for {person_id}")
                    
                    # Show sample field values
                    for i, vector in enumerate(person_data[:5]):
                        metadata = vector.get('metadata', {})
                        text = vector.get('text', '')[:100]
                        field_value = metadata.get('field_value', 'N/A')
                        logger.info(f"🔍 Person Vector {i+1}: {metadata.get('field_name', 'N/A')} = {field_value}")
                else:
                    result['person_data'] = {
                        'person_id': person_id,
                        'total_vectors': 0,
                        'message': f'No person data found for {person_id}'
                    }
            else:
                # List all available person_ids
                person_files = list(self.person_storage_dir.glob("person_*_data.json"))
                available_persons = [f.stem.replace('person_', '').replace('_data', '') for f in person_files]
                result['person_data'] = {
                    'available_person_ids': available_persons,
                    'total_persons': len(available_persons)
                }
                logger.info(f"👥 Found {len(available_persons)} persons: {available_persons}")
            
            return result
                
        except Exception as e:
            logger.error(f"❌ Failed to debug stored data: {e}")
            return {'success': False, 'error': str(e)}

# Global service instance
interpreted_form_service = InterpretedFormService()
