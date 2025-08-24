import os
import uuid
import json
import aiofiles
import csv
import pandas as pd
import asyncio
from typing import List, Dict, Any, Optional
from pathlib import Path
import logging
from docx import Document
import PyPDF2
import re
import unicodedata
from io import BytesIO, StringIO
from dotenv import load_dotenv

# Optional imports for enhanced PDF processing
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logger.warning("PyMuPDF not available - using alternative PDF processors")

try:
    from pdfminer.high_level import extract_text as pdfminer_extract_text
    from pdfminer.layout import LAParams
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False

# Load environment variables
load_dotenv()

from app.services.gemini_service import gemini_service
from app.services.pinecone_service import pinecone_service
from app.database import document_repo

# OLD OCR SERVICE IMPORTS - COMMENTED OUT FOR GOOGLE VISION MIGRATION
# Simple OCR Service Integration
# try:
#     from app.services.simple_ocr_service import simple_ocr_service
#     OCR_AVAILABLE = simple_ocr_service.is_available()
#     logger = logging.getLogger(__name__)
#     if OCR_AVAILABLE:
#         logger.info("✅ Simple OCR Service integrated successfully")
#     else:
#         logger.warning("⚠️ No OCR engines available")
# except ImportError as e:
#     OCR_AVAILABLE = False
#     logger = logging.getLogger(__name__)
#     logger.warning(f"⚠️ OCR Service not available: {e}")

# NEW GOOGLE VISION OCR SERVICE
try:
    from app.services.google_vision_ocr_service import google_vision_ocr_service
    OCR_AVAILABLE = google_vision_ocr_service.is_available()
    logger = logging.getLogger(__name__)
    if OCR_AVAILABLE:
        logger.info("✅ Google Vision OCR Service integrated successfully")
    else:
        logger.warning("⚠️ Google Vision OCR not available")
except ImportError as e:
    OCR_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.warning(f"⚠️ Google Vision OCR Service not available: {e}")

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.upload_dir = Path(os.getenv("UPLOAD_DIR", "./uploads"))
        self.upload_dir.mkdir(exist_ok=True)
        # ENHANCED: Optimized chunk sizes for better information capture
        self.chunk_size = 1500  # Increased for more context
        self.chunk_overlap = 400  # Increased overlap for better continuity
        self.min_chunk_size = 300  # Minimum viable chunk size
        
        # OCR Integration Configuration
        self.ocr_enabled = OCR_AVAILABLE and os.getenv("OCR_ENABLED", "true").lower() == "true"
        self.force_ocr = os.getenv("FORCE_OCR", "false").lower() == "true"
        self.ocr_fallback_threshold = 100  # If text extraction yields < 100 chars, try OCR
        self._last_ocr_metadata = None  # Temporary storage for OCR metadata
        
        if self.ocr_enabled:
            logger.info("🔍 Google Vision OCR integration enabled - intelligent scanned document detection active")
        else:
            logger.info("📄 Google Vision OCR integration disabled - text-based documents only")
    
    async def process_uploaded_file(
        self, 
        file_content: bytes, 
        filename: str, 
        device_id: str
    ) -> Dict[str, Any]:
        """Process uploaded file and store in vector database"""
        try:
            total_start_time = asyncio.get_event_loop().time()
            logger.info(f"🚀 Starting to process document: {filename} for device: {device_id}")
            logger.info(f"📊 File size: {len(file_content)} bytes")
            
            # Generate unique document ID
            document_id = str(uuid.uuid4())
            
            # STEP 1: Extract text from file
            logger.info("📄 STEP 1/5: Starting text extraction...")
            text_start_time = asyncio.get_event_loop().time()
            
            text_content = await self._extract_text(file_content, filename)
            if not text_content:
                raise ValueError("Could not extract text from file")
            
            text_duration = asyncio.get_event_loop().time() - text_start_time
            logger.info(f"✅ STEP 1 COMPLETE: Text extraction took {text_duration:.2f}s")
            logger.info(f"📝 Extracted {len(text_content)} characters")
            
            # Add debug info for OCR results
            if hasattr(self, '_last_ocr_metadata') and self._last_ocr_metadata:
                logger.info(f"🔍 OCR was used: {self._last_ocr_metadata.get('ocr_required', False)}")
                if self._last_ocr_metadata.get('total_processing_time'):
                    logger.info(f"⏱️ OCR processing time: {self._last_ocr_metadata['total_processing_time']:.2f}s")
            
            # Show first 200 characters of extracted text for debugging
            preview_text = text_content[:200].replace('\n', ' ').replace('\r', ' ')
            logger.info(f"📄 Text preview: {preview_text}...")
            
            # Check if text seems to be from OCR (common patterns)
            if any(pattern in text_content.lower() for pattern in ['[ocr_content]', 'scanned', 'image']):
                logger.info("🔍 Text appears to be from OCR processing")
            
            # Count paragraphs/sections for chunking info
            paragraph_count = len([p for p in text_content.split('\n\n') if p.strip()])
            line_count = len([l for l in text_content.split('\n') if l.strip()])
            logger.info(f"📊 Document structure: {paragraph_count} paragraphs, {line_count} lines")

            # STEP 2: Create chunks
            logger.info("📦 STEP 2/5: Starting chunking process...")
            chunk_start_time = asyncio.get_event_loop().time()
            
            chunks = self._create_chunks(text_content)
            if not chunks:
                raise ValueError("No chunks were created from the document")
            
            chunk_duration = asyncio.get_event_loop().time() - chunk_start_time
            logger.info(f"✅ STEP 2 COMPLETE: Chunking took {chunk_duration:.2f}s")
            logger.info(f"📦 Created {len(chunks)} chunks from document")

            # STEP 3: Generate embeddings and store in Pinecone
            logger.info("🔗 STEP 3/5: Starting vector storage process...")
            vector_start_time = asyncio.get_event_loop().time()
            
            try:
                await asyncio.wait_for(
                    self._store_chunks_in_pinecone(chunks, document_id, device_id, filename),
                    timeout=120.0  # 2-minute timeout for entire vector storage process
                )
                
                vector_duration = asyncio.get_event_loop().time() - vector_start_time
                logger.info(f"✅ STEP 3 COMPLETE: Vector storage took {vector_duration:.2f}s")
                
            except asyncio.TimeoutError:
                logger.error("❌ STEP 3 FAILED: Vector storage timed out after 2 minutes")
                raise ValueError("Document processing timed out during vector storage. Please try with a smaller document or check network connectivity.")

            # STEP 4: Store metadata in MongoDB
            logger.info("💾 STEP 4/5: Storing document metadata...")
            metadata_start_time = asyncio.get_event_loop().time()
            document_metadata = {
                "document_id": document_id,
                "filename": filename,
                "file_size": len(file_content),
                "file_type": Path(filename).suffix.lower(),
                "device_id": device_id,
                "chunk_count": len(chunks),
                "processed": True,
                "processing_method": "enhanced_extraction",
                "text_length": len(text_content),
                "chunk_statistics": {
                    "total_chunks": len(chunks),
                    "avg_chunk_size": sum(len(chunk["content"]) for chunk in chunks) / len(chunks) if chunks else 0,
                    "high_importance_chunks": sum(1 for chunk in chunks if chunk.get("importance_score", 0) > 0.7)
                }
            }
            
            # Add OCR metadata if available
            if hasattr(self, '_last_ocr_metadata') and self._last_ocr_metadata:
                document_metadata["ocr_processing"] = self._last_ocr_metadata
                # Clear the temporary metadata
                self._last_ocr_metadata = None
            
            await document_repo.create_document(document_metadata)
            
            metadata_duration = asyncio.get_event_loop().time() - metadata_start_time
            logger.info(f"✅ STEP 4 COMPLETE: Metadata storage took {metadata_duration:.2f}s")
            
            # STEP 5: Save file to disk (optional, for backup)
            logger.info("💾 STEP 5/5: Saving file backup...")
            file_start_time = asyncio.get_event_loop().time()
            
            file_path = self.upload_dir / f"{document_id}_{filename}"
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(file_content)
            
            file_duration = asyncio.get_event_loop().time() - file_start_time
            total_duration = asyncio.get_event_loop().time() - total_start_time
            
            logger.info(f"✅ STEP 5 COMPLETE: File backup took {file_duration:.2f}s")
            logger.info(f"🎉 PROCESSING COMPLETE: Total time {total_duration:.2f}s")
            logger.info(f"📊 TIMING BREAKDOWN:")
            logger.info(f"   📄 Text extraction: {text_duration:.2f}s")
            logger.info(f"   📦 Chunking: {chunk_duration:.2f}s") 
            logger.info(f"   🔗 Vector storage: {vector_duration:.2f}s")
            logger.info(f"   💾 Metadata: {metadata_duration:.2f}s")
            logger.info(f"   💾 File backup: {file_duration:.2f}s")
            logger.info(f"✅ Successfully processed document {filename} for device {device_id} - Created {len(chunks)} chunks")
            
            return {
                "document_id": document_id,
                "filename": filename,
                "device_id": device_id,
                "chunks_created": len(chunks),
                "status": "success"
            }
            
        except Exception as e:
            logger.error(f"❌ Failed to process document {filename}: {e}")
            raise
    
    async def _extract_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from different file types with intelligent OCR integration"""
        try:
            file_extension = Path(filename).suffix.lower()
            logger.info(f"📄 Extracting text from {filename} (type: {file_extension})")
            
            extracted_text = ""
            
            if file_extension == '.txt':
                extracted_text = file_content.decode('utf-8')
            
            elif file_extension == '.pdf':
                extracted_text = await self._extract_text_from_pdf_with_ocr(file_content, filename)
            
            elif file_extension == '.docx':
                extracted_text = self._extract_text_from_docx(file_content)
            
            elif file_extension == '.md':
                extracted_text = file_content.decode('utf-8')
            
            elif file_extension == '.csv':
                extracted_text = self._extract_text_from_csv(file_content)
            
            elif file_extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                # Image files - always use OCR
                extracted_text = await self._extract_text_from_image(file_content, filename)
            
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Clean up the extracted text
            cleaned_text = extracted_text.strip()
            if not cleaned_text:
                raise ValueError(f"No text content found in file {filename}")
            
            logger.info(f"✅ Extracted {len(cleaned_text)} characters from {filename}")
            return cleaned_text
                
        except Exception as e:
            logger.error(f"❌ Failed to extract text from {filename}: {e}")
            raise
    
    def _extract_text_from_pdf(self, file_content: bytes) -> str:
        """Enhanced PDF text extraction using multiple methods for better accuracy"""
        try:
            logger.info("🔍 Starting enhanced PDF text extraction...")
            pdf_file = BytesIO(file_content)
            
            # Method 1: Try pdfplumber (best for layout preservation)
            text_plumber = self._extract_with_pdfplumber(pdf_file)
            if text_plumber and len(text_plumber.strip()) > 100:
                logger.info("✅ Successfully extracted text using pdfplumber")
                return self._clean_extracted_text(text_plumber)
            
            # Method 2: Try PyMuPDF (good for most PDFs)
            pdf_file.seek(0)
            text_pymupdf = self._extract_with_pymupdf(pdf_file)
            if text_pymupdf and len(text_pymupdf.strip()) > 100:
                logger.info("✅ Successfully extracted text using PyMuPDF")
                return self._clean_extracted_text(text_pymupdf)
            
            # Method 3: Try pdfminer (good for complex layouts)
            pdf_file.seek(0)
            text_pdfminer = self._extract_with_pdfminer(pdf_file)
            if text_pdfminer and len(text_pdfminer.strip()) > 100:
                logger.info("✅ Successfully extracted text using pdfminer")
                return self._clean_extracted_text(text_pdfminer)
            
            # Method 4: Fallback to PyPDF2 (basic extraction)
            pdf_file.seek(0)
            text_pypdf2 = self._extract_with_pypdf2(pdf_file)
            if text_pypdf2 and len(text_pypdf2.strip()) > 50:
                logger.info("⚠️ Using PyPDF2 as fallback")
                return self._clean_extracted_text(text_pypdf2)
            
            # If all methods fail but we got some text, use the best available
            all_texts = [text_plumber, text_pymupdf, text_pdfminer, text_pypdf2]
            best_text = max((text for text in all_texts if text), key=len, default="")
            
            if best_text and len(best_text.strip()) > 20:
                logger.warning("⚠️ Using partial text extraction result")
                return self._clean_extracted_text(best_text)
            
            # Return empty string instead of raising exception to allow OCR fallback
            logger.warning("⚠️ No readable text could be extracted using traditional methods - will try OCR")
            return ""
            
        except Exception as e:
            logger.error(f"❌ Failed to extract text from PDF: {e}")
            # Return empty string instead of re-raising to allow OCR fallback
            return ""
    
    async def _extract_text_from_pdf_with_ocr(self, file_content: bytes, filename: str) -> str:
        """
        Intelligent PDF text extraction with OCR fallback
        First tries traditional text extraction, then uses OCR for scanned documents
        """
        try:
            extraction_start_time = asyncio.get_event_loop().time()
            logger.info(f"🔍 Starting intelligent PDF processing for {filename}")
            
            # Step 1: Try traditional text extraction first
            traditional_text = ""
            traditional_start_time = asyncio.get_event_loop().time()
            try:
                traditional_text = self._extract_text_from_pdf(file_content)
                text_length = len(traditional_text.strip())
                
                traditional_end_time = asyncio.get_event_loop().time()
                traditional_duration = traditional_end_time - traditional_start_time
                logger.info(f"📄 Traditional extraction took {traditional_duration:.2f}s, yielded {text_length} characters")
                
                # If we got substantial text, we're done
                if text_length >= self.ocr_fallback_threshold:
                    logger.info(f"✅ Sufficient text extracted traditionally - no OCR needed")
                    return traditional_text
                
                # If we got some text but not much, it might be a mixed document
                if text_length > 20:
                    logger.info(f"⚠️ Limited text found ({text_length} chars) - checking if OCR can improve")
                    
                    if self.ocr_enabled and not self.force_ocr:
                        # Try OCR and compare results
                        try:
                            logger.info("🔍 Starting Google Vision OCR comparison...")
                            ocr_comparison_start = asyncio.get_event_loop().time()
                            
                            ocr_text, ocr_metadata = await google_vision_ocr_service.process_document(
                                file_content, filename
                            )
                            
                            ocr_comparison_end = asyncio.get_event_loop().time()
                            ocr_comparison_duration = ocr_comparison_end - ocr_comparison_start
                            logger.info(f"🔍 OCR comparison took {ocr_comparison_duration:.2f}s")
                            
                            # Store OCR metadata
                            self._last_ocr_metadata = ocr_metadata
                            
                            if ocr_metadata.get('ocr_required', False):
                                ocr_length = len(ocr_text.strip())
                                logger.info(f"🔍 OCR yielded {ocr_length} additional characters")
                                
                                # If OCR extracted significantly more content, use it
                                if ocr_length > text_length * 1.5:
                                    logger.info("✅ OCR extracted significantly more content - using OCR result")
                                    return ocr_text
                                elif ocr_length > 100:
                                    # Combine traditional text with OCR text
                                    logger.info("🔄 Combining traditional text with OCR results")
                                    combined_text = f"{traditional_text}\n\n[OCR_EXTRACTED_CONTENT]\n{ocr_text}"
                                    return combined_text
                        except Exception as ocr_error:
                            logger.warning(f"⚠️ OCR fallback failed: {ocr_error}")
                    
                    # Return traditional text if OCR didn't help
                    return traditional_text
                
            except Exception as traditional_error:
                logger.warning(f"⚠️ Traditional PDF extraction failed: {traditional_error}")
                # Don't re-raise - continue to OCR fallback
                traditional_text = ""
            
            # Step 2: Traditional extraction failed or yielded minimal text - try OCR
            if not self.ocr_enabled:
                logger.error("❌ PDF appears to be scanned but Google Vision OCR is disabled")
                raise ValueError("Document appears to be scanned/image-based but Google Vision OCR is not available")
            
            logger.info("🔍 Traditional extraction insufficient - attempting OCR processing")
            
            try:
                # Use Google Vision OCR processing with timeout protection
                logger.info("⏰ Starting Google Vision OCR processing with 5-minute timeout...")
                ocr_start_time = asyncio.get_event_loop().time()
                
                ocr_text, ocr_metadata = await asyncio.wait_for(
                    google_vision_ocr_service.process_document(file_content, filename),
                    timeout=300.0  # 5-minute timeout for entire OCR process
                )
                
                ocr_end_time = asyncio.get_event_loop().time()
                ocr_duration = ocr_end_time - ocr_start_time
                logger.info(f"✅ OCR processing completed in {ocr_duration:.2f}s")
                
                # Store OCR metadata for later use in document metadata
                self._last_ocr_metadata = ocr_metadata
                
                if not ocr_metadata.get('ocr_required', False):
                    logger.warning("⚠️ Google Vision OCR service determined no OCR needed, but traditional extraction failed")
                    return "[DOCUMENT_PROCESSED] This document appears to contain scanned content but could not be processed with Google Vision OCR. Please verify the document quality and try again."
                
                ocr_length = len(ocr_text.strip())
                logger.info(f"✅ Google Vision OCR extraction completed - {ocr_length} characters extracted")
                
                # Log Google Vision OCR processing details
                if 'ocr_result' in ocr_metadata:
                    ocr_result = ocr_metadata['ocr_result']
                    logger.info(f"📊 Google Vision OCR Details: Method={ocr_result.get('method_used', 'google_vision')}, "
                               f"Confidence={ocr_result.get('confidence', 0):.2f}, "
                               f"Pages={ocr_result.get('page_count', 1)}, "
                               f"Time={ocr_result.get('processing_time', 0):.2f}s")
                
                if ocr_length < 20:
                    logger.warning("⚠️ OCR extraction yielded minimal text")
                    # Still return what we got rather than failing
                    return ocr_text or "[DOCUMENT_PROCESSED] This document was processed but minimal text was extracted. The document may be primarily images or very low quality scans."
                
                return ocr_text
                
            except asyncio.TimeoutError:
                logger.error("❌ Google Vision OCR processing timed out after 5 minutes")
                # Return empty string to prevent error message chunks
                raise ValueError("Document processing timed out during Google Vision OCR. This document may be too large or complex for processing.")
            
            except Exception as ocr_error:
                logger.error(f"❌ Google Vision OCR processing failed: {ocr_error}")
                return "[DOCUMENT_PROCESSED] This document could not be processed due to Google Vision OCR errors. Please verify the document format and quality."
            
        except Exception as e:
            logger.error(f"❌ Intelligent PDF processing failed for {filename}: {e}")
            raise
            raise
    
    async def _extract_text_from_image(self, file_content: bytes, filename: str) -> str:
        """Extract text from image files using OCR"""
        try:
            if not self.ocr_enabled:
                raise ValueError("Image files require Google Vision OCR but OCR is disabled")
            
            logger.info(f"🖼️ Processing image file {filename} with Google Vision OCR")
            
            # Use Google Vision OCR service to process the image
            ocr_text, ocr_metadata = await google_vision_ocr_service.process_document(
                file_content, filename
            )
            
            # Store OCR metadata for later use
            self._last_ocr_metadata = ocr_metadata
            
            text_length = len(ocr_text.strip())
            logger.info(f"✅ OCR extracted {text_length} characters from image {filename}")
            
            # Log Google Vision OCR processing details
            if ocr_metadata:
                logger.info(f"📊 Google Vision OCR Details: Method={ocr_metadata.get('method', 'google_vision')}, "
                           f"Time={ocr_metadata.get('total_processing_time', 0):.2f}s")
            
            return ocr_text
            
        except Exception as e:
            logger.error(f"❌ Failed to extract text from image {filename}: {e}")
            raise
    
    def _extract_with_pdfplumber(self, pdf_file: BytesIO) -> str:
        """Extract text using pdfplumber (best for layout preservation)"""
        try:
            if not PDFPLUMBER_AVAILABLE:
                logger.warning("⚠️ pdfplumber not available, skipping this extraction method")
                return ""
                
            text_parts = []
            
            with pdfplumber.open(pdf_file) as pdf:
                logger.info(f"📄 PDF has {len(pdf.pages)} pages")
                
                for page_num, page in enumerate(pdf.pages):
                    try:
                        # Extract text with layout preservation
                        page_text = page.extract_text(
                            x_tolerance=3,
                            y_tolerance=3,
                            layout=True,
                            x_density=7.25,
                            y_density=13
                        )
                        
                        if page_text and page_text.strip():
                            # Validate text quality before processing
                            if self._is_text_quality_good(page_text):
                                # Clean up spacing and formatting
                                cleaned_text = re.sub(r'\n\s*\n', '\n\n', page_text)
                                cleaned_text = re.sub(r' +', ' ', cleaned_text)
                                text_parts.append(cleaned_text)
                                logger.debug(f"📄 Extracted {len(page_text)} chars from page {page_num + 1}")
                            else:
                                logger.warning(f"⚠️ Poor text quality detected on page {page_num + 1}, skipping")
                        
                        # Also try table extraction for structured data
                        tables = page.extract_tables()
                        for table in tables:
                            if table:
                                table_text = "\n".join([" | ".join([str(cell) if cell else "" for cell in row]) for row in table])
                                if table_text.strip():
                                    text_parts.append(f"\n[TABLE DATA]\n{table_text}\n[/TABLE DATA]\n")
                        
                    except Exception as e:
                        logger.warning(f"⚠️ Failed to extract from page {page_num + 1} with pdfplumber: {e}")
                        continue
            
            if text_parts:
                full_text = "\n\n".join(text_parts)
                logger.info(f"✅ pdfplumber extracted {len(full_text)} characters from {len(text_parts)} pages")
                return full_text
            
            return ""
            
        except Exception as e:
            logger.warning(f"⚠️ pdfplumber extraction failed: {e}")
            return ""
    
    def _extract_with_pymupdf(self, pdf_file: BytesIO) -> str:
        """Extract text using PyMuPDF (good for most PDFs)"""
        try:
            if not PYMUPDF_AVAILABLE:
                logger.warning("⚠️ PyMuPDF not available, skipping this extraction method")
                return ""
                
            text_parts = []
            
            doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
            logger.info(f"📄 PDF has {len(doc)} pages")
            
            for page_num in range(len(doc)):
                try:
                    page = doc.load_page(page_num)
                    
                    # Extract text with layout preservation
                    text_dict = page.get_text("dict")
                    page_text = self._process_pymupdf_blocks(text_dict)
                    
                    if page_text and page_text.strip():
                        text_parts.append(page_text)
                        logger.debug(f"📄 Extracted {len(page_text)} chars from page {page_num + 1}")
                    
                except Exception as e:
                    logger.warning(f"⚠️ Failed to extract from page {page_num + 1} with PyMuPDF: {e}")
                    continue
            
            doc.close()
            
            if text_parts:
                full_text = "\n\n".join(text_parts)
                logger.info(f"✅ PyMuPDF extracted {len(full_text)} characters from {len(text_parts)} pages")
                return full_text
            
            return ""
            
        except Exception as e:
            logger.warning(f"⚠️ PyMuPDF extraction failed: {e}")
            return ""
    
    def _process_pymupdf_blocks(self, text_dict: dict) -> str:
        """Process PyMuPDF text blocks to preserve layout"""
        try:
            text_parts = []
            
            for block in text_dict.get("blocks", []):
                if "lines" in block:  # Text block
                    block_text = []
                    for line in block["lines"]:
                        line_text = []
                        for span in line.get("spans", []):
                            if "text" in span:
                                text = span["text"].strip()
                                if text:
                                    line_text.append(text)
                        if line_text:
                            block_text.append(" ".join(line_text))
                    
                    if block_text:
                        text_parts.append("\n".join(block_text))
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            logger.warning(f"⚠️ Failed to process PyMuPDF blocks: {e}")
            return ""
    
    def _extract_with_pdfminer(self, pdf_file: BytesIO) -> str:
        """Extract text using pdfminer (good for complex layouts)"""
        try:
            if not PDFMINER_AVAILABLE:
                logger.warning("⚠️ pdfminer not available, skipping this extraction method")
                return ""
            
            # Configure layout analysis parameters
            laparams = LAParams(
                char_margin=2.0,
                line_margin=0.5,
                word_margin=0.1,
                boxes_flow=0.5,
                detect_vertical=True,
                all_texts=False
            )
            
            text = pdfminer_extract_text(
                pdf_file,
                laparams=laparams
            )
            
            if text and text.strip():
                logger.info(f"✅ pdfminer extracted {len(text)} characters")
                return text
            
            return ""
            
        except Exception as e:
            logger.warning(f"⚠️ pdfminer extraction failed: {e}")
            return ""
    
    def _extract_with_pypdf2(self, pdf_file: BytesIO) -> str:
        """Extract text using PyPDF2 (basic fallback)"""
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            if len(pdf_reader.pages) == 0:
                return ""
            
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text)
                        logger.debug(f"📄 PyPDF2 extracted text from page {page_num + 1}")
                except Exception as e:
                    logger.warning(f"⚠️ Failed to extract from page {page_num + 1} with PyPDF2: {e}")
                    continue
            
            if text_parts:
                full_text = "\n\n".join(text_parts)
                logger.info(f"✅ PyPDF2 extracted {len(full_text)} characters from {len(text_parts)} pages")
                return full_text
            
            return ""
            
        except Exception as e:
            logger.warning(f"⚠️ PyPDF2 extraction failed: {e}")
            return ""
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        try:
            if not text:
                return ""
            
            # Normalize unicode characters
            text = unicodedata.normalize('NFKD', text)
            
            # Remove or replace problematic characters
            text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f]', ' ', text)
            
            # Enhanced fix for common OCR errors and encoding issues
            replacements = {
                # Smart quotes and dashes - MORE COMPREHENSIVE
                'â€™': "'", 'â€˜': "'", ''': "'", ''': "'",
                'â€œ': '"', 'â€\x9d': '"', '"': '"', '"': '"',
                'â€"': '—', 'â€"': '–', '–': '-', '—': '-',
                'â€¦': '...', '…': '...',
                
                # Non-breaking spaces and similar issues - EXPANDED
                'Â ': ' ', 'Â': ' ', '\xa0': ' ', '\u00a0': ' ',
                '\u2000': ' ', '\u2001': ' ', '\u2002': ' ', '\u2003': ' ',
                '\u2004': ' ', '\u2005': ' ', '\u2006': ' ', '\u2007': ' ',
                '\u2008': ' ', '\u2009': ' ', '\u200a': ' ', '\u200b': '',
                
                # Common bullet points and symbols
                'â€¢': '•', 'â—': '•', '•': '•',
                'â–ª': '▪', 'â–«': '▫', '▪': '▪', '▫': '▫',
                'â€º': '›', 'â€¹': '‹', '›': '>', '‹': '<',
                
                # Currency and special symbols
                'â‚¬': '€', 'Â£': '£', 'Â¥': '¥',
                'Â®': '®', 'Â©': '©', 'â„¢': '™',
                '®': '®', '©': '©', '™': 'TM',
                
                # Accented characters (common encoding issues) - EXPANDED
                'Ã¡': 'á', 'Ã©': 'é', 'Ã­': 'í', 'Ã³': 'ó', 'Ãº': 'ú',
                'Ã ': 'à', 'Ã¨': 'è', 'Ã¬': 'ì', 'Ã²': 'ò', 'Ã¹': 'ù',
                'Ã¢': 'â', 'Ãª': 'ê', 'Ã®': 'î', 'Ã´': 'ô', 'Ã»': 'û',
                'Ã¤': 'ä', 'Ã«': 'ë', 'Ã¯': 'ï', 'Ã¶': 'ö', 'Ã¼': 'ü',
                'Ã±': 'ñ', 'Ã§': 'ç', 'Ã¿': 'ÿ',
                
                # Additional problematic sequences - MUCH MORE COMPREHENSIVE
                'ï¿½': '',    # Replacement character (usually garbled)
                'â–': '-',    # Various dash issues
                'â€': '',     # Common prefix for encoding issues
                'â€ ': ' ',   # Another variant
                'â€\x9c': '"', 'â€\x9d': '"',  # More quote variants
                'â€\x98': "'", 'â€\x99': "'",  # More quote variants
                'â€\x93': '-', 'â€\x94': '-',  # More dash variants
                'â€\xa6': '...', # Ellipsis variant
                
                # Remove complex garbled sequences
                'â€™s': "'s",  # Possessive apostrophe
                'â€œthe': '"the',  # Quote before word
                'â€\x9cthe': '"the',  # Another quote variant
                
                # Table and form artifacts
                '|': ' | ',  # Keep table separators readable
                
                # Remove obviously corrupted sequences
                'â€š': ',', 'â€ž': '"', 'â€°': '%',
                'âˆ': '', 'â•': '', 'â–': '',
                'â—': '', 'â˜': '', 'â™': '',
                'âš': '', 'â›': '', 'âœ': '',
                'â': '', 'âž': '', 'âŸ': '',
            }
            
            # Apply all replacements
            for old, new in replacements.items():
                text = text.replace(old, new)
            
            # Additional aggressive cleaning for remaining artifacts
            # Remove any remaining â€ sequences that weren't caught above
            text = re.sub(r'â€[^\w\s]', '', text)  # Remove â€ followed by special chars
            text = re.sub(r'â€\w{1,2}', '', text)  # Remove â€ followed by 1-2 chars
            
            # Remove sequences that are likely encoding artifacts
            # More aggressive removal of garbled sequences
            text = re.sub(r'[^\w\s.,;:!?()&@#$%^*+=|\\/<>[\]{}"\'`~-]{3,}', ' ', text)
            
            # Remove isolated special characters that might be artifacts
            text = re.sub(r'\b[^\w\s.,;:!?()-]\b', ' ', text)
            
            # Clean up multiple encoding artifacts in sequence
            text = re.sub(r'(â€|Â|ï¿½|â–|â‚¬|â„¢|Ã){2,}', ' ', text)
            
            # Remove excessive whitespace while preserving paragraph breaks
            text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
            text = re.sub(r'\n[ \t]+', '\n', text)  # Remove leading whitespace on lines
            text = re.sub(r'[ \t]+\n', '\n', text)  # Remove trailing whitespace on lines
            text = re.sub(r'\n{3,}', '\n\n', text)  # Multiple newlines to double newline
            
            # Remove page numbers and headers/footers that might interfere
            lines = text.split('\n')
            cleaned_lines = []
            
            for line in lines:
                line = line.strip()
                
                # Skip likely page numbers (standalone numbers)
                if re.match(r'^\d+$', line) and len(line) <= 3:
                    continue
                
                # Skip lines that are just page markers
                if re.match(r'^(page|página)\s*\d+', line.lower()):
                    continue
                
                # Skip very short lines that might be artifacts (but keep field labels with colons)
                if len(line) < 3 and ':' not in line:
                    continue
                
                cleaned_lines.append(line)
            
            # Rejoin and do final cleanup
            text = '\n'.join(cleaned_lines)
            text = text.strip()
            
            # Log cleaning results
            logger.debug(f"🧹 Text cleaning: {len(text)} characters after cleanup")
            
            return text
            
        except Exception as e:
            logger.warning(f"⚠️ Text cleaning failed: {e}")
            return text  # Return original text if cleaning fails
    
    def _is_text_quality_good(self, text: str) -> bool:
        """Check if extracted text has good quality (not corrupted/garbled)"""
        try:
            if not text or len(text.strip()) < 5:
                return False
            
            # Special handling for table data - be more lenient
            is_table_data = '[TABLE DATA]' in text or '|' in text or text.count(' | ') > 3
            
            # Check for reasonable ASCII character ratio (less strict for tables)
            printable_chars = sum(1 for c in text if c.isprintable() and ord(c) < 256)
            ascii_ratio = printable_chars / len(text)
            min_ascii_ratio = 0.6 if is_table_data else 0.7  # More lenient for tables
            if ascii_ratio < min_ascii_ratio:
                logger.debug(f"❌ Text quality poor: ASCII ratio {ascii_ratio:.2f} < {min_ascii_ratio}")
                return False
            
            # Check for excessive encoding artifacts (less strict for tables)
            artifact_patterns = [
                'â€', 'â€™', 'â€œ', 'â€\x9d', 'Â', 'ï¿½', 'â–', 'â‚¬',
                'â€˜', 'â€"', 'â€"', 'â€¦', 'â€¢', 'â—', 'â–ª', 'â–«',
                'â€º', 'â€¹', 'Â®', 'Â©', 'â„¢', 'Ã¡', 'Ã©', 'Ã­',
                'Ã³', 'Ãº', 'Ã ', 'Ã¨', 'Ã¬', 'Ã²', 'Ã¹', 'Ã¢',
                'Ãª', 'Ã®', 'Ã´', 'Ã»', 'Ã¤', 'Ã«', 'Ã¯', 'Ã¶',
                'Ã¼', 'Ã±', 'Ã§', 'Ã¿', 'âˆ', 'â•', 'â–', 'â—',
                'â˜', 'â™', 'âš', 'â›', 'âœ', 'â', 'âž', 'âŸ'
            ]
            artifact_count = sum(text.count(pattern) for pattern in artifact_patterns)
            max_artifacts = len(text) / 50 if is_table_data else len(text) / 100  # More lenient for tables
            if artifact_count > max_artifacts:
                logger.debug(f"❌ Text quality poor: {artifact_count} artifacts in {len(text)} chars (max: {max_artifacts:.0f})")
                return False
            
            # Check for reasonable word structure (more lenient)
            words = text.split()
            if len(words) < 2:  # Reduced from 3
                return False
            
            # Skip detailed checks for table data
            if is_table_data:
                logger.debug(f"✅ Table data accepted: {len(text)} chars, {len(words)} words")
                return True
            
            # Check average word length (should be reasonable)
            avg_word_length = sum(len(word.strip('.,;:!?()|')) for word in words) / len(words)
            if avg_word_length < 2.0 or avg_word_length > 15:  # More lenient range
                logger.debug(f"❌ Text quality poor: avg word length {avg_word_length}")
                return False
            
            # Check for too many very short words (likely artifacts) - more lenient
            short_words = sum(1 for word in words if len(word.strip('.,;:!?()|')) <= 1)
            if short_words / len(words) > 0.6:  # Increased threshold
                logger.debug(f"❌ Text quality poor: {short_words}/{len(words)} words are too short")
                return False
            
            # Check for reasonable character distribution - more lenient
            alpha_chars = sum(1 for c in text if c.isalpha())
            if alpha_chars / len(text) < 0.25:  # Reduced from 0.4 - allow more numbers/symbols
                logger.debug(f"❌ Text quality poor: only {alpha_chars/len(text)*100:.1f}% alphabetic")
                return False
            
            # Check for excessive repetitive patterns - more lenient
            repetitive_pattern_count = len(re.findall(r'(.)\1{5,}', text))  # 6+ same chars (was 4+)
            if repetitive_pattern_count > len(text) / 100:  # Increased threshold
                logger.debug(f"❌ Text quality poor: {repetitive_pattern_count} repetitive patterns")
                return False
            
            logger.debug(f"✅ Text quality good: {len(text)} chars, {len(words)} words, {artifact_count} artifacts")
            return True
            
        except Exception as e:
            logger.warning(f"⚠️ Failed to check text quality: {e}")
            return True  # Default to good quality on error
    
    def _extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_file = BytesIO(file_content)
            doc = Document(docx_file)
            
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            if not text_parts:
                raise ValueError("No text content found in DOCX file")
            
            full_text = "\n".join(text_parts)
            logger.info(f"✅ Extracted text from DOCX with {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
            return full_text.strip()
            
        except Exception as e:
            logger.error(f"❌ Failed to extract text from DOCX: {e}")
            raise
    
    def _extract_text_from_csv(self, file_content: bytes) -> str:
        """Extract and structure text from CSV file for optimal chunking"""
        try:
            logger.info("📊 Starting CSV text extraction...")
            
            # Try different encodings to read the CSV
            encodings = ['utf-8', 'utf-8-sig', 'iso-8859-1', 'cp1252']
            df = None
            
            for encoding in encodings:
                try:
                    csv_text = file_content.decode(encoding)
                    
                    # Try different delimiters
                    delimiters = [',', ';', '\t', '|']
                    
                    for delimiter in delimiters:
                        try:
                            df = pd.read_csv(StringIO(csv_text), delimiter=delimiter)
                            
                            # Check if parsing was successful (should have multiple columns)
                            if len(df.columns) > 1 and len(df) > 0:
                                logger.info(f"✅ Successfully parsed CSV with delimiter '{delimiter}' and encoding '{encoding}'")
                                break
                        except:
                            continue
                    
                    if df is not None and len(df.columns) > 1:
                        break
                        
                except UnicodeDecodeError:
                    continue
            
            if df is None or len(df.columns) <= 1:
                raise ValueError("Could not parse CSV file with any supported encoding or delimiter")
            
            # Clean column names
            df.columns = df.columns.str.strip()
            
            # Fill NaN values with empty strings for better processing
            df = df.fillna('')
            
            logger.info(f"📊 CSV parsed successfully: {len(df)} rows, {len(df.columns)} columns")
            logger.info(f"📊 Columns: {list(df.columns)}")
            
            # Structure the CSV data for optimal chunking and retrieval
            text_parts = []
            
            # 1. Add header information as first chunk
            header_info = f"[CSV_HEADER] CSV Dataset Information\n"
            header_info += f"Dataset contains {len(df)} records with {len(df.columns)} columns\n"
            header_info += f"Columns: {', '.join(df.columns)}\n\n"
            
            # 2. Add individual record data - THIS IS THE KEY FIX
            for idx, (_, row) in enumerate(df.iterrows()):
                record_text = f"[CSV_RECORD_{idx+1}] Device Record {idx+1}:\n"
                
                for col in df.columns:
                    value = str(row[col]).strip()
                    if value:  # Only include non-empty values
                        record_text += f"{col}: {value}\n"
                    else:
                        record_text += f"{col}: [EMPTY_FIELD]\n"
                
                # Add the record as a complete chunk
                record_text += f"\nThis is a complete record from the CSV dataset.\n"
                text_parts.append(record_text)
            
            # 3. Add column pattern information for better matching
            for col in df.columns:
                col_data = df[col].dropna().astype(str)
                non_empty_data = col_data[col_data.str.strip() != ''].tolist()
                
                if non_empty_data:
                    col_description = f"[COLUMN_{col.replace(' ', '_').upper()}] Column Information for '{col}':\n"
                    col_description += f"Sample values: {', '.join(non_empty_data)}\n"
                    col_description += f"Total entries: {len(non_empty_data)}\n"
                    col_description += f"Data pattern: {self._detect_csv_column_pattern(non_empty_data)}\n"
                    col_description += f"All values: {'; '.join(non_empty_data)}\n\n"
                    text_parts.append(col_description)
            
            # Combine all parts
            full_text = "\n".join(text_parts)
            
            logger.info(f"✅ CSV text extraction completed: {len(full_text)} characters generated")
            return full_text.strip()
            
        except Exception as e:
            logger.error(f"❌ Failed to extract text from CSV: {e}")
            raise
    
    def _detect_csv_column_pattern(self, values: List[str]) -> str:
        """Detect the data pattern in a CSV column"""
        if not values:
            return "empty"
        
        # Check for common patterns
        patterns = {
            'date': r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}',
            'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            'phone': r'[\+]?[1-9]?[\d\s\-\(\)]{7,15}',
            'number': r'^\d+(\.\d+)?$',
            'currency': r'[\$£€¥]?\d+(\.\d{2})?',
            'percentage': r'\d+(\.\d+)?%',
            'url': r'https?://[^\s]+',
            'id_code': r'^[A-Z0-9\-_]+$',
            'version': r'v?\d+\.\d+(\.\d+)?'
        }
        
        for pattern_name, pattern in patterns.items():
            matches = sum(1 for v in values if re.search(pattern, str(v), re.IGNORECASE))
            if matches > len(values) * 0.5:  # More than 50% match
                return pattern_name
        
        # Check for specific content types
        if any('model' in str(v).lower() for v in values):
            return 'model_number'
        elif any('name' in str(v).lower() for v in values):
            return 'name'
        elif any('address' in str(v).lower() or 'street' in str(v).lower() for v in values):
            return 'address'
        elif any('manufacturer' in str(v).lower() or 'company' in str(v).lower() for v in values):
            return 'manufacturer'
        
        return 'text'
    
    def _create_chunks(self, text: str) -> List[Dict[str, Any]]:
        """Create optimized overlapping chunks from text with better boundary detection and enhanced coverage"""
        try:
            if not text or len(text.strip()) == 0:
                logger.warning("⚠️ Empty or whitespace-only text provided for chunking")
                return []
            
            # Clean and prepare text for chunking
            cleaned_text = self._prepare_text_for_chunking(text)
            
            # ENHANCED: Detect if this is CSV content and use specialized chunking
            if self._is_csv_content(cleaned_text):
                return self._create_csv_aware_chunks(cleaned_text)
            
            chunks = []
            start = 0
            chunk_id = 0
            text_length = len(cleaned_text)
            
            logger.info(f"📊 Creating chunks from text of length {text_length} characters")
            
            # SMART CHUNKING: Adjust parameters based on document length
            if text_length <= 1000:
                # Small documents: create fewer, larger chunks
                chunk_size = max(text_length // 2, 300)  # At most 2 chunks
                overlap = min(chunk_size // 4, 100)  # 25% overlap max
                logger.info(f"� Small document detected - using optimized chunking")
            elif text_length <= 5000:
                # Medium documents: moderate chunking
                chunk_size = 1000
                overlap = 200
                logger.info(f"📄 Medium document detected - using moderate chunking")
            else:
                # Large documents: standard chunking
                chunk_size = self.chunk_size
                overlap = self.chunk_overlap
                logger.info(f"📚 Large document detected - using standard chunking")
            
            logger.info(f"🔧 Chunk size: {chunk_size}, overlap: {overlap}")
            
            # If text is very small, still try to create meaningful chunks
            if text_length <= 500:
                logger.info("📄 Small document - creating optimized chunks")
                # Even for small documents, try to create multiple chunks if there are clear separators
                paragraphs = self._split_by_paragraphs(cleaned_text)
                if len(paragraphs) > 1:
                    logger.info(f"📄 Small document has {len(paragraphs)} paragraphs - creating multiple chunks")
                    chunks = []
                    for i, paragraph in enumerate(paragraphs):
                        if paragraph.strip():
                            chunk_data = {
                                "chunk_id": f"chunk_{i}",
                                "content": paragraph.strip(),
                                "start_index": 0,
                                "end_index": len(paragraph.strip()),
                                "word_count": len(paragraph.strip().split()),
                                "content_type": self._classify_content_type(paragraph.strip()),
                                "has_structured_data": self._contains_structured_data(paragraph.strip()),
                                "contains_fields": self._contains_form_fields(paragraph.strip()),
                                "importance_score": 0.8,
                                "entity_density": self._calculate_entity_density(paragraph.strip()),
                                "information_richness": self._calculate_information_richness(paragraph.strip()),
                                "semantic_keywords": self._extract_semantic_keywords(paragraph.strip()),
                                "position_info": {"relative_position": i / len(paragraphs), "section": f"paragraph_{i}"},
                                "coverage_info": {"covers_start": i == 0, "covers_end": i == len(paragraphs) - 1}
                            }
                            chunks.append(chunk_data)
                    
                    if chunks:
                        logger.info(f"✅ Created {len(chunks)} chunks from small document with paragraphs")
                        return chunks
                
                # Fallback to single chunk for very small content
                chunk_data = {
                    "chunk_id": "chunk_0",
                    "content": cleaned_text,
                    "start_index": 0,
                    "end_index": len(cleaned_text),
                    "word_count": len(cleaned_text.split()),
                    "content_type": self._classify_content_type(cleaned_text),
                    "has_structured_data": self._contains_structured_data(cleaned_text),
                    "contains_fields": self._contains_form_fields(cleaned_text),
                    "importance_score": 0.8,  # Higher score for single chunk
                    "entity_density": self._calculate_entity_density(cleaned_text),
                    "information_richness": self._calculate_information_richness(cleaned_text),
                    "semantic_keywords": self._extract_semantic_keywords(cleaned_text),
                    "position_info": {"relative_position": 0.0, "section": "single"},
                    "coverage_info": {"covers_start": True, "covers_end": True}
                }
                
                logger.info(f"✅ Created 1 optimized chunk from small document")
                return [chunk_data]
            
            # ENHANCED: Create comprehensive chunks with improved coverage
            while start < text_length:
                end = min(start + chunk_size, text_length)
                
                # Find the best boundary for splitting
                chunk_text, actual_end = self._find_optimal_chunk_boundary(
                    cleaned_text, start, end, text_length
                )
                
                # Only add meaningful chunks
                chunk_content = chunk_text.strip()
                if self._is_valid_chunk(chunk_content):
                    # Enhance chunk with metadata
                    chunk_metadata = self._extract_chunk_metadata(chunk_content)
                    
                    # ENHANCED: Better context preservation and keyword extraction
                    enhanced_metadata = self._enhance_chunk_metadata(chunk_content, chunk_id, start, actual_end)
                    chunk_metadata.update(enhanced_metadata)
                    
                    chunks.append({
                        "chunk_id": chunk_id,
                        "content": chunk_content,
                        "start_index": start,
                        "end_index": actual_end,
                        "word_count": len(chunk_content.split()),
                        "has_structured_data": chunk_metadata["has_structured_data"],
                        "contains_fields": chunk_metadata["contains_fields"],
                        "content_type": chunk_metadata["content_type"],
                        "importance_score": chunk_metadata.get("importance_score", 0.5),
                        "semantic_keywords": chunk_metadata.get("semantic_keywords", []),
                        "entity_density": chunk_metadata.get("entity_density", 0.0),
                        "information_richness": chunk_metadata.get("information_richness", 0.0),
                        "chunk_quality_score": chunk_metadata.get("chunk_quality_score", 0.5),
                        "coverage_info": {
                            "chunk_position": f"{chunk_id}/{chunk_id}",  # Will be updated later
                            "document_coverage": f"{start}-{actual_end}",
                            "total_length": len(cleaned_text)
                        }
                    })
                    logger.debug(f"📦 Created chunk {chunk_id}: {len(chunk_content)} chars, type: {chunk_metadata['content_type']}, importance: {chunk_metadata.get('importance_score', 0.5):.2f}")
                    chunk_id += 1
                elif chunk_content:
                    logger.debug(f"⏭️ Skipped invalid chunk ({len(chunk_content)} chars): {chunk_content[:50]}...")
                
                # Calculate next start position with smart overlap
                next_start = self._calculate_next_start(
                    start, actual_end, cleaned_text, self.chunk_overlap
                )
                
                # FIXED: Prevent character-by-character sliding window
                if next_start <= start:
                    # If overlap calculation fails, advance by at least chunk_size - overlap
                    # This ensures we make meaningful progress
                    min_advance = max(chunk_size - overlap, 200)
                    next_start = start + min_advance
                else:
                    start = next_start
                
                if start >= text_length:
                    break
            
            # ENHANCED: Post-process chunks for better coverage
            enhanced_chunks = self._post_process_chunks(chunks, cleaned_text)
            
            logger.info(f"✅ Created {len(enhanced_chunks)} enhanced chunks from document")
            
            if len(enhanced_chunks) == 0:
                logger.error(f"❌ ZERO CHUNKS CREATED! Text length: {text_length}")
                logger.error(f"❌ First 500 chars of text: {cleaned_text[:500]}")
                logger.error(f"❌ Text is all whitespace: {cleaned_text.isspace()}")
                logger.error(f"❌ Text stripped length: {len(cleaned_text.strip())}")
            
            return enhanced_chunks
            
        except Exception as e:
            logger.error(f"❌ Failed to create chunks: {e}")
            raise
    
    def _is_csv_content(self, text: str) -> bool:
        """Detect if the text content is from a CSV file"""
        csv_indicators = [
            "[CSV_HEADER]", "[CSV_RECORD_", "[COLUMN_", 
            "Dataset contains", "records with", "columns"
        ]
        return any(indicator in text for indicator in csv_indicators)
    
    def _create_csv_aware_chunks(self, text: str) -> List[Dict[str, Any]]:
        """Create chunks specifically optimized for CSV content"""
        logger.info("📊 Using CSV-aware chunking strategy")
        
        chunks = []
        chunk_id = 0
        
        # Split text into logical sections
        sections = text.split("\n\n")
        
        for section in sections:
            section = section.strip()
            if not section:
                continue
                
            # Determine section type
            if section.startswith("[CSV_HEADER]"):
                content_type = "csv_header"
                importance_score = 0.9
            elif section.startswith("[CSV_RECORD_"):
                content_type = "csv_record"
                importance_score = 1.0
            elif section.startswith("[COLUMN_"):
                content_type = "csv_column_info"
                importance_score = 0.8
            else:
                content_type = "csv_general"
                importance_score = 0.6
            
            # Extract metadata for this chunk
            chunk_metadata = self._extract_chunk_metadata(section)
            
            chunks.append({
                "chunk_id": chunk_id,
                "content": section,
                "start_index": 0,  # Not critical for CSV
                "end_index": len(section),
                "word_count": len(section.split()),
                "has_structured_data": True,
                "contains_fields": ":" in section,
                "content_type": content_type,
                "importance_score": importance_score,
                "semantic_keywords": chunk_metadata.get("semantic_keywords", []),
                "entity_density": chunk_metadata.get("entity_density", 0.0),
                "information_richness": importance_score,
                "chunk_quality_score": importance_score,
                "coverage_info": {
                    "chunk_position": f"{chunk_id}/{len(sections)}",
                    "document_coverage": f"section_{chunk_id}",
                    "total_length": len(text)
                }
            })
            
            logger.debug(f"📦 Created CSV chunk {chunk_id}: {len(section)} chars, type: {content_type}")
            chunk_id += 1
        
        logger.info(f"✅ Created {len(chunks)} CSV-aware chunks")
        return chunks

    def _prepare_text_for_chunking(self, text: str) -> str:
        """Prepare text for optimal chunking"""
        try:
            # Normalize line breaks
            text = re.sub(r'\r\n', '\n', text)
            text = re.sub(r'\r', '\n', text)
            
            # Preserve important formatting markers
            text = re.sub(r'\n\n+', '\n\n', text)  # Multiple newlines to double
            
            # Mark important sections for better chunking
            # Mark field labels (important for form processing)
            text = re.sub(r'^([A-Za-z][A-Za-z\s]*:)\s*$', r'\1 [FIELD_LABEL]', text, flags=re.MULTILINE)
            
            # Mark structured data patterns
            text = re.sub(r'(\[TABLE DATA\].*?\[/TABLE DATA\])', r'\1 [STRUCTURED_CONTENT]', text, flags=re.DOTALL)
            
            return text.strip()
            
        except Exception as e:
            logger.warning(f"⚠️ Failed to prepare text for chunking: {e}")
            return text
    
    def _find_optimal_chunk_boundary(self, text: str, start: int, end: int, text_length: int) -> tuple:
        """Find the optimal boundary for chunk splitting"""
        try:
            if end >= text_length:
                return text[start:end], end
            
            chunk_text = text[start:end]
            
            # Look for natural boundaries in order of preference
            boundaries = [
                (r'\n\n', 2),  # Paragraph breaks (highest priority)
                (r'\.\s+[A-Z]', 2),  # Sentence endings followed by capital letters
                (r'\.\n', 2),  # Sentence endings at line breaks
                (r'\n', 1),  # Line breaks
                (r'\.\s', 2),  # Sentence endings
                (r';\s', 2),  # Semicolons
                (r',\s', 2),  # Commas
                (r'\s', 1),  # Any whitespace
            ]
            
            # Find the best boundary within the last 25% of the chunk
            search_start = max(start + int(self.chunk_size * 0.75), start + self.chunk_size // 2)
            search_text = text[search_start:end]
            
            for pattern, offset in boundaries:
                matches = list(re.finditer(pattern, search_text))
                if matches:
                    # Use the last match (closest to end)
                    last_match = matches[-1]
                    boundary_pos = search_start + last_match.start() + offset
                    return text[start:boundary_pos], boundary_pos
            
            # No good boundary found, use original end
            return chunk_text, end
            
        except Exception as e:
            logger.warning(f"⚠️ Failed to find optimal boundary: {e}")
            return text[start:end], end
    
    def _is_valid_chunk(self, chunk_content: str) -> bool:
        """Determine if a chunk is valid and worth storing"""
        try:
            if not chunk_content or len(chunk_content.strip()) < 10:  # Reduced from 20
                return False
            
            # Check for special content types that should be preserved
            is_table_data = '[TABLE DATA]' in chunk_content or chunk_content.count('|') > 3
            is_structured_content = '[STRUCTURED_CONTENT]' in chunk_content
            has_field_markers = ':' in chunk_content and any(field in chunk_content.lower() for field in ['name', 'number', 'date', 'model', 'manufacturer', 'format', 'effective'])
            
            # Be very lenient with table data and structured content
            if is_table_data or is_structured_content or has_field_markers:
                logger.debug(f"✅ Accepting structured content: table={is_table_data}, structured={is_structured_content}, fields={has_field_markers}")
                return True
            
            # Check for minimum word count - more lenient
            words = chunk_content.split()
            if len(words) < 3:  # Reduced from 5
                return False
            
            # Avoid chunks that are mostly special characters or numbers - more lenient
            alphanumeric_ratio = sum(c.isalnum() for c in chunk_content) / len(chunk_content)
            if alphanumeric_ratio < 0.3:  # Reduced from 0.5
                return False
            
            # Check for reasonable character distribution (avoid garbled text) - more lenient
            ascii_ratio = sum(1 for c in chunk_content if ord(c) < 128) / len(chunk_content)
            if ascii_ratio < 0.7:  # Reduced from 0.8
                return False
            
            # Avoid chunks with too many special encoding characters - more lenient
            encoding_artifacts = ['â€', 'Â', 'ï¿½', 'â–', 'â€œ', 'â€\x9d']
            artifact_count = sum(chunk_content.count(artifact) for artifact in encoding_artifacts)
            if artifact_count > 5:  # Increased from 3
                return False
            
            # Check for reasonable word length distribution - more lenient
            if words:
                avg_word_length = sum(len(word) for word in words) / len(words)
                if avg_word_length < 1.5 or avg_word_length > 20:  # More lenient range
                    return False
                
                # Check for too many very short or very long words - more lenient
                short_words = sum(1 for word in words if len(word) <= 2)
                long_words = sum(1 for word in words if len(word) > 25)
                if short_words / len(words) > 0.8 or long_words / len(words) > 0.15:  # More lenient
                    return False
            
            # Avoid chunks that are just repeated characters - more lenient
            unique_chars = len(set(chunk_content.lower().replace(' ', '').replace('\n', '')))
            if unique_chars < 5:  # Reduced from 10
                return False
            
            # Check for reasonable sentence structure - more lenient
            sentences = re.split(r'[.!?]+', chunk_content)
            if len(sentences) > 1:
                avg_sentence_length = sum(len(s.split()) for s in sentences if s.strip()) / max(1, len([s for s in sentences if s.strip()]))
                if avg_sentence_length < 2 or avg_sentence_length > 150:  # More lenient range
                    return False
            
            return True
            
        except Exception as e:
            logger.warning(f"⚠️ Failed to validate chunk: {e}")
            return True  # Default to valid on error
    
    def _extract_chunk_metadata(self, chunk_content: str) -> Dict[str, Any]:
        """Extract metadata about chunk content for better processing"""
        try:
            metadata = {
                "has_structured_data": False,
                "contains_fields": False,
                "content_type": "text"
            }
            
            # Check for structured data
            if any(marker in chunk_content for marker in ['[TABLE DATA]', '[STRUCTURED_CONTENT]']):
                metadata["has_structured_data"] = True
                metadata["content_type"] = "structured"
            
            # Check for form fields
            if re.search(r'.*:\s*[_\[\{]|.*:\s*$', chunk_content, re.MULTILINE):
                metadata["contains_fields"] = True
                if metadata["content_type"] == "text":
                    metadata["content_type"] = "form"
            
            # Check for lists or enumerations
            if re.search(r'^\s*[\d\w]\.\s', chunk_content, re.MULTILINE):
                if metadata["content_type"] == "text":
                    metadata["content_type"] = "list"
            
            return metadata
            
        except Exception as e:
            logger.warning(f"⚠️ Failed to extract chunk metadata: {e}")
            return {"has_structured_data": False, "contains_fields": False, "content_type": "text"}
    
    def _calculate_next_start(self, current_start: int, current_end: int, text: str, overlap: int) -> int:
        """Calculate the next start position with smart overlap"""
        try:
            basic_next_start = max(current_end - overlap, current_start + 1)
            
            # Try to start at a natural boundary within the overlap region
            search_start = max(basic_next_start - 50, current_start + 1)
            search_end = min(basic_next_start + 50, len(text))
            search_text = text[search_start:search_end]
            
            # Look for paragraph or sentence boundaries
            boundaries = [r'\n\n', r'\.\s+[A-Z]', r'\n', r'\.\s']
            
            for pattern in boundaries:
                matches = list(re.finditer(pattern, search_text))
                if matches:
                    # Find the match closest to our basic start position
                    target_pos = basic_next_start - search_start
                    best_match = min(matches, key=lambda m: abs(m.start() - target_pos))
                    return search_start + best_match.end()
            
            return basic_next_start
            
        except Exception as e:
            logger.warning(f"⚠️ Failed to calculate smart overlap: {e}")
            return max(current_end - overlap, current_start + 1)
    
    async def _store_chunks_in_pinecone(
        self, 
        chunks: List[Dict[str, Any]], 
        document_id: str, 
        device_id: str, 
        filename: str
    ):
        """Generate embeddings and store chunks in Pinecone with enhanced metadata"""
        try:
            logger.info(f"🔗 Storing {len(chunks)} chunks in vector database for document {filename}")
            
            vectors = []
            
            for i, chunk in enumerate(chunks):
                try:
                    logger.info(f"🔄 Processing chunk {i+1}/{len(chunks)}")
                    
                    # Prepare text for embedding (clean version)
                    embedding_text = self._prepare_text_for_embedding(chunk["content"])
                    
                    # Generate embedding for cleaned chunk with timeout
                    logger.debug(f"📡 Generating embedding for chunk {i+1}...")
                    start_time = asyncio.get_event_loop().time()
                    
                    try:
                        embedding = await asyncio.wait_for(
                            gemini_service.get_embedding(embedding_text),
                            timeout=10.0  # 10-second timeout per embedding
                        )
                        
                        embed_time = asyncio.get_event_loop().time() - start_time
                        logger.debug(f"✅ Embedding generated in {embed_time:.2f}s for chunk {i+1}")
                        
                    except asyncio.TimeoutError:
                        logger.warning(f"⏰ Embedding generation timed out for chunk {i+1}, using fallback")
                        embedding = self._generate_fallback_embedding(embedding_text)
                    except Exception as embed_error:
                        logger.warning(f"❌ Embedding generation failed for chunk {i+1}: {embed_error}, using fallback")
                        embedding = self._generate_fallback_embedding(embedding_text)
                    
                    # Create enhanced metadata
                    metadata = {
                        "document_id": document_id,
                        "chunk_id": chunk["chunk_id"],
                        "content": chunk["content"][:2000],  # Increased storage for better context
                        "filename": filename,
                        "device_id": device_id,
                        "start_index": chunk["start_index"],
                        "end_index": chunk["end_index"],
                        "word_count": chunk.get("word_count", 0),
                        "content_type": chunk.get("content_type", "text"),
                        "has_structured_data": chunk.get("has_structured_data", False),
                        "contains_fields": chunk.get("contains_fields", False),
                        "text_length": len(chunk["content"]),
                        "extraction_quality": self._assess_extraction_quality(chunk["content"]),
                        # ENHANCED: More comprehensive metadata for better retrieval
                        "importance_score": chunk.get("importance_score", 0.5),
                        "entity_density": chunk.get("entity_density", 0.0),
                        "information_richness": chunk.get("information_richness", 0.0),
                        "semantic_keywords": ' '.join(chunk.get("semantic_keywords", [])),
                        "position_info": json.dumps(chunk.get("position_info", {})),
                        "coverage_info": json.dumps(chunk.get("coverage_info", {})),
                        # Add searchable keywords for better retrieval
                        "keywords": self._extract_keywords(chunk["content"]),
                        "has_numbers": bool(re.search(r'\d', chunk["content"])),
                        "has_dates": bool(re.search(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}', chunk["content"])),
                        "has_technical_terms": self._has_technical_terms(chunk["content"]),
                        "has_form_fields": bool(re.search(r'[A-Za-z\s]+:\s*(?:$|_|\.\.\.)', chunk["content"])),
                        "chunk_quality_score": self._calculate_chunk_quality_score(chunk["content"])
                    }
                    
                    # Create vector with enhanced metadata
                    vector = {
                        "id": f"{document_id}_{chunk['chunk_id']}",
                        "values": embedding,
                        "metadata": metadata
                    }
                    vectors.append(vector)
                    
                    if (i + 1) % 10 == 0:
                        logger.debug(f"📊 Processed {i + 1}/{len(chunks)} embeddings")
                        
                except Exception as e:
                    logger.error(f"❌ Failed to create embedding for chunk {i}: {e}")
                    # Continue with other chunks instead of failing completely
                    continue
            
            # Store in Pinecone
            if vectors:
                await pinecone_service.upsert_vectors(vectors, device_id)
                logger.info(f"✅ Successfully stored {len(vectors)} vectors in Pinecone for device {device_id}")
                
                # Log quality statistics
                avg_quality = sum(v["metadata"]["extraction_quality"] for v in vectors) / len(vectors)
                structured_count = sum(1 for v in vectors if v["metadata"]["has_structured_data"])
                field_count = sum(1 for v in vectors if v["metadata"]["contains_fields"])
                
                logger.info(f"📊 Quality stats - Avg: {avg_quality:.2f}, Structured: {structured_count}, Fields: {field_count}")
            else:
                raise ValueError("No vectors were created for storage")
            
        except Exception as e:
            logger.error(f"❌ Failed to store chunks in Pinecone: {e}")
            raise
    
    def _prepare_text_for_embedding(self, text: str) -> str:
        """Prepare text for embedding generation"""
        try:
            # Remove embedding markers we added for chunking
            text = re.sub(r'\s*\[FIELD_LABEL\]', '', text)
            text = re.sub(r'\s*\[STRUCTURED_CONTENT\]', '', text)
            
            # Clean up excessive whitespace
            text = re.sub(r'\s+', ' ', text)
            text = text.strip()
            
            return text
            
        except Exception as e:
            logger.warning(f"⚠️ Failed to prepare text for embedding: {e}")
            return text
    
    def _assess_extraction_quality(self, text: str) -> float:
        """Assess the quality of extracted text (0.0 to 1.0)"""
        try:
            if not text:
                return 0.0
            
            quality_score = 1.0
            
            # Check for encoding issues
            if any(char in text for char in ['â€™', 'â€œ', 'â€\x9d', 'Â ']):
                quality_score -= 0.2
            
            # Check for excessive special characters
            special_char_ratio = sum(1 for c in text if not c.isalnum() and c not in ' \n\t.,;:!?-()[]{}') / len(text)
            if special_char_ratio > 0.3:
                quality_score -= 0.3
            
            # Check for reasonable word distribution
            words = text.split()
            if words:
                avg_word_length = sum(len(word) for word in words) / len(words)
                if avg_word_length < 2 or avg_word_length > 15:
                    quality_score -= 0.2
            
            # Check for readable sentence structure
            sentences = re.split(r'[.!?]+', text)
            if sentences:
                avg_sentence_length = sum(len(sentence.split()) for sentence in sentences) / len(sentences)
                if avg_sentence_length < 3 or avg_sentence_length > 50:
                    quality_score -= 0.1
            
            return max(0.0, min(1.0, quality_score))
            
        except Exception as e:
            logger.warning(f"⚠️ Failed to assess extraction quality: {e}")
            return 0.5  # Default to medium quality
    
    def _extract_keywords(self, text: str) -> str:
        """Extract important keywords from text for better searchability"""
        try:
            # Convert to lowercase for processing
            text_lower = text.lower()
            
            # Common important terms in medical/technical documents
            important_terms = {
                # Medical device terms
                'device', 'medical', 'equipment', 'instrument', 'apparatus',
                'monitor', 'sensor', 'probe', 'catheter', 'implant',
                'diagnosis', 'treatment', 'therapy', 'procedure',
                
                # Document terms
                'specification', 'manual', 'guide', 'instruction', 'protocol',
                'standard', 'requirement', 'compliance', 'validation',
                'certificate', 'approval', 'registration', 'license',
                
                # Technical terms
                'model', 'serial', 'version', 'revision', 'configuration',
                'parameter', 'setting', 'calibration', 'measurement',
                'accuracy', 'precision', 'range', 'limit', 'threshold',
                
                # Company/regulatory terms
                'manufacturer', 'supplier', 'vendor', 'distributor',
                'fda', 'ce', 'iso', 'iec', 'astm', 'ansi',
                'regulation', 'directive', 'standard', 'guideline'
            }
            
            # Find matching terms
            found_terms = []
            for term in important_terms:
                if term in text_lower:
                    found_terms.append(term)
            
            # Also extract potential model numbers, document numbers, etc.
            numbers = re.findall(r'\b[A-Z0-9]{2,}[-]?[A-Z0-9]*\b', text)
            found_terms.extend(numbers[:5])  # Limit to first 5 numbers
            
            return ' '.join(found_terms[:10])  # Limit to first 10 keywords
            
        except Exception as e:
            logger.warning(f"⚠️ Failed to extract keywords: {e}")
            return ""
    
    def _has_technical_terms(self, text: str) -> bool:
        """Check if text contains technical terms"""
        try:
            text_lower = text.lower()
            technical_indicators = [
                # Medical terms
                'medical', 'clinical', 'diagnostic', 'therapeutic', 'surgical',
                'patient', 'physician', 'hospital', 'healthcare',
                
                # Technical terms
                'specification', 'parameter', 'calibration', 'measurement',
                'accuracy', 'precision', 'frequency', 'voltage', 'current',
                'temperature', 'pressure', 'humidity', 'sterilization',
                
                # Regulatory terms
                'compliance', 'validation', 'verification', 'regulation',
                'standard', 'requirement', 'guideline', 'protocol',
                
                # Document types
                'manual', 'guide', 'instruction', 'procedure', 'checklist'
            ]
            
            return any(term in text_lower for term in technical_indicators)
            
        except Exception as e:
            logger.warning(f"⚠️ Failed to check technical terms: {e}")
            return False
    
    def _enhance_chunk_metadata(self, chunk_content: str, chunk_id: int, start_index: int, end_index: int) -> Dict[str, Any]:
        """Enhanced metadata extraction for better retrieval"""
        try:
            metadata = {}
            
            # Calculate importance score based on content richness
            metadata["importance_score"] = self._calculate_importance_score(chunk_content)
            
            # Extract semantic keywords for better matching
            metadata["semantic_keywords"] = self._extract_semantic_keywords(chunk_content)
            
            # Calculate entity density (names, numbers, technical terms)
            metadata["entity_density"] = self._calculate_entity_density(chunk_content)
            
            # Calculate information richness
            metadata["information_richness"] = self._calculate_information_richness(chunk_content)
            
            # Calculate chunk quality score
            metadata["chunk_quality_score"] = self._calculate_chunk_quality_score(chunk_content)
            
            # Identify chunk position in document
            metadata["position_info"] = {
                "chunk_id": chunk_id,
                "relative_position": start_index / max(end_index, 1),
                "is_beginning": start_index < 1000,
                "is_middle": 1000 <= start_index <= end_index - 1000,
                "is_end": start_index > end_index - 1000
            }
            
            return metadata
            
        except Exception as e:
            logger.warning(f"⚠️ Failed to enhance chunk metadata: {e}")
            return {}
    
    def _calculate_importance_score(self, text: str) -> float:
        """Calculate importance score based on content indicators"""
        try:
            score = 0.5  # Base score
            text_lower = text.lower()
            
            # High importance indicators
            high_importance_terms = [
                'device name', 'model number', 'serial number', 'manufacturer',
                'document number', 'version', 'date', 'specification',
                'requirements', 'standards', 'compliance', 'approval',
                'certification', 'generic name', 'intended use'
            ]
            
            for term in high_importance_terms:
                if term in text_lower:
                    score += 0.1
            
            # Form field indicators (very important for template filling)
            if ':' in text and any(field in text_lower for field in ['name', 'number', 'date', 'model', 'manufacturer']):
                score += 0.2
            
            # Technical data indicators
            if any(indicator in text for indicator in [':', ';', '(', ')', '[', ']', '{', '}']):
                score += 0.05
            
            # Presence of numbers (often important data)
            import re
            numbers = re.findall(r'\b\d+(?:\.\d+)?\b', text)
            if numbers:
                score += min(len(numbers) * 0.02, 0.15)
            
            # Uppercase abbreviations (often important)
            abbreviations = re.findall(r'\b[A-Z]{2,}\b', text)
            if abbreviations:
                score += min(len(abbreviations) * 0.03, 0.1)
            
            return min(score, 1.0)  # Cap at 1.0
            
        except Exception as e:
            logger.warning(f"⚠️ Failed to calculate importance score: {e}")
            return 0.5
    
    def _extract_semantic_keywords(self, text: str) -> List[str]:
        """Extract semantic keywords for better retrieval"""
        try:
            import re
            
            keywords = set()
            text_lower = text.lower()
            
            # Domain-specific terms
            domain_terms = [
                # Device types
                'pulse oximeter', 'oximeter', 'monitor', 'sensor', 'probe',
                'catheter', 'implant', 'stent', 'pacemaker', 'defibrillator',
                
                # Medical terms
                'medical device', 'diagnostic', 'therapeutic', 'surgical',
                'clinical', 'patient', 'physician', 'hospital',
                
                # Document types
                'dmf', 'device master file', 'specification', 'manual',
                'guide', 'instruction', 'protocol', 'standard',
                
                # Regulatory terms
                'fda', 'ce mark', 'iso', 'iec', 'compliance', 'validation',
                'verification', 'approval', 'certification', 'registration',
                
                # Technical terms
                'model', 'version', 'serial', 'manufacturer', 'supplier',
                'accuracy', 'precision', 'calibration', 'measurement'
            ]
            
            for term in domain_terms:
                if term in text_lower:
                    keywords.add(term)
            
            # Extract numbers and codes (often important identifiers)
            numbers = re.findall(r'\b[A-Z0-9]{2,}[-/]?[A-Z0-9]*\b', text)
            keywords.update(numbers[:5])  # Limit to first 5
            
            # Extract capitalized terms (proper nouns, brands)
            capitalized = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', text)
            keywords.update([term.lower() for term in capitalized[:5]])
            
            return list(keywords)[:10]  # Limit to 10 keywords
            
        except Exception as e:
            logger.warning(f"⚠️ Failed to extract semantic keywords: {e}")
            return []
    
    def _calculate_entity_density(self, text: str) -> float:
        """Calculate density of named entities and important identifiers"""
        try:
            import re
            
            words = text.split()
            if not words:
                return 0.0
            
            entity_count = 0
            
            # Count capitalized words (potential proper nouns)
            entity_count += len(re.findall(r'\b[A-Z][a-z]+\b', text))
            
            # Count numbers
            entity_count += len(re.findall(r'\b\d+\b', text))
            
            # Count codes/identifiers
            entity_count += len(re.findall(r'\b[A-Z0-9]{2,}\b', text))
            
            # Count technical abbreviations
            entity_count += len(re.findall(r'\b[A-Z]{2,}\b', text))
            
            return min(entity_count / len(words), 1.0)
            
        except Exception as e:
            logger.warning(f"⚠️ Failed to calculate entity density: {e}")
            return 0.0
    
    def _calculate_information_richness(self, text: str) -> float:
        """Calculate how information-rich the text is"""
        try:
            # Base richness on various factors
            richness = 0.0
            
            # Sentence structure diversity
            sentences = text.split('.')
            if len(sentences) > 1:
                avg_sentence_length = sum(len(s.split()) for s in sentences) / len(sentences)
                if 5 <= avg_sentence_length <= 25:  # Good sentence length
                    richness += 0.2
            
            # Punctuation diversity (indicates structured content)
            unique_punct = set(c for c in text if c in '.,;:!?()[]{}')
            richness += min(len(unique_punct) * 0.05, 0.3)
            
            # Vocabulary diversity
            words = text.lower().split()
            if words:
                unique_words = len(set(words))
                vocabulary_ratio = unique_words / len(words)
                richness += min(vocabulary_ratio, 0.3)
            
            # Presence of structured data indicators
            if any(indicator in text for indicator in [':', '=', '->', '=>', '|']):
                richness += 0.2
            
            return min(richness, 1.0)
            
        except Exception as e:
            logger.warning(f"⚠️ Failed to calculate information richness: {e}")
            return 0.5
    
    def _post_process_chunks(self, chunks: List[Dict[str, Any]], full_text: str) -> List[Dict[str, Any]]:
        """Post-process chunks to ensure comprehensive coverage"""
        try:
            if not chunks:
                return chunks
            
            # Sort chunks by importance score (highest first)
            chunks.sort(key=lambda x: x.get('importance_score', 0.5), reverse=True)
            
            # Ensure we have good coverage of important content
            enhanced_chunks = []
            
            for chunk in chunks:
                enhanced_chunks.append(chunk)
                
                # Add position context for better understanding
                chunk['coverage_info'] = {
                    'total_chunks': len(chunks),
                    'chunk_rank_by_importance': enhanced_chunks.index(chunk) + 1,
                    'contains_critical_info': chunk.get('importance_score', 0) > 0.7
                }
            
            # Log coverage statistics
            high_importance_count = sum(1 for chunk in enhanced_chunks if chunk.get('importance_score', 0) > 0.7)
            logger.info(f"📊 Chunk coverage: {high_importance_count}/{len(enhanced_chunks)} high-importance chunks")
            
            # Ensure minimum coverage of document
            if len(enhanced_chunks) < 3 and len(full_text) > 2000:
                logger.warning(f"⚠️ Low chunk count ({len(enhanced_chunks)}) for document size ({len(full_text)} chars)")
            
            return enhanced_chunks
            
        except Exception as e:
            logger.warning(f"⚠️ Failed to post-process chunks: {e}")
            return chunks
    
    def _calculate_chunk_quality_score(self, content: str) -> float:
        """Calculate overall quality score for a chunk"""
        try:
            quality_score = 0.0
            
            # Text coherence (sentence structure)
            sentences = content.split('.')
            if len(sentences) > 1:
                avg_length = sum(len(s.split()) for s in sentences if s.strip()) / max(1, len([s for s in sentences if s.strip()]))
                if 5 <= avg_length <= 30:
                    quality_score += 0.3
            
            # Information density
            words = content.split()
            if words:
                # Good word length distribution
                avg_word_length = sum(len(word) for word in words) / len(words)
                if 3 <= avg_word_length <= 8:
                    quality_score += 0.2
                
                # Vocabulary richness
                unique_words = len(set(word.lower() for word in words))
                if unique_words / len(words) > 0.6:  # Good vocabulary diversity
                    quality_score += 0.2
            
            # Structured content indicators
            if any(indicator in content for indicator in [':', ';', '(', ')', '[', ']']):
                quality_score += 0.1
            
            # Technical content indicators
            if self._has_technical_terms(content):
                quality_score += 0.1
            
            # Form field indicators (important for template filling)
            if re.search(r'[A-Za-z\s]+:\s*(?:$|_|\.\.\.)', content):
                quality_score += 0.1
            
            return min(quality_score, 1.0)
            
        except Exception as e:
            logger.warning(f"⚠️ Failed to calculate chunk quality score: {e}")
            return 0.5
    
    async def delete_document(self, document_id: str, device_id: str) -> bool:
        """Delete document and all its chunks"""
        try:
            # Get document metadata to verify it exists
            document = await document_repo.get_document_by_id(document_id)
            if not document:
                logger.warning(f"⚠️ Document {document_id} not found in database")
                # Continue with cleanup attempt anyway
            
            # Method 1: Delete all vectors for this document using metadata filtering (more reliable)
            logger.info(f"🗑️ Deleting all vectors for document {document_id} from device {device_id}")
            deletion_success = await pinecone_service.delete_document_vectors(document_id, device_id)
            
            # Method 2: Fallback - try to delete by chunk IDs if metadata filtering failed
            if not deletion_success and document and "chunk_count" in document:
                logger.info(f"🔄 Fallback: Attempting deletion by chunk IDs for document {document_id}")
                chunk_ids = [f"{document_id}_{i}" for i in range(document["chunk_count"])]
                deletion_success = await pinecone_service.delete_vectors(chunk_ids, device_id)
            
            # Delete from MongoDB
            await document_repo.delete_document(document_id)
            
            # Delete file from disk
            try:
                if document:
                    file_path = self.upload_dir / f"{document_id}_{document['filename']}"
                    if file_path.exists():
                        file_path.unlink()
                        logger.info(f"🗑️ Deleted file from disk: {file_path}")
            except Exception as e:
                logger.warning(f"⚠️ Could not delete file from disk: {e}")
            
            if deletion_success:
                logger.info(f"✅ Successfully deleted document {document_id} and all its chunks for device {device_id}")
            else:
                logger.warning(f"⚠️ Document {document_id} metadata deleted, but vector cleanup may have failed")
            
            return True  # Return True even if vector deletion failed, since metadata is cleaned up
            
        except Exception as e:
            logger.error(f"❌ Failed to delete document {document_id}: {e}")
            return False
    
    def _generate_fallback_embedding(self, text: str) -> List[float]:
        """Generate a simple hash-based embedding as fallback when Gemini API fails"""
        try:
            import hashlib
            import numpy as np
            
            # Create a deterministic embedding based on text hash
            text_hash = hashlib.md5(text.encode()).hexdigest()
            
            # Convert hash to a 1024-dimensional vector (to match Pinecone index)
            # Split the hash into chunks and create a repeatable pattern
            hash_chunks = [text_hash[i:i+8] for i in range(0, len(text_hash), 8)]
            
            # Create embedding vector
            embedding = []
            for i in range(1024):
                # Use different parts of the hash to create variety
                chunk_idx = i % len(hash_chunks)
                char_idx = i % len(hash_chunks[chunk_idx])
                
                # Convert hex character to float between -1 and 1
                hex_val = int(hash_chunks[chunk_idx][char_idx], 16)
                normalized_val = (hex_val - 7.5) / 7.5  # Normalize to [-1, 1]
                embedding.append(normalized_val)
            
            # Add some text-based features
            text_len_factor = min(len(text) / 1000.0, 1.0)  # Normalize text length
            word_count_factor = min(len(text.split()) / 100.0, 1.0)  # Normalize word count
            
            # Adjust first few dimensions based on text characteristics
            embedding[0] = text_len_factor
            embedding[1] = word_count_factor
            embedding[2] = 1.0 if any(char.isdigit() for char in text) else -1.0
            embedding[3] = 1.0 if any(char.isupper() for char in text) else -1.0
            
            logger.debug(f"📊 Generated fallback embedding with {len(embedding)} dimensions")
            return embedding
            
        except Exception as e:
            logger.error(f"❌ Failed to generate fallback embedding: {e}")
            # Return zero vector as last resort
            return [0.0] * 1024

    def _split_by_paragraphs(self, text: str) -> List[str]:
        """Split text into paragraphs for better chunking"""
        try:
            # Split by double newlines first (paragraphs)
            paragraphs = text.split('\n\n')
            
            # If no double newlines, try single newlines
            if len(paragraphs) <= 1:
                paragraphs = text.split('\n')
            
            # Filter out empty paragraphs and very short ones
            filtered_paragraphs = []
            for para in paragraphs:
                para = para.strip()
                if para and len(para) > 20:  # Must have at least 20 characters
                    filtered_paragraphs.append(para)
            
            return filtered_paragraphs
            
        except Exception as e:
            logger.warning(f"Failed to split by paragraphs: {e}")
            return [text]

    def _classify_content_type(self, text: str) -> str:
        """
        Classify the type of content based on text patterns
        """
        try:
            text_lower = text.lower()
            
            # Check for structured data patterns
            if any(keyword in text_lower for keyword in ['name:', 'model:', 'serial:', 'part number:', 'specifications:']):
                return "device_specification"
            elif any(keyword in text_lower for keyword in ['patient:', 'doctor:', 'diagnosis:', 'treatment:', 'medication:']):
                return "medical_record"
            elif any(keyword in text_lower for keyword in ['invoice', 'receipt', 'payment', 'total:', 'amount:']):
                return "financial_document"
            elif any(keyword in text_lower for keyword in ['manual', 'instructions', 'how to', 'step', 'procedure']):
                return "instruction_manual"
            elif text.count('\n') > 20 and '|' in text:
                return "tabular_data"
            elif text.count(':') > 5:
                return "structured_list"
            else:
                return "general_document"
                
        except Exception as e:
            logger.warning(f"Failed to classify content type: {e}")
            return "general_document"
    
    def _contains_structured_data(self, text: str) -> bool:
        """
        Check if text contains structured data patterns
        """
        try:
            # Look for common structured patterns
            patterns = [
                r'\w+:\s*\w+',  # key:value pairs
                r'\|\s*\w+\s*\|',  # table columns
                r'^\s*\d+\.\s+',  # numbered lists
                r'^\s*-\s+',  # bullet points
                r'\w+\s*=\s*\w+',  # assignments
            ]
            
            for pattern in patterns:
                if re.search(pattern, text, re.MULTILINE):
                    return True
            
            return False
            
        except Exception:
            return False
    
    def _contains_form_fields(self, text: str) -> bool:
        """
        Check if text contains form field patterns
        """
        try:
            form_indicators = [
                'name:', 'email:', 'phone:', 'address:', 'date:',
                'model number:', 'serial number:', 'part number:',
                'manufacturer:', 'description:', 'specifications:',
                '___', '[    ]', '______'  # Common form field patterns
            ]
            
            text_lower = text.lower()
            form_field_count = sum(1 for indicator in form_indicators if indicator in text_lower)
            
            return form_field_count >= 2
            
        except Exception:
            return False
    
    def _calculate_entity_density(self, text: str) -> float:
        """
        Calculate the density of named entities in the text
        """
        try:
            words = text.split()
            if not words:
                return 0.0
            
            # Simple heuristic for entity detection
            entity_patterns = [
                r'[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*',  # Proper nouns
                r'\b[A-Z]{2,}\b',  # Acronyms
                r'\b\d+(?:\.\d+)?\s*[A-Za-z]+\b',  # Numbers with units
                r'\b[A-Z]\d+[A-Z]?\d*\b',  # Model numbers
            ]
            
            entity_count = 0
            for pattern in entity_patterns:
                entity_count += len(re.findall(pattern, text))
            
            return min(entity_count / len(words), 1.0)
            
        except Exception:
            return 0.0
    
    def _calculate_information_richness(self, text: str) -> float:
        """
        Calculate information richness score based on vocabulary diversity and structure
        """
        try:
            if not text:
                return 0.0
            
            words = text.split()
            if len(words) < 5:
                return 0.5
            
            # Calculate vocabulary diversity
            unique_words = set(word.lower() for word in words)
            diversity_score = len(unique_words) / len(words)
            
            # Bonus for structured content
            structure_bonus = 0.1 if self._contains_structured_data(text) else 0.0
            
            # Bonus for technical terms
            technical_terms = ['specification', 'model', 'serial', 'device', 'system', 'component']
            technical_bonus = 0.1 if any(term in text.lower() for term in technical_terms) else 0.0
            
            richness_score = diversity_score + structure_bonus + technical_bonus
            return min(richness_score, 1.0)
            
        except Exception:
            return 0.5
    
    def _extract_semantic_keywords(self, text: str) -> list:
        """
        Extract semantic keywords from text
        """
        try:
            if not text:
                return []
            
            # Simple keyword extraction based on frequency and importance
            words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
            
            # Common stop words to filter out
            stop_words = {'the', 'and', 'for', 'are', 'but', 'not', 'you', 'all', 'can', 'had', 'her', 'was', 'one', 'our', 'out', 'day', 'get', 'has', 'him', 'his', 'how', 'man', 'new', 'now', 'old', 'see', 'two', 'way', 'who', 'boy', 'did', 'its', 'let', 'put', 'say', 'she', 'too', 'use'}
            
            # Count word frequencies
            word_freq = {}
            for word in words:
                if word not in stop_words and len(word) > 2:
                    word_freq[word] = word_freq.get(word, 0) + 1
            
            # Get top keywords
            sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
            keywords = [word for word, freq in sorted_words[:10] if freq > 1]
            
            return keywords[:5]  # Return top 5 keywords
            
        except Exception:
            return []

# Global instance
document_processor = DocumentProcessor()
