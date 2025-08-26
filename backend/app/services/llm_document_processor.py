"""
LLM-powered document structure analysis and content processing
Handles intelligent analysis of document structure and content removal
"""

import logging
from typing import Dict, List, Tuple, Set
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.text.paragraph import Paragraph
from .gemini_service import gemini_service

logger = logging.getLogger(__name__)

class LLMDocumentProcessor:
    """LLM-powered document structure analysis and processing"""
    
    def __init__(self):
        self.gemini = gemini_service
    
    async def match_headings_with_toc(self, doc: Document) -> Tuple[Set[str], Dict[str, int]]:
        """
        Match headings in document with table of contents and determine their levels.
        
        Args:
            doc: The Word document to analyze
            
        Returns:
            Tuple containing:
            - Set of headings that should be preserved
            - Dictionary mapping heading text to its level
        """
        try:
            # Extract TOC and document headings
            toc_entries = self._extract_toc_entries(doc)
            doc_headings = []
            
            # Collect headings from paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text and self._is_likely_heading(para):
                    heading_info = {
                        'text': text,
                        'style': para.style.name if para.style else None,
                        'properties': self._extract_paragraph_formatting(para)
                    }
                    doc_headings.append(heading_info)
            
            # Create LLM prompt to match headings
            prompt = f"""Analyze these document headings and match them with the table of contents entries.
Determine which headings are structural and should be preserved in a blank template.

Table of Contents Entries:
{self._format_list_with_dashes(toc_entries)}

Document Headings:
{self._format_heading_info(doc_headings)}

Consider:
1. Exact matches between TOC and document headings
2. Numbered sections (e.g., "1.2.3 Section Name")
3. Heading formatting and styles
4. Section hierarchy and structure

Return a list of the headings that should be preserved, one per line."""

            # Get LLM analysis
            response = await self.gemini.generate_response(prompt)
            
            # Process response to get preserved headings
            preserved_headings = set()
            for line in response.strip().split('\n'):
                heading = line.strip()
                if heading and any(h['text'] == heading for h in doc_headings):
                    preserved_headings.add(heading)
            
            # Determine heading levels
            heading_levels = {}
            for heading in preserved_headings:
                # Try to extract level from numbering
                match = re.match(r'^((?:\d+\.)*\d+)?\s*(.+)$', heading)
                if match:
                    number, title = match.groups()
                    if number:
                        level = len(number.rstrip('.').split('.'))
                        heading_levels[heading] = level
                        continue
                
                # Default to level 2 if no numbering found
                heading_levels[heading] = 2
            
            return preserved_headings, heading_levels
            
        except Exception as e:
            logger.error(f"❌ Heading matching failed: {e}")
            return set(), {}
    
    def _format_list_with_dashes(self, items: List[str]) -> str:
        """Format list items with dashes"""
        return '\n'.join(f"- {item}" for item in items)
    
    def _format_heading_info(self, headings: List[Dict]) -> str:
        """Format heading information for LLM prompt"""
        formatted = []
        for heading in headings:
            info = []
            if heading.get('style'):
                info.append(f"Style: {heading['style']}")
            
            properties = heading.get('properties', {})
            if any(run.get('bold') for run in properties.get('runs', [])):
                info.append("Bold")
            if any(run.get('font_size') and run.get('font_size') >= Pt(12) for run in properties.get('runs', [])):
                info.append("Large Font")
            
            formatted.append(f"- {heading['text']}")
            if info:
                formatted[-1] += f" ({', '.join(info)})"
        
        return '\n'.join(formatted)
    
    async def analyze_document_structure(
        self, 
        doc: Document
    ) -> Tuple[Dict[str, Dict], List[Dict], Dict[str, str]]:
        """
        Analyze complete document structure including:
        1. Heading hierarchy and numbering from TOC
        2. Content sections that should be preserved/removed
        3. Table structures to maintain
        
        Returns:
        - heading_structure: Complete heading hierarchy with numbering
        - preserved_tables: Table structures to maintain
        - section_content: Map of headings to their preserved content
        """
        try:
            # Extract and analyze TOC structure
            toc_entries = self._extract_toc_entries(doc)
            heading_structure = self._extract_heading_structure(toc_entries)
            
            # Analyze document content
            doc_sections = self._extract_document_sections(doc)
            
            # Match headings with TOC and process content
            heading_matches = await self._match_and_process_content(
                heading_structure,
                doc_sections
            )
            
            # Analyze tables
            preserved_tables = await self.analyze_table_structure(doc)
            
            return heading_structure, preserved_tables, heading_matches
            
        except Exception as e:
            logger.error(f"❌ Document structure analysis failed: {e}")
            return {}, [], {}
    
    def _extract_toc_entries(self, doc: Document) -> List[str]:
        """Extract table of contents entries from document"""
        toc_entries = []
        in_toc = False
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # Check for TOC start indicators
            if not in_toc and self._is_toc_start(text):
                in_toc = True
                continue
            
            # Check for TOC end
            if in_toc and self._is_toc_end(text):
                in_toc = False
                continue
            
            # Process TOC entry
            if in_toc and text:
                # Clean up the entry
                entry = self._clean_toc_entry(text)
                if entry:
                    toc_entries.append(entry)
        
        return toc_entries
    
    def _extract_heading_structure(self, headings: List[str]) -> Dict:
        """Extract numbering and hierarchical structure from headings"""
        structure = {}
        
        for heading in headings:
            # Extract numbering if present (e.g., "1.2.3 Heading Title")
            match = re.match(r'^((?:\d+\.)*\d+)?\s*(.+)$', heading)
            if match:
                number, title = match.groups()
                if number:
                    parts = number.rstrip('.').split('.')
                    current = structure
                    for i, part in enumerate(parts):
                        if part not in current:
                            current[part] = {
                                'title': title if i == len(parts)-1 else '',
                                'number': number.rstrip('.'),
                                'level': i + 1,
                                'children': {}
                            }
                        current = current[part]['children']
                else:
                    # Handle unnumbered headings
                    structure[title] = {
                        'title': title,
                        'number': '',
                        'level': 1,
                        'children': {}
                    }
        
        return structure
    
    def _extract_document_sections(self, doc: Document) -> List[Dict]:
        """Extract sections with their content from document"""
        sections = []
        current_section = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            if self._is_likely_heading(para):
                # Save previous section if exists
                if current_section:
                    sections.append(current_section)
                
                # Start new section
                current_section = {
                    'heading': text,
                    'content': [],
                    'formatting': self._extract_paragraph_formatting(para)
                }
            elif current_section:
                current_section['content'].append(text)
        
        # Add last section
        if current_section:
            sections.append(current_section)
        
        return sections
    
    async def _match_and_process_content(
        self,
        heading_structure: Dict,
        doc_sections: List[Dict]
    ) -> Dict[str, str]:
        """Match TOC headings with document sections and process content"""
        try:
            # Prepare LLM prompt
            prompt = f"""Analyze these document sections and match them with the table of contents structure.
For each section, determine what content should be preserved in a blank template.

Table of Contents Structure:
{self._format_structure(heading_structure)}

Document Sections:
{self._format_sections(doc_sections)}

For each section:
1. Match it with the corresponding TOC heading
2. Analyze the content to determine:
   - What structural elements to preserve
   - What descriptive content to remove
   - What placeholders or blank fields to maintain
3. Return only the essential structure with descriptions removed

Return a JSON object mapping section headings to their processed content:
{{
    "heading": {{
        "matched_toc_entry": "1.1 Example Heading",
        "preserved_content": "processed content with descriptions removed"
    }},
    ...
}}"""

            # Get LLM analysis
            response = await self.gemini.generate_response(prompt)
            
            # Parse response
            import json
            processed_sections = json.loads(response)
            
            return processed_sections
            
        except Exception as e:
            logger.error(f"❌ Failed to process sections: {e}")
            return {}
    
    def _format_structure(self, structure: Dict, indent=0) -> str:
        """Format heading structure for LLM prompt"""
        result = []
        for key, data in structure.items():
            prefix = "  " * indent
            number = data.get('number', '')
            title = data.get('title', '')
            if title:
                result.append(f"{prefix}{number} {title}")
            
            # Process children
            children = data.get('children', {})
            if children:
                result.extend(self._format_structure(children, indent + 1).split('\n'))
        
        return '\n'.join(result)
    
    def _format_sections(self, sections: List[Dict]) -> str:
        """Format document sections for LLM prompt"""
        result = []
        for section in sections:
            result.append(f"HEADING: {section['heading']}")
            result.append("CONTENT:")
            result.extend(section['content'])
            result.append("-" * 40)
        return '\n'.join(result)
    
    def _extract_paragraph_formatting(self, para: Paragraph) -> Dict:
        """Extract formatting information from paragraph"""
        formatting = {
            'style': para.style.name if para.style else None,
            'alignment': para.alignment,
            'runs': []
        }
        
        for run in para.runs:
            run_format = {
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'color': run.font.color.rgb if run.font.color.rgb else None
            }
            formatting['runs'].append(run_format)
        
        return formatting
    
    def _is_toc_start(self, text: str) -> bool:
        """Check if text indicates start of table of contents"""
        text_lower = text.lower()
        toc_indicators = [
            'table of contents',
            'contents',
            'index',
            'table des matières',
            's. no. contents page no'
        ]
        return any(indicator in text_lower for indicator in toc_indicators)
    
    def _is_toc_end(self, text: str) -> bool:
        """Check if text indicates end of table of contents"""
        if not text.strip():
            return False
        
        text_lower = text.lower()
        section_starters = [
            'introduction',
            'background',
            'overview',
            'purpose',
            'scope',
            'executive summary'
        ]
        
        return any(starter in text_lower for starter in section_starters)
    
    def _clean_toc_entry(self, text: str) -> str:
        """Clean up table of contents entry text"""
        # Remove page numbers and dots
        text = re.sub(r'\s*\.{2,}\s*\d+\s*$', '', text)
        text = re.sub(r'\s+\d+\s*$', '', text)
        
        # Clean up whitespace while preserving numbering
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text
    
    def _is_likely_heading(self, paragraph: Paragraph) -> bool:
        """Check if paragraph is likely a heading"""
        if not paragraph.text.strip():
            return False
        
        # Check formatting
        if paragraph.style and 'heading' in paragraph.style.name.lower():
            return True
            
        for run in paragraph.runs:
            if run.bold or (run.font.size and run.font.size >= Pt(12)):
                return True
        
        # Check patterns
        text = paragraph.text.strip()
        heading_patterns = [
            r'^\d+\.\s+[A-Z]',  # Numbered sections
            r'^\d+\.\d+\s+[A-Z]',  # Subsections
            r'^[A-Z][^.!?]+$',  # All caps or title case
            r'^(?:section|chapter)\s+\d+'  # Explicit markers
        ]
        
        return any(re.match(pattern, text) for pattern in heading_patterns)
    
    async def analyze_table_structure(self, doc: Document) -> List[Dict]:
        """Analyze tables to determine which to preserve"""
        preserved_tables = []
        
        for i, table in enumerate(doc.tables):
            table_data = self._extract_table_data(table)
            if await self._should_preserve_table(table_data):
                preserved_tables.append({
                    'index': i,
                    'structure': table_data,
                    'preserve': True
                })
        
        return preserved_tables
    
    def _extract_table_data(self, table) -> Dict:
        """Extract table data and structure"""
        data = {
            'headers': [],
            'sample_rows': []
        }
        
        if not table.rows:
            return data
            
        # Extract headers
        header_row = table.rows[0]
        data['headers'] = [cell.text.strip() for cell in header_row.cells]
        
        # Extract sample rows
        for row in table.rows[1:3]:  # Just first couple rows
            row_data = [cell.text.strip() for cell in row.cells]
            data['sample_rows'].append(row_data)
            
        return data
    
    async def _should_preserve_table(self, table_data: Dict) -> bool:
        """Use LLM to determine if table should be preserved"""
        try:
            prompt = f"""Analyze this table structure and determine if it should be preserved in a blank template.

Headers: {table_data['headers']}
Sample Rows: {table_data['sample_rows']}

Consider:
1. Is this a form or data entry table?
2. Does it provide important structure?
3. Would removing it make the template less usable?

Return ONLY 'true' or 'false'"""

            response = await self.gemini.generate_response(prompt)
            return response.lower().strip() == 'true'
            
        except Exception as e:
            logger.error(f"❌ Failed to analyze table: {e}")
            return True  # Preserve by default
