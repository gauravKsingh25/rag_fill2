"""
Deterministic Document Filler Service
Implements the explicit mapping approach to solve field confusion issues.
Uses python-docx for precise control over formatting and bold insertion.
"""

import re
from typing import Dict, List, Tuple, Optional, Any
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import logging
from pathlib import Path

from .template_mapping_service import TemplateMappingService, TemplateConfig

logger = logging.getLogger(__name__)

class DeterministicDocumentFiller:
    """
    Deterministic document filler using explicit field mapping.
    Solves the field confusion problem by using deterministic insertion
    based on template-specific mapping configurations.
    """
    
    def __init__(self):
        self.mapping_service = TemplateMappingService()
    
    def fill_document_deterministic(self, template_content: str, input_data: Dict[str, Any], 
                                  template_name: Optional[str] = None) -> Tuple[bool, str, Optional[Document]]:
        """
        Fill document using deterministic mapping approach
        Returns (success, message, filled_document)
        """
        try:
            # Auto-detect template if not specified
            if not template_name:
                template_name = self.mapping_service.detect_template_type(template_content)
                if not template_name:
                    return False, "Could not detect template type. Please specify template name.", None
            
            logger.info(f"Filling document using template: {template_name}")
            
            # Validate input data
            is_valid, errors, validated_data = self.mapping_service.validate_input_data(template_name, input_data)
            if not is_valid:
                error_msg = f"Input validation failed:\n" + "\n".join(errors)
                return False, error_msg, None
            
            # Get replacement mappings
            replacements = self.mapping_service.create_replacement_mapping(template_name, validated_data)
            
            # Fill document based on template type
            if template_name == "sales_tax_affidavit":
                doc = self._fill_sales_tax_affidavit(validated_data)
            elif template_name == "bond_and_bail":
                doc = self._fill_bond_and_bail(validated_data)
            elif template_name == "income_tax_extension":
                doc = self._fill_income_tax_extension(validated_data)
            else:
                return False, f"No filler implementation for template: {template_name}", None
            
            success_msg = f"Successfully filled {template_name} with {len(replacements)} field replacements"
            logger.info(success_msg)
            
            return True, success_msg, doc
            
        except Exception as e:
            error_msg = f"Error filling document: {str(e)}"
            logger.error(error_msg, exc_info=True)
            return False, error_msg, None
    
    def _fill_sales_tax_affidavit(self, data: Dict[str, Any]) -> Document:
        """Fill Sales Tax Affidavit template with validated data"""
        doc = Document()
        
        # Title
        title = doc.add_heading('BEFORE THE HONBLE MEMBER-TRIBUNAL ', 1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.add_text(data.get('tribunal_member', ''))
        title_run.bold = True
        
        doc.add_paragraph()
        
        # Reference line
        ref_para = doc.add_paragraph("Ref : In the case of ")
        ref_para.add_run(data.get('firm_name', '')).bold = True
        
        # Assessment year
        assess_para = doc.add_paragraph("Assessment Year ")
        assess_para.add_run(data.get('assessment_year', '')).bold = True
        
        doc.add_paragraph()
        
        # Affidavit heading
        affidavit_heading = doc.add_heading('AFFIDAVIT', 2)
        affidavit_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Deponent information
        deponent_para = doc.add_paragraph("Affidavit of Mr. ")
        deponent_para.add_run(data.get('deponent_name', '')).bold = True
        deponent_para.add_run(" S/o Mr. ")
        deponent_para.add_run(data.get('deponent_father_name', '')).bold = True
        deponent_para.add_run(", aged about ")
        deponent_para.add_run(str(data.get('deponent_age', ''))).bold = True
        deponent_para.add_run(" years R/o ")
        deponent_para.add_run(data.get('deponent_address', '')).bold = True
        
        doc.add_paragraph()
        
        # Affirmation text
        doc.add_paragraph("I do hereby solemnly affirm and declare as under:")
        
        doc.add_paragraph()
        
        # Financial details
        financial_para = doc.add_paragraph()
        financial_para.add_run("(a) Admitted turn over Rs. ")
        financial_para.add_run(data.get('admitted_turnover', '')).bold = True
        
        financial_para2 = doc.add_paragraph()
        financial_para2.add_run("(b) Assessed turn over Rs. ")
        financial_para2.add_run(data.get('assessed_turnover', '')).bold = True
        
        financial_para3 = doc.add_paragraph()
        financial_para3.add_run("(c) Disputed turn over Rs. ")
        financial_para3.add_run(data.get('disputed_turnover', '')).bold = True
        
        financial_para4 = doc.add_paragraph()
        financial_para4.add_run("(d) Disputed tax Rs. ")
        financial_para4.add_run(data.get('disputed_tax', '')).bold = True
        
        doc.add_paragraph()
        
        # Prayer/Appeal section
        appeal_text = data.get('appeal_request', f"It is therefore prayed that the disputed tax amount of Rs. {data.get('disputed_tax', '')} may be waived.")
        prayer_para = doc.add_paragraph("Prayer: ")
        prayer_para.add_run(appeal_text).bold = True
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Signature section
        sig_para = doc.add_paragraph("Signature of Deponent")
        doc.add_paragraph()
        doc.add_paragraph("Date: _______________")
        doc.add_paragraph("Place: _______________")
        
        return doc
    
    def _fill_bond_and_bail(self, data: Dict[str, Any]) -> Document:
        """Fill Bond and Bail Bond template with validated data"""
        doc = Document()
        
        # Title
        title = doc.add_heading('BOND AND BAIL BOND UNDER CRIMINAL PROCEDURE CODE 1973 AFTER ARREST UNDER A WARRANT', 1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Bond text
        bond_para1 = doc.add_paragraph("I (name) ")
        bond_para1.add_run(data.get('accused_name', '')).bold = True
        bond_para1.add_run(" of ")
        bond_para1.add_run(data.get('accused_address', '')).bold = True
        bond_para1.add_run(" do hereby bind myself in the sum of Rs. ")
        bond_para1.add_run(data.get('forfeiture_amount', '')).bold = True
        bond_para1.add_run(" to the State, that I shall attend on the ")
        bond_para1.add_run(str(data.get('appearance_day', ''))).bold = True
        bond_para1.add_run(" day of ")
        bond_para1.add_run(data.get('appearance_month', '')).bold = True
        bond_para1.add_run(", ")
        bond_para1.add_run(str(data.get('appearance_year', ''))).bold = True
        bond_para1.add_run(" before the ")
        bond_para1.add_run(data.get('court_name', '')).bold = True
        
        doc.add_paragraph()
        
        # Magistrate section
        magistrate_para = doc.add_paragraph("or before such other Court or Magistrate as the case may be transferred to, or before the District Magistrate of ")
        magistrate_para.add_run(data.get('magistrate_name', '')).bold = True
        magistrate_para.add_run(" to answer to the charge of ")
        magistrate_para.add_run(data.get('charge_details', '')).bold = True
        
        doc.add_paragraph()
        
        # Forfeiture clause
        doc.add_paragraph("And I bind myself to forfeit the said sum if I fail to attend as above bound.")
        
        doc.add_paragraph()
        
        # Surety section (if provided)
        if data.get('surety_name'):
            surety_para = doc.add_paragraph("And I ")
            surety_para.add_run(data.get('surety_name', '')).bold = True
            surety_para.add_run(" of ")
            surety_para.add_run(data.get('surety_address', '')).bold = True
            surety_para.add_run(" do hereby declare myself surety for the appearance of the said ")
            surety_para.add_run(data.get('accused_name', '')).bold = True
            surety_para.add_run(" and I bind myself in the sum of Rs. ")
            surety_para.add_run(data.get('surety_forfeiture_amount', data.get('forfeiture_amount', ''))).bold = True
            surety_para.add_run(" that he shall attend as above bound.")
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Signature sections
        doc.add_paragraph("Signature of the accused: _______________________")
        doc.add_paragraph()
        if data.get('surety_name'):
            doc.add_paragraph("Signature of the surety: _______________________")
            doc.add_paragraph()
        
        doc.add_paragraph("Date: _______________")
        doc.add_paragraph("Place: _______________")
        
        return doc
    
    def _fill_income_tax_extension(self, data: Dict[str, Any]) -> Document:
        """Fill Income Tax Extension Affidavit template with validated data"""
        doc = Document()
        
        # Title
        title = doc.add_heading('BEFORE THE INCOME TAX OFFICER, ', 1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.add_text(data.get('income_tax_officer', ''))
        title_run.bold = True
        
        doc.add_paragraph()
        
        # Matter section
        matter_para = doc.add_paragraph("In the matter of ")
        matter_para.add_run(data.get('company_name', '')).bold = True
        
        doc.add_paragraph()
        
        # Affidavit heading
        affidavit_heading = doc.add_heading('AFFIDAVIT FOR EXTENDING TIME', 2)
        affidavit_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Deponent information
        deponent_para = doc.add_paragraph("I, ")
        deponent_para.add_run(data.get('deponent_name', '')).bold = True
        deponent_para.add_run(" S/o ")
        deponent_para.add_run(data.get('deponent_father_name', '')).bold = True
        deponent_para.add_run(", aged ")
        deponent_para.add_run(str(data.get('deponent_age', ''))).bold = True
        deponent_para.add_run(" years, resident of ")
        deponent_para.add_run(data.get('deponent_address', '')).bold = True
        deponent_para.add_run(", do hereby solemnly affirm and declare as under:")
        
        doc.add_paragraph()
        
        # Numbered paragraphs
        para1 = doc.add_paragraph("1. That the original due date for filing return was ")
        para1.add_run(data.get('due_date_original', '')).bold = True
        para1.add_run(".")
        
        para2 = doc.add_paragraph("2. That notice under section ")
        para2.add_run(data.get('notice_section', '')).bold = True
        para2.add_run(" was received on ")
        para2.add_run(data.get('notice_date', '')).bold = True
        para2.add_run(".")
        
        para3 = doc.add_paragraph("3. That the books of accounts were closed on ")
        para3.add_run(data.get('accounts_closed_date', '')).bold = True
        para3.add_run(".")
        
        para4 = doc.add_paragraph("4. That extension was applied for till ")
        para4.add_run(data.get('extension_applied_till', '')).bold = True
        para4.add_run(".")
        
        para5 = doc.add_paragraph("5. That Form No. ")
        para5.add_run(data.get('form_no', '')).bold = True
        para5.add_run(" was filed on ")
        para5.add_run(data.get('form_filed_date', '')).bold = True
        para5.add_run(" with receipt no. ")
        para5.add_run(data.get('receipt_no', '')).bold = True
        para5.add_run(".")
        
        para6 = doc.add_paragraph("6. That the return was subsequently filed on ")
        para6.add_run(data.get('return_filed_date', '')).bold = True
        para6.add_run(".")
        
        if data.get('officer_verbal_order_date'):
            para7 = doc.add_paragraph("7. That verbal order was given by the officer on ")
            para7.add_run(data.get('officer_verbal_order_date', '')).bold = True
            para7.add_run(".")
        
        doc.add_paragraph()
        
        # Prayer
        doc.add_paragraph("I therefore pray that the extension of time may be granted and the return may be accepted.")
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Signature section
        doc.add_paragraph("Signature of Deponent: _______________________")
        doc.add_paragraph()
        doc.add_paragraph("Date: _______________")
        doc.add_paragraph("Place: _______________")
        
        return doc
    
    def save_filled_document(self, doc: Document, output_path: str, filename: str) -> str:
        """Save the filled document to specified path"""
        try:
            full_path = Path(output_path) / filename
            full_path.parent.mkdir(parents=True, exist_ok=True)
            
            doc.save(str(full_path))
            logger.info(f"Document saved to: {full_path}")
            return str(full_path)
            
        except Exception as e:
            logger.error(f"Error saving document: {str(e)}")
            raise
    
    def get_template_preview(self, template_name: str) -> Optional[str]:
        """Get a preview of what fields will be filled for a template"""
        config = self.mapping_service.get_template_config(template_name)
        if not config:
            return None
        
        preview_lines = [
            f"Template: {config.template_description}",
            f"Fields to be filled: {len(config.field_mappings)}",
            "",
            "Required fields:"
        ]
        
        for mapping in config.field_mappings:
            status = "REQUIRED" if mapping.required else "OPTIONAL"
            preview_lines.append(f"  • {mapping.template_location} → {mapping.json_key} ({mapping.field_type.value}) [{status}]")
        
        return "\n".join(preview_lines)
