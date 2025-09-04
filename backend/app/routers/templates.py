from fastapi import APIRouter, HTTPException, UploadFile, File, Form, BackgroundTasks
from fastapi.responses import FileResponse
from typing import List, Dict, Any
import uuid
import logging
import re
from pathlib import Path
import tempfile
import os
from datetime import datetime
from docx import Document

from app.models import TemplateRequest, TemplateResponse
from app.services.gemini_service import gemini_service
from app.services.pinecone_service import pinecone_service
from app.services.csv_processor import CSVProcessor
from app.routers.devices import get_device
from app.services import gcs_service
from app.database import mongodb
from app.routers.file_history import add_file_to_history
from app.services.file_cleanup_service import file_cleanup_service

router = APIRouter()
logger = logging.getLogger(__name__)

# Initialize CSV processor
csv_processor = CSVProcessor()

@router.post("/upload-and-fill", response_model=TemplateResponse)
async def upload_and_fill_template(
    device_id: str = Form(...),
    file: UploadFile = File(...),
    filling_mode: str = Form("general")
):
    """Upload a template file and fill it with device knowledge"""
    try:
        # Verify device exists
        await get_device(device_id)
        
        # Validate file type (only .docx for now)
        if not file.filename.endswith('.docx'):
            raise HTTPException(
                status_code=400,
                detail="Only .docx template files are supported"
            )
        
        # Read template content
        template_content = await file.read()
        
        # Process template
        filled_template_path, filled_fields, missing_fields = await process_template(
            template_content=template_content,
            filename=file.filename,
            device_id=device_id,
            filling_mode=filling_mode
        )
        
        return TemplateResponse(
            filled_template_url=f"/api/templates/download/{Path(filled_template_path).name}",
            filled_fields=filled_fields,
            missing_fields=missing_fields
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Failed to process template: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to process template: {e}")

async def process_template(
    template_content: bytes,
    filename: str,
    device_id: str,
    filling_mode: str = "general"
) -> tuple[str, Dict[str, str], List[str]]:
    """Process template and fill placeholders using enhanced question-based approach"""
    try:
        # Create temporary file for processing
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file.write(template_content)
            temp_file_path = temp_file.name
        
        # Load document
        doc = Document(temp_file_path)
        
        # Extract all text to analyze placeholders
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        full_text += paragraph.text + "\n"
        
        # Filter out table of contents, headers, footers before processing
        filtered_text = gemini_service._filter_template_content(full_text)
        
        # Extract missing fields using enhanced pattern matching on filtered content
        missing_field_info = await extract_missing_fields_enhanced(filtered_text)
        
        filled_fields = {}
        missing_fields = []
        print(missing_field_info)

        logger.info(f"🔍 Found {len(missing_field_info)} fields to fill: {[field['field_name'] for field in missing_field_info]}")
        
        # For each missing field, create targeted questions and search
        for field_info in missing_field_info:
            field_name = field_info['field_name']
            field_context = field_info['context']
            field_pattern = field_info['pattern']
            
            logger.info(f"🔍 Processing field: {field_name}")
            print(f"🔍 Field context: {field_context[:200]}...")  # Print first 200 chars of context

            # ENHANCED: Generate comprehensive targeted questions for this field
            questions = await gemini_service.generate_field_questions(field_name, field_context)
            print(f"🔍 Generated questions: {questions}")
            
            # ENHANCED: Comprehensive multi-query search approach
            all_query_vectors = []
            
            # Generate embeddings for all questions
            for question in questions:
                query_embedding = await gemini_service.get_embedding(question)
                all_query_vectors.append(query_embedding)
            
            # Also add direct field name and context embeddings
            field_embedding = await gemini_service.get_embedding(field_name)
            all_query_vectors.append(field_embedding)
            
            # Add context-aware embedding
            context_query = f"{field_name} information from {field_context[:100]}"
            context_embedding = await gemini_service.get_embedding(context_query)
            all_query_vectors.append(context_embedding)
            
            # ENHANCED: Use comprehensive search for maximum document coverage
            comprehensive_results = await pinecone_service.comprehensive_search(
                query_vectors=all_query_vectors,
                device_id=device_id,
                top_k_per_query=10,  # More results per query
                final_top_k=20       # More final results for comprehensive analysis
            )
            
            if comprehensive_results:
                # Extract high-quality context documents
                context_docs = []
                high_importance_docs = []
                
                for result in comprehensive_results:
                    content = result.content
                    metadata = result.metadata
                    
                    if len(content) > 50:  # Ensure meaningful content
                        context_docs.append(content)
                        
                        # Separate high-importance content
                        importance_score = metadata.get('importance_score', 0.5)
                        if importance_score > 0.7 or metadata.get('has_form_fields', False):
                            high_importance_docs.append(content)
                
                # Prioritize high-importance documents but include comprehensive context
                final_context_docs = high_importance_docs[:10] + context_docs[:15]
                final_context_docs = list(dict.fromkeys(final_context_docs))  # Remove duplicates while preserving order
                
                if len(final_context_docs) >= 5:  # Ensure sufficient context
                    # Use enhanced field filling with comprehensive context analysis
                    field_value = await gemini_service.fill_template_field_enhanced(
                        field_name=field_name,
                        field_context=field_context,
                        context_docs=final_context_docs,
                        questions=questions,
                        device_id=device_id
                    )
                    
                    if field_value and field_value.strip():
                        filled_fields[field_name] = field_value.strip()
                        logger.info(f"✅ Filled field '{field_name}': {field_value.strip()[:50]}...")
                        print(f"✅ Filled '{field_name}' with: {field_value.strip()}")
                    else:
                        missing_fields.append(field_name)
                        logger.warning(f"❌ Could not fill field: {field_name} (AI could not extract value)")
                        print(f"❌ Could not extract value for: {field_name}")
                else:
                    missing_fields.append(field_name)
                    logger.warning(f"❌ Could not fill field: {field_name} (insufficient context documents: {len(final_context_docs)})")
                    print(f"❌ Insufficient context for: {field_name} (only {len(final_context_docs)} docs)")
            else:
                missing_fields.append(field_name)
                logger.warning(f"❌ No search results for field: {field_name}")
                print(f"❌ No search results for: {field_name}")
        
        # Replace placeholders in document with enhanced pattern matching
        replacement_count = 0
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            updated_text = original_text
            
            for field_info in missing_field_info:
                field_name = field_info['field_name']
                field_pattern = field_info['pattern']
                pattern_type = field_info['pattern_type']
                
                if field_name in filled_fields:
                    value = filled_fields[field_name]
                    
                    # Enhanced replacement based on pattern type
                    if pattern_type == 'COLON_FIELD':
                        # For standard colon fields, append the value after the colon
                        if field_pattern in updated_text:
                            updated_text = updated_text.replace(field_pattern, f"{field_pattern} {value}")
                            replacement_count += 1
                    elif pattern_type == 'COLON_FIELD_UNDERLINE':
                        # Replace "Field: ___" with "Field: Value"
                        pattern_match = re.search(rf'{re.escape(field_name)}:\s*_+', updated_text)
                        if pattern_match:
                            updated_text = updated_text.replace(pattern_match.group(0), f"{field_name}: {value}")
                            replacement_count += 1
                    elif pattern_type == 'COLON_FIELD_DOTS':
                        # Replace "Field: ..." with "Field: Value"
                        pattern_match = re.search(rf'{re.escape(field_name)}:\s*\.+', updated_text)
                        if pattern_match:
                            updated_text = updated_text.replace(pattern_match.group(0), f"{field_name}: {value}")
                            replacement_count += 1
                    elif pattern_type in ['COLON_FIELD_BRACKET', 'COLON_FIELD_BRACE', 'COLON_FIELD_ANGLE']:
                        # Replace "Field: [placeholder]" with "Field: Value"
                        pattern_match = re.search(rf'{re.escape(field_name)}:\s*[\[\{{<].*?[\]\}}>]', updated_text)
                        if pattern_match:
                            updated_text = updated_text.replace(pattern_match.group(0), f"{field_name}: {value}")
                            replacement_count += 1
                    elif pattern_type == 'COLON_FIELD_PLACEHOLDER':
                        # Replace "Field: TBD" with "Field: Value"
                        pattern_match = re.search(rf'{re.escape(field_name)}:\s*(?:TBD|TBC|TODO|XXX|--|\?+)', updated_text)
                        if pattern_match:
                            updated_text = updated_text.replace(pattern_match.group(0), f"{field_name}: {value}")
                            replacement_count += 1
                    elif pattern_type in ['COLON_FIELD_END', 'COLON_FIELD_INLINE']:
                        # Handle different colon field variations
                        colon_pattern = f"{field_name}:"
                        if colon_pattern in updated_text:
                            # Replace "Field Name:" with "Field Name: Value"
                            updated_text = updated_text.replace(colon_pattern, f"{colon_pattern} {value}")
                            replacement_count += 1
                    elif pattern_type.startswith('IMPLICIT_'):
                        # Handle implicit fields (similar to colon fields)
                        colon_pattern = f"{field_name}:"
                        if colon_pattern in updated_text:
                            updated_text = updated_text.replace(colon_pattern, f"{colon_pattern} {value}")
                            replacement_count += 1
                    elif pattern_type in ['BRACKET_PLACEHOLDER', 'INSTRUCTION_BRACKET']:
                        # Replace bracketed placeholders
                        if field_pattern in updated_text:
                            updated_text = updated_text.replace(field_pattern, value)
                            replacement_count += 1
                    elif pattern_type in ['BRACE_PLACEHOLDER', 'INSTRUCTION_BRACE']:
                        # Replace brace placeholders
                        if field_pattern in updated_text:
                            updated_text = updated_text.replace(field_pattern, value)
                            replacement_count += 1
                    elif pattern_type in ['ANGLE_PLACEHOLDER', 'INSTRUCTION_ANGLE']:
                        # Replace angle bracket placeholders
                        if field_pattern in updated_text:
                            updated_text = updated_text.replace(field_pattern, value)
                            replacement_count += 1
                    elif pattern_type in ['VERY_LONG_UNDERLINE', 'LONG_UNDERLINE', 'SHORT_UNDERLINE']:
                        # For underlines, replace with value but preserve some formatting
                        if field_pattern in updated_text:
                            # Keep the format but replace underlines with value
                            if len(value) <= len(field_pattern):
                                # Value fits within underlines
                                centered_value = value.center(len(field_pattern))
                                updated_text = updated_text.replace(field_pattern, centered_value)
                            else:
                                # Value is longer than underlines
                                updated_text = updated_text.replace(field_pattern, value)
                            replacement_count += 1
                    elif pattern_type in ['DATE_UNDERLINE', 'DATE_FORMAT', 'DATE_FORMAT_US', 'DATE_FORMAT_LOWER', 'DATE_FORMAT_US_LOWER']:
                        # For date fields, format appropriately
                        if field_pattern in updated_text:
                            # Try to format as date if possible
                            formatted_value = format_date_value(value, pattern_type)
                            updated_text = updated_text.replace(field_pattern, formatted_value)
                            replacement_count += 1
                    elif pattern_type in ['MISSING_MARKER', 'TO_BE_FILLED_MARKER', 'FILL_IN_MARKER', 'TBD_MARKER', 'TBC_MARKER', 'PLACEHOLDER_MARKER']:
                        # Replace explicit markers
                        if field_pattern in updated_text:
                            updated_text = updated_text.replace(field_pattern, value)
                            replacement_count += 1
                    else:
                        # Standard replacement for other pattern types
                        if field_pattern in updated_text:
                            updated_text = updated_text.replace(field_pattern, value)
                            replacement_count += 1
            
            # Update paragraph if text changed
            if updated_text != original_text:
                paragraph.text = updated_text
                logger.info(f"🔄 Updated paragraph: {original_text[:50]}... -> {updated_text[:50]}...")
        
        # Also check tables for missing fields (tables are separate from paragraphs in docx)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text
                        updated_text = original_text
                        
                        for field_info in missing_field_info:
                            field_name = field_info['field_name']
                            field_pattern = field_info['pattern']
                            pattern_type = field_info['pattern_type']
                            
                            if field_name in filled_fields and field_pattern in updated_text:
                                value = filled_fields[field_name]
                                
                                # Apply same replacement logic as for paragraphs
                                if pattern_type.startswith('COLON_FIELD'):
                                    if pattern_type == 'COLON_FIELD_UNDERLINE':
                                        pattern_match = re.search(rf'{re.escape(field_name)}:\s*_+', updated_text)
                                        if pattern_match:
                                            updated_text = updated_text.replace(pattern_match.group(0), f"{field_name}: {value}")
                                            replacement_count += 1
                                    elif pattern_type == 'COLON_FIELD_DOTS':
                                        pattern_match = re.search(rf'{re.escape(field_name)}:\s*\.+', updated_text)
                                        if pattern_match:
                                            updated_text = updated_text.replace(pattern_match.group(0), f"{field_name}: {value}")
                                            replacement_count += 1
                                    else:
                                        updated_text = updated_text.replace(field_pattern, f"{field_pattern} {value}")
                                        replacement_count += 1
                                else:
                                    updated_text = updated_text.replace(field_pattern, value)
                                    replacement_count += 1
                        
                        if updated_text != original_text:
                            paragraph.text = updated_text
                            logger.info(f"🔄 Updated table cell: {original_text[:30]}... -> {updated_text[:30]}...")
        
        logger.info(f"🔄 Made {replacement_count} replacements in document")
        
        # Save filled template
        output_dir = Path("./filled_templates")
        output_dir.mkdir(exist_ok=True)
        
        filled_filename = f"filled_{uuid.uuid4().hex}_{filename}"
        filled_path = output_dir / filled_filename
        
        doc.save(str(filled_path))
        
        # Clean up temp file
        os.unlink(temp_file_path)

        # Try to upload the filled template to GCS and add to file history (not favorites)
        try:
            if gcs_service.is_available():
                # Use the centralized file history function
                await add_file_to_history(
                    filename=filled_filename,
                    file_path=str(filled_path),
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    file_type="filled_template",
                    gcs_folder="filled_templates"
                )
                logger.info(f"✅ Added filled template to file history: {filled_filename}")
        except Exception as e:
            logger.warning(f"Could not upload filled template to GCS or persist metadata: {e}")
        
        logger.info(f"✅ Template processed: {len(filled_fields)} fields filled, {len(missing_fields)} missing")
        logger.info(f"✅ Filled fields: {list(filled_fields.keys())}")
        logger.info(f"❌ Missing fields: {missing_fields}")
        
        return str(filled_path), filled_fields, missing_fields
        
    except Exception as e:
        logger.error(f"❌ Failed to process template: {e}")
        raise

async def extract_missing_fields_enhanced(template_content: str) -> List[Dict[str, str]]:
    """Extract missing fields with comprehensive pattern matching and context analysis, focusing on main content"""
    try:
        missing_fields = []
        
        # Define comprehensive patterns for missing fields (order matters - more specific first)
        patterns = [
            # Explicit missing markers
            (r'\[MISSING\]', 'MISSING_MARKER'),
            (r'\[TO BE FILLED\]', 'TO_BE_FILLED_MARKER'),
            (r'\[FILL\s*IN\]', 'FILL_IN_MARKER'),
            (r'\[TBD\]', 'TBD_MARKER'),
            (r'\[TBC\]', 'TBC_MARKER'),
            (r'\[PLACEHOLDER\]', 'PLACEHOLDER_MARKER'),
            (r'\[INSERT\s*[^\]]*\]', 'INSERT_MARKER'),
            (r'\[ENTER\s*[^\]]*\]', 'ENTER_MARKER'),
            (r'\[ADD\s*[^\]]*\]', 'ADD_MARKER'),
            
            # Bracketed placeholders (more specific patterns first)
            (r'\[(?:Enter|Insert|Add|Type|Fill|Complete|Specify|Provide)\s+[^\]]*\]', 'INSTRUCTION_BRACKET'),
            (r'\[[A-Za-z][^\]]*\]', 'BRACKET_PLACEHOLDER'),
            
            # Curly braces placeholders
            (r'\{(?:Enter|Insert|Add|Type|Fill|Complete|Specify|Provide)\s+[^}]*\}', 'INSTRUCTION_BRACE'),
            (r'\{[A-Za-z][^}]*\}', 'BRACE_PLACEHOLDER'),
            
            # Angle bracket placeholders
            (r'<(?:Enter|Insert|Add|Type|Fill|Complete|Specify|Provide)\s+[^>]*>', 'INSTRUCTION_ANGLE'),
            (r'<[A-Za-z][^>]*>', 'ANGLE_PLACEHOLDER'),
            
            # Enhanced underlines and dots patterns
            (r'_{10,}', 'VERY_LONG_UNDERLINE'),    # Very long underlines (10+ chars)
            (r'_{5,9}', 'LONG_UNDERLINE'),         # Long underlines (5-9 chars)
            (r'_{3,4}', 'SHORT_UNDERLINE'),        # Short underlines (3-4 chars)
            (r'\.{6,}', 'VERY_LONG_DOTS'),         # Very long dots
            (r'\.{4,5}', 'LONG_DOTS'),             # Long dots
            (r'\.{3}', 'THREE_DOTS'),              # Exactly three dots
            
            # Enhanced form field patterns with colons - FOCUS ON THESE
            (r'[A-Za-z][A-Za-z\s\(\)/&\-,\.]*:\s*_+', 'COLON_FIELD_UNDERLINE'),    # "Field: ___"
            (r'[A-Za-z][A-Za-z\s\(\)/&\-,\.]*:\s*\.+', 'COLON_FIELD_DOTS'),        # "Field: ..."
            (r'[A-Za-z][A-Za-z\s\(\)/&\-,\.]*:\s*\[.*?\]', 'COLON_FIELD_BRACKET'), # "Field: [value]"
            (r'[A-Za-z][A-Za-z\s\(\)/&\-,\.]*:\s*\{.*?\}', 'COLON_FIELD_BRACE'),   # "Field: {value}"
            (r'[A-Za-z][A-Za-z\s\(\)/&\-,\.]*:\s*$', 'COLON_FIELD_END'),           # "Field: " at end of line
            (r'[A-Za-z][A-Za-z\s\(\)/&\-,\.]*:\s*(?=\s|$)', 'COLON_FIELD_INLINE'), # "Field: " followed by space
            
            # Table cell patterns
            (r'\|\s*\|\s*\|', 'EMPTY_TABLE_CELL'),     # Empty table cells
            (r'\|\s*_+\s*\|', 'TABLE_UNDERLINE_CELL'), # Table cells with underlines
            (r'\|\s*\.+\s*\|', 'TABLE_DOTS_CELL'),     # Table cells with dots
            
            # Enhanced date patterns
            (r'__/__/____', 'DATE_UNDERLINE'),
            (r'_/_/__', 'SHORT_DATE_UNDERLINE'),
            (r'DD/MM/YYYY', 'DATE_FORMAT'),
            (r'MM/DD/YYYY', 'DATE_FORMAT_US'),
            (r'dd/mm/yyyy', 'DATE_FORMAT_LOWER'),
            (r'mm/dd/yyyy', 'DATE_FORMAT_US_LOWER'),
            (r'Date:\s*$', 'DATE_FIELD'),
            (r'Created:\s*$', 'CREATED_FIELD'),
            (r'Modified:\s*$', 'MODIFIED_FIELD'),
            (r'Approved:\s*$', 'APPROVED_FIELD'),
            
            # Enhanced signature patterns
            (r'Signature:\s*$', 'SIGNATURE_FIELD'),
            (r'Signed:\s*$', 'SIGNED_FIELD'),
            (r'By:\s*$', 'BY_FIELD'),
            (r'Authorized:\s*$', 'AUTHORIZED_FIELD'),
            (r'Approved by:\s*$', 'APPROVED_BY_FIELD'),
            
            # Enhanced number patterns
            (r'No\.?\s*:\s*$', 'NUMBER_FIELD'),
            (r'#\s*:\s*$', 'HASH_NUMBER_FIELD'),
            (r'ID\s*:\s*$', 'ID_FIELD'),
            (r'Code\s*:\s*$', 'CODE_FIELD'),
            (r'Reference\s*:\s*$', 'REFERENCE_FIELD'),
            
            # Common template field patterns
            (r'Name\s*:\s*$', 'NAME_FIELD'),
            (r'Model\s*:\s*$', 'MODEL_FIELD'),
            (r'Version\s*:\s*$', 'VERSION_FIELD'),
            (r'Type\s*:\s*$', 'TYPE_FIELD'),
            (r'Description\s*:\s*$', 'DESCRIPTION_FIELD'),
            (r'Manufacturer\s*:\s*$', 'MANUFACTURER_FIELD'),
            (r'Company\s*:\s*$', 'COMPANY_FIELD'),
            (r'Address\s*:\s*$', 'ADDRESS_FIELD'),
            (r'Phone\s*:\s*$', 'PHONE_FIELD'),
            (r'Email\s*:\s*$', 'EMAIL_FIELD'),
        ]
        
        lines = template_content.split('\n')
        
        # Process each line for patterns
        for line_num, line in enumerate(lines):
            original_line = line
            
            # Skip lines that look like TOC entries or headers/footers
            if is_toc_or_header_line(line):
                continue
            
            # First, handle colon-based fields specially (MAIN FOCUS)
            colon_fields = extract_colon_fields(line, line_num, lines)
            missing_fields.extend(colon_fields)
            
            # Then handle other patterns
            for pattern, pattern_type in patterns:
                if pattern_type.startswith('COLON_'):
                    continue  # Already handled above
                    
                matches = re.finditer(pattern, line, re.IGNORECASE)
                for match in matches:
                    matched_text = match.group()
                    
                    # Extract field name from context
                    field_name = extract_field_name_from_context_enhanced(
                        line, match.start(), matched_text, pattern_type
                    )
                    
                    # Skip if field name is too generic or empty
                    if not field_name or len(field_name.strip()) < 2:
                        continue
                        
                    # Skip common false positives
                    if field_name.lower() in ['the', 'and', 'for', 'with', 'from', 'page', 'section', 'of', 'in', 'to', 'on', 'at']:
                        continue
                    
                    # Get surrounding context (more context for better understanding)
                    context_lines = []
                    for i in range(max(0, line_num - 3), min(len(lines), line_num + 4)):
                        if i < len(lines) and not is_toc_or_header_line(lines[i]):
                            context_lines.append(lines[i].strip())
                    context = ' '.join(context_lines)
                    
                    missing_fields.append({
                        'field_name': field_name,
                        'pattern': matched_text,
                        'context': context,
                        'line': line.strip(),
                        'pattern_type': pattern_type,
                        'line_number': line_num,
                        'position': match.start()
                    })
        
        # ENHANCED: Also look for implicit missing fields (fields that should be filled but appear empty)
        implicit_fields = extract_implicit_missing_fields(template_content, lines)
        missing_fields.extend(implicit_fields)
        
        # Remove duplicates and rank by importance
        unique_fields = {}
        for field in missing_fields:
            key = field['field_name'].lower().strip()
            
            # Prefer more specific pattern types
            if key not in unique_fields or is_better_pattern_type(field['pattern_type'], unique_fields[key]['pattern_type']):
                unique_fields[key] = field
                unique_fields[key]['field_name'] = field['field_name']  # Keep original case
        
        # Sort by line number for logical processing order
        result = sorted(unique_fields.values(), key=lambda x: x.get('line_number', 0))
        
        logger.info(f"🔍 Extracted {len(result)} unique fields from {len(missing_fields)} total matches")
        return result
        
    except Exception as e:
        logger.error(f"❌ Failed to extract missing fields: {e}")
        return []

def extract_implicit_missing_fields(template_content: str, lines: List[str]) -> List[Dict[str, str]]:
    """Extract implicit missing fields - fields that appear to need filling but don't have obvious placeholders"""
    try:
        implicit_fields = []
        
        # Patterns for fields that might have implicit missing values
        implicit_patterns = [
            # Fields with empty values or minimal placeholders
            (r'([A-Za-z][A-Za-z\s\(\)/&\-,\.]*?):\s*(?:N/A|TBD|TBC|TODO|XXX|--|\?+)\s*$', 'IMPLICIT_PLACEHOLDER'),
            (r'([A-Za-z][A-Za-z\s\(\)/&\-,\.]*?):\s*(?:\[.*?\]|\{.*?\}|<.*?>)\s*$', 'IMPLICIT_BRACKETED'),
            
            # Common medical device template fields that are often empty
            (r'(Generic\s+[Nn]ame):\s*$', 'IMPLICIT_GENERIC_NAME'),
            (r'(Device\s+[Nn]ame):\s*$', 'IMPLICIT_DEVICE_NAME'),
            (r'(Product\s+[Nn]ame):\s*$', 'IMPLICIT_PRODUCT_NAME'),
            (r'(Model\s+(?:No|Number|Name)):\s*$', 'IMPLICIT_MODEL'),
            (r'(Document\s+(?:No|Number)):\s*$', 'IMPLICIT_DOC_NO'),
            (r'(Serial\s+(?:No|Number)):\s*$', 'IMPLICIT_SERIAL'),
            (r'(Part\s+(?:No|Number)):\s*$', 'IMPLICIT_PART_NO'),
            (r'(Catalog\s+(?:No|Number)):\s*$', 'IMPLICIT_CATALOG_NO'),
            (r'(Reference\s+(?:No|Number)):\s*$', 'IMPLICIT_REFERENCE'),
            (r'(Manufacturer):\s*$', 'IMPLICIT_MANUFACTURER'),
            (r'(Company\s+[Nn]ame):\s*$', 'IMPLICIT_COMPANY'),
            (r'(Supplier):\s*$', 'IMPLICIT_SUPPLIER'),
            (r'(Distributor):\s*$', 'IMPLICIT_DISTRIBUTOR'),
            
            # Regulatory and quality fields
            (r'(Regulatory\s+[Ss]tatus):\s*$', 'IMPLICIT_REGULATORY'),
            (r'(Classification):\s*$', 'IMPLICIT_CLASSIFICATION'),
            (r'(Risk\s+[Cc]lass):\s*$', 'IMPLICIT_RISK_CLASS'),
            (r'(Intended\s+[Uu]se):\s*$', 'IMPLICIT_INTENDED_USE'),
            (r'(Indications?):\s*$', 'IMPLICIT_INDICATION'),
            (r'(Contraindications?):\s*$', 'IMPLICIT_CONTRAINDICATION'),
            
            # Technical specifications
            (r'(Power\s+[Rs]equirement):\s*$', 'IMPLICIT_POWER'),
            (r'(Operating\s+[Tt]emperature):\s*$', 'IMPLICIT_TEMP'),
            (r'(Storage\s+[Tt]emperature):\s*$', 'IMPLICIT_STORAGE_TEMP'),
            (r'(Humidity):\s*$', 'IMPLICIT_HUMIDITY'),
            (r'(Dimensions?):\s*$', 'IMPLICIT_DIMENSIONS'),
            (r'(Weight):\s*$', 'IMPLICIT_WEIGHT'),
            (r'(Material):\s*$', 'IMPLICIT_MATERIAL'),
            
            # Dates and versions
            (r'(Issue\s+[Dd]ate):\s*$', 'IMPLICIT_ISSUE_DATE'),
            (r'(Revision\s+[Dd]ate):\s*$', 'IMPLICIT_REVISION_DATE'),
            (r'(Effective\s+[Dd]ate):\s*$', 'IMPLICIT_EFFECTIVE_DATE'),
            (r'(Version):\s*$', 'IMPLICIT_VERSION'),
            (r'(Revision):\s*$', 'IMPLICIT_REVISION'),
        ]
        
        for line_num, line in enumerate(lines):
            if is_toc_or_header_line(line):
                continue
                
            for pattern, pattern_type in implicit_patterns:
                matches = re.finditer(pattern, line, re.IGNORECASE)
                for match in matches:
                    field_name = match.group(1).strip()
                    matched_text = match.group(0).strip()
                    
                    # Get context
                    context_lines = []
                    for i in range(max(0, line_num - 2), min(len(lines), line_num + 3)):
                        if i < len(lines) and not is_toc_or_header_line(lines[i]):
                            context_lines.append(lines[i].strip())
                    context = ' '.join(context_lines)
                    
                    implicit_fields.append({
                        'field_name': field_name,
                        'pattern': matched_text,
                        'context': context,
                        'line': line.strip(),
                        'pattern_type': pattern_type,
                        'line_number': line_num,
                        'position': match.start()
                    })
        
        return implicit_fields
        
    except Exception as e:
        logger.error(f"❌ Failed to extract implicit missing fields: {e}")
        return []

def is_toc_or_header_line(line: str) -> bool:
    """Check if a line is part of TOC, header, or footer and should be skipped"""
    line_lower = line.lower().strip()
    
    # Empty lines are ok
    if not line_lower:
        return False
    
    # Check for TOC patterns
    toc_patterns = [
        r'table\s+of\s+contents',
        r'contents',
        r'^\d+\.\s*.+\.\.\.\.\s*\d+$',  # TOC entry with dots
        r'^\d+\.\d+\s*.+\s+\d+\s*$',   # TOC sub-entry  
        r'^[A-Z\s]+\s+\d+\s*$',        # All caps with page number
        r'page\s+\d+',
        r'^\s*\d+\s*$',                 # Page number alone
        r'\.{3,}',                      # Dot leaders
    ]
    
    # Check for header/footer patterns
    header_footer_patterns = [
        r'header',
        r'footer', 
        r'confidential',
        r'proprietary',
        r'copyright',
        r'©\s*\d{4}',
        r'revision\s+\d+',
        r'version\s+\d+',
    ]
    
    for pattern in toc_patterns + header_footer_patterns:
        if re.search(pattern, line_lower):
            return True
    
    return False

def extract_colon_fields(line: str, line_num: int, all_lines: List[str]) -> List[Dict[str, str]]:
    """Extract fields that end with colons (form fields) - MAIN FOCUS for template filling"""
    colon_fields = []
    
    try:
        # Skip this line if it's part of TOC or headers
        if is_toc_or_header_line(line):
            return colon_fields
        
        # Enhanced patterns for fields ending with colon and different suffixes
        colon_patterns = [
            # Standard colon field patterns
            (r'([A-Za-z][A-Za-z\s\(\)/&\-,\.]*?):\s*_+\s*$', 'COLON_FIELD_UNDERLINE'),     # "Field: ___"
            (r'([A-Za-z][A-Za-z\s\(\)/&\-,\.]*?):\s*\.+\s*$', 'COLON_FIELD_DOTS'),         # "Field: ..."
            (r'([A-Za-z][A-Za-z\s\(\)/&\-,\.]*?):\s*\[.*?\]\s*$', 'COLON_FIELD_BRACKET'),  # "Field: [value]"
            (r'([A-Za-z][A-Za-z\s\(\)/&\-,\.]*?):\s*\{.*?\}\s*$', 'COLON_FIELD_BRACE'),    # "Field: {value}"
            (r'([A-Za-z][A-Za-z\s\(\)/&\-,\.]*?):\s*<.*?>\s*$', 'COLON_FIELD_ANGLE'),     # "Field: <value>"
            (r'([A-Za-z][A-Za-z\s\(\)/&\-,\.]*?):\s*(?:TBD|TBC|TODO|XXX|--|\?+)\s*$', 'COLON_FIELD_PLACEHOLDER'), # "Field: TBD"
            (r'([A-Za-z][A-Za-z\s\(\)/&\-,\.]*?):\s*$', 'COLON_FIELD_END'),              # "Field: " at end of line
        ]
        
        for pattern, pattern_type in colon_patterns:
            matches = re.finditer(pattern, line.strip())
            
            for match in matches:
                field_text = match.group(1).strip()
                
                # Clean up common prefixes and patterns
                field_text = re.sub(r'^\d+[\.\)]\s*', '', field_text)  # Remove numbering
                field_text = re.sub(r'^[a-z]\)\s*', '', field_text)    # Remove a), b), c) numbering
                field_text = re.sub(r'^\W+', '', field_text)           # Remove leading non-word chars
                
                # Skip very short or common words that are likely not real fields
                if len(field_text) < 2 or field_text.lower() in ['the', 'and', 'for', 'with', 'from', 'page', 'section', 'of', 'in', 'to', 'on', 'at']:
                    continue
                
                # Skip TOC-like entries even if they have colons
                if re.search(r'\d+\s*$', field_text) or re.search(r'page|section|chapter', field_text.lower()):
                    continue
                
                # Additional validation for meaningful field names
                if not re.search(r'[A-Za-z]{2,}', field_text):  # Must contain at least 2 consecutive letters
                    continue
                
                # Get context from surrounding lines (excluding TOC/header lines)
                context_lines = []
                for i in range(max(0, line_num - 2), min(len(all_lines), line_num + 3)):
                    if i < len(all_lines) and not is_toc_or_header_line(all_lines[i]):
                        context_lines.append(all_lines[i].strip())
                context = ' '.join(context_lines)
                
                # Determine the pattern based on what follows the colon
                matched_pattern = match.group(0).strip()
                
                colon_fields.append({
                    'field_name': field_text,
                    'pattern': matched_pattern,
                    'context': context,
                    'line': line.strip(),
                    'pattern_type': pattern_type,
                    'line_number': line_num,
                    'position': match.start()
                })
    
    except Exception as e:
        logger.error(f"❌ Failed to extract colon fields: {e}")
    
    return colon_fields

def is_better_pattern_type(new_type: str, existing_type: str) -> bool:
    """Determine if a new pattern type is better than existing one"""
    # Priority order (higher number = better)
    priority = {
        # Explicit markers (highest priority)
        'MISSING_MARKER': 15,
        'TO_BE_FILLED_MARKER': 14,
        'FILL_IN_MARKER': 13,
        'INSERT_MARKER': 12,
        'ENTER_MARKER': 12,
        'ADD_MARKER': 12,
        
        # Colon fields with specific patterns (very high priority)
        'COLON_FIELD_UNDERLINE': 11,
        'COLON_FIELD_DOTS': 11,
        'COLON_FIELD_BRACKET': 10,
        'COLON_FIELD_BRACE': 10,
        'COLON_FIELD': 9,
        'COLON_FIELD_END': 8,
        'COLON_FIELD_INLINE': 8,
        
        # Instruction patterns (high priority)
        'INSTRUCTION_BRACKET': 7,
        'INSTRUCTION_BRACE': 7,
        'INSTRUCTION_ANGLE': 7,
        
        # Implicit fields (medium-high priority)
        'IMPLICIT_GENERIC_NAME': 6,
        'IMPLICIT_DEVICE_NAME': 6,
        'IMPLICIT_PRODUCT_NAME': 6,
        'IMPLICIT_MODEL': 6,
        'IMPLICIT_DOC_NO': 6,
        'IMPLICIT_MANUFACTURER': 6,
        'IMPLICIT_PLACEHOLDER': 5,
        'IMPLICIT_BRACKETED': 5,
        
        # Standard placeholders (medium priority)
        'BRACKET_PLACEHOLDER': 4,
        'BRACE_PLACEHOLDER': 4,
        'ANGLE_PLACEHOLDER': 4,
        
        # Specific field types (medium priority)
        'DATE_FIELD': 4,
        'SIGNATURE_FIELD': 4,
        'NUMBER_FIELD': 4,
        'NAME_FIELD': 4,
        'MODEL_FIELD': 4,
        'MANUFACTURER_FIELD': 4,
        
        # Generic markers (lower priority)
        'TBD_MARKER': 3,
        'TBC_MARKER': 3,
        'PLACEHOLDER_MARKER': 3,
        
        # Table patterns (lower priority)
        'TABLE_UNDERLINE_CELL': 3,
        'TABLE_DOTS_CELL': 3,
        'EMPTY_TABLE_CELL': 2,
        
        # Underlines and dots (lower priority)
        'VERY_LONG_UNDERLINE': 2,
        'LONG_UNDERLINE': 2,
        'SHORT_UNDERLINE': 1,
        'VERY_LONG_DOTS': 2,
        'LONG_DOTS': 1,
        'THREE_DOTS': 1,
        
        # Date format patterns (lower priority than explicit date fields)
        'DATE_UNDERLINE': 2,
        'DATE_FORMAT': 2,
        'DATE_FORMAT_US': 2,
        'DATE_FORMAT_LOWER': 2,
        'DATE_FORMAT_US_LOWER': 2,
    }
    
    new_priority = priority.get(new_type, 0)
    existing_priority = priority.get(existing_type, 0)
    
    return new_priority > existing_priority

def format_date_value(value: str, pattern_type: str) -> str:
    """Format date values according to the expected pattern"""
    try:
        # If value is already in a date format, return as is
        if re.match(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}', value):
            return value
        
        # If it's a date word, try to format it
        from datetime import datetime
        import re
        
        # Try to extract date components from text
        date_patterns = [
            r'(\d{1,2})[/-](\d{1,2})[/-](\d{2,4})',  # DD/MM/YYYY or MM/DD/YYYY
            r'(\d{4})[/-](\d{1,2})[/-](\d{1,2})',    # YYYY/MM/DD
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, value)
            if match:
                if pattern_type == 'DATE_FORMAT_US':
                    return f"{match.group(1)}/{match.group(2)}/{match.group(3)}"
                else:
                    return f"{match.group(1)}/{match.group(2)}/{match.group(3)}"
        
        # If no date pattern found, return original value
        return value
        
    except:
        return value

def extract_field_name_from_context_enhanced(
    line: str, 
    match_position: int, 
    matched_text: str, 
    pattern_type: str
) -> str:
    """Enhanced field name extraction with better context understanding"""
    try:
        # For colon fields, the field name is already in the matched text
        if pattern_type.startswith('COLON_FIELD'):
            # Extract field name from patterns like "Field Name:" or "Field Name: ___"
            colon_match = re.search(r'([A-Za-z][A-Za-z\s\(\)/&\-,\.]*?):', matched_text)
            if colon_match:
                return clean_field_name(colon_match.group(1))
            return matched_text.rstrip(':').strip()
        
        # For implicit fields, extract the field name
        if pattern_type.startswith('IMPLICIT_'):
            if ':' in matched_text:
                return clean_field_name(matched_text.split(':')[0])
            return clean_field_name(matched_text)
        
        # Look for field names before the placeholder
        before_text = line[:match_position].strip()
        after_text = line[match_position + len(matched_text):].strip()
        
        # Enhanced patterns for field names
        field_patterns = [
            # Most specific patterns first
            (r'([A-Za-z][A-Za-z\s\(\)/&\-,\.]*?):\s*$', 'FIELD_COLON'),                    # "Field Name: [MISSING]"
            (r'(\b[A-Z][A-Za-z\s]*(?:No|Number|Name|Date|ID|Code|Model|Type|Version))\s*$', 'FIELD_KEYWORD'), # "Document Number [MISSING]"
            (r'(\b[A-Z][A-Za-z\s]{2,})\s*$', 'FIELD_CAPITALIZED'),                        # "Generic Name [MISSING]"
            (r'([A-Za-z][A-Za-z\s]{2,})\s*$', 'FIELD_GENERAL'),                           # Any text before placeholder
        ]
        
        # Try to extract from before text
        for pattern, _ in field_patterns:
            match = re.search(pattern, before_text)
            if match:
                field_name = match.group(1).strip()
                field_name = clean_field_name(field_name)
                if field_name and len(field_name) > 1:
                    return field_name
        
        # If no good match before, try after text for certain patterns
        if pattern_type in ['VERY_LONG_UNDERLINE', 'LONG_UNDERLINE', 'SHORT_UNDERLINE']:
            after_patterns = [
                (r'^([A-Za-z][A-Za-z\s]*)', 'AFTER_FIELD'),  # Text after underlines
            ]
            for pattern, _ in after_patterns:
                match = re.search(pattern, after_text)
                if match:
                    field_name = match.group(1).strip()
                    field_name = clean_field_name(field_name)
                    if field_name and len(field_name) > 1:
                        return field_name
        
        # For instruction patterns, try to extract meaningful names from the instruction
        if 'INSTRUCTION' in pattern_type:
            # Extract from patterns like "[Enter your name]" -> "Name"
            instruction_patterns = [
                (r'(?:Enter|Insert|Add|Type|Fill|Complete|Specify|Provide)\s+(?:your\s+)?(.+)', 'INSTRUCTION_EXTRACT'),
                (r'(.+?)\s+(?:here|field|value)', 'INSTRUCTION_FIELD'),
                (r'(.+)', 'INSTRUCTION_GENERIC'),
            ]
            
            # Remove brackets/braces first
            instruction_text = re.sub(r'[\[\]{}()<>]', '', matched_text).strip()
            
            for pattern, _ in instruction_patterns:
                match = re.search(pattern, instruction_text, re.IGNORECASE)
                if match:
                    extracted = match.group(1).strip()
                    if len(extracted) > 2 and not any(word in extracted.lower() for word in ['enter', 'insert', 'add', 'type', 'fill']):
                        return clean_field_name(extracted)
        
        # Generate descriptive name based on pattern type and context
        return generate_field_name_from_pattern(matched_text, pattern_type, line)
        
    except Exception as e:
        logger.error(f"❌ Failed to extract field name from context: {e}")
        return matched_text

def clean_field_name(field_name: str) -> str:
    """Clean and standardize field names"""
    try:
        # Remove common prefixes (numbering, bullets, etc.)
        field_name = re.sub(r'^\d+[\.\)]\s*', '', field_name)  # Remove "1. ", "2) "
        field_name = re.sub(r'^[a-z]\)\s*', '', field_name)    # Remove "a) ", "b) "
        field_name = re.sub(r'^\W+', '', field_name)           # Remove leading non-word chars
        
        # Remove common suffixes that don't add value
        field_name = re.sub(r'\s*(is|are|was|were)\s*$', '', field_name, flags=re.IGNORECASE)
        
        # Standardize spacing
        field_name = re.sub(r'\s+', ' ', field_name).strip()
        
        # Title case for better readability
        if field_name and not field_name.isupper():
            field_name = field_name.title()
        
        return field_name
    except:
        return field_name

def generate_field_name_from_pattern(matched_text: str, pattern_type: str, line: str) -> str:
    """Generate descriptive field names based on pattern type"""
    try:
        # Create descriptive names based on pattern type
        pattern_names = {
            'MISSING_MARKER': 'Missing Information',
            'TO_BE_FILLED_MARKER': 'To Be Filled',
            'FILL_IN_MARKER': 'Fill In',
            'TBD_MARKER': 'To Be Determined',
            'TBC_MARKER': 'To Be Confirmed',
            'PLACEHOLDER_MARKER': 'Placeholder',
            'BRACKET_PLACEHOLDER': 'Information',
            'BRACE_PLACEHOLDER': 'Information',
            'ANGLE_PLACEHOLDER': 'Information', 
            'INSTRUCTION_BRACKET': 'Instruction Field',
            'INSTRUCTION_ANGLE': 'Instruction Field',
            'LONG_UNDERLINE': 'Signature',
            'SHORT_UNDERLINE': 'Field',
            'LONG_DOTS': 'Information',
            'THREE_DOTS': 'Continuation',
            'DATE_UNDERLINE': 'Date',
            'DATE_FORMAT': 'Date',
            'DATE_FORMAT_US': 'Date',
            'DATE_FIELD': 'Date',
            'SIGNATURE_FIELD': 'Signature',
            'SIGNED_FIELD': 'Signature',
            'BY_FIELD': 'Signed By',
            'NUMBER_FIELD': 'Number',
            'HASH_NUMBER_FIELD': 'Number',
            'EMPTY_TABLE_CELL': 'Table Data',
        }
        
        base_name = pattern_names.get(pattern_type, 'Field')
        
        # Try to make it more specific based on context
        line_lower = line.lower()
        
        # Check for context clues in the line
        if 'name' in line_lower:
            return 'Name'
        elif 'date' in line_lower:
            return 'Date'
        elif 'number' in line_lower or 'no.' in line_lower or '#' in line_lower:
            return 'Number'
        elif 'model' in line_lower:
            return 'Model'
        elif 'version' in line_lower:
            return 'Version'
        elif 'serial' in line_lower:
            return 'Serial Number'
        elif 'manufacturer' in line_lower:
            return 'Manufacturer'
        elif 'generic' in line_lower:
            return 'Generic Name'
        elif 'document' in line_lower:
            return 'Document'
        elif 'signature' in line_lower or 'sign' in line_lower:
            return 'Signature'
        elif 'address' in line_lower:
            return 'Address'
        elif 'phone' in line_lower or 'tel' in line_lower:
            return 'Phone'
        elif 'email' in line_lower:
            return 'Email'
        
        return base_name
        
    except:
        return 'Field'

def extract_field_name_from_context(line: str, match_position: int, matched_text: str) -> str:
    """Legacy function - keeping for compatibility"""
    return extract_field_name_from_context_enhanced(line, match_position, matched_text, 'UNKNOWN')

@router.get("/download/{filename}")
async def download_filled_template(filename: str, background_tasks: BackgroundTasks):
    """Download a filled template file and schedule cleanup after download"""
    try:
        file_path = Path("./filled_templates") / filename
        
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="Template file not found")
        
        # Schedule file cleanup after download (30 seconds delay to ensure download completes)
        background_tasks.add_task(
            file_cleanup_service.schedule_file_cleanup,
            str(file_path),
            30  # 30 seconds delay
        )
        
        logger.info(f"📥 Downloading template: {filename} (cleanup scheduled)")
        
        return FileResponse(
            path=str(file_path),
            filename=filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Failed to download template: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to download template: {e}")

@router.post("/analyze")
async def analyze_template(
    device_id: str = Form(...),
    file: UploadFile = File(...)
):
    """Analyze a template to show what fields can be filled"""
    try:
        # Verify device exists
        await get_device(device_id)
        
        # Validate file type
        if not file.filename.endswith('.docx'):
            raise HTTPException(
                status_code=400,
                detail="Only .docx template files are supported"
            )
        
        # Read and analyze template
        template_content = await file.read()
        
        # Create temporary file for analysis
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file.write(template_content)
            temp_file_path = temp_file.name
        
        try:
            # Load document and extract text
            doc = Document(temp_file_path)
            full_text = ""
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            full_text += paragraph.text + "\n"
            
            # Filter out unwanted content before field extraction
            filtered_text = gemini_service._filter_template_content(full_text)
            
            # Extract placeholder fields from filtered content
            placeholder_fields = await gemini_service.extract_template_fields(filtered_text)
            
            # For each field, check if we have relevant information
            field_analysis = {}
            
            for field in placeholder_fields:
                # Search for information related to this field
                query_embedding = await gemini_service.get_embedding(f"information about {field}")
                
                search_results = await pinecone_service.search_vectors(
                    query_vector=query_embedding,
                    device_id=device_id,
                    top_k=3
                )
                
                field_analysis[field] = {
                    "can_fill": len(search_results) > 0,
                    
                    "confidence": search_results[0].score if search_results else 0,
                    "sources": len(search_results)
                }
            
            analysis_result = {
                "device_id": device_id,
                "template_filename": file.filename,
                "analysis_type": "Template Field Analysis",
                "total_fields": len(placeholder_fields),
                "fillable_fields": len([f for f, a in field_analysis.items() if a["can_fill"]]),
                "field_analysis": field_analysis
            }
            
            logger.info(f"✅ Template analysis completed: {len([f for f, a in field_analysis.items() if a['can_fill']])}/{len(placeholder_fields)} fields can be filled")
            
            # Note: We don't add analysis results to file history since it's just analysis data, not a file
            # The analysis results are returned directly to the frontend for display
            
            return analysis_result
            
        finally:
            # Clean up temp file
            os.unlink(temp_file_path)
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Failed to analyze template: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to analyze template: {e}")

@router.post("/upload-and-fill-csv")
async def upload_and_fill_csv(
    device_id: str = Form(...),
    file: UploadFile = File(...),
    filling_mode: str = Form("general")
):
    """Upload a CSV file and fill empty cells with device knowledge"""
    try:
        # Verify device exists
        await get_device(device_id)
        
        # Validate file type
        if not file.filename.endswith('.csv'):
            raise HTTPException(
                status_code=400,
                detail="Only .csv files are supported"
            )
        
        # Read CSV content
        csv_content = await file.read()
        
        # Process CSV using RAG knowledge base
        result = await csv_processor.process_csv_file(
            csv_content=csv_content,
            filename=file.filename,
            device_id=device_id,
            filling_mode=filling_mode
        )
        
        # Add the filled CSV to file history (not favorites)
        if result.get('success') and result.get('filled_csv_url'):
            try:
                filled_filename = f"filled_{file.filename}"
                # Extract the actual filename from the download URL
                download_url = result.get('filled_csv_url', '')
                if '/download-csv/' in download_url:
                    actual_filename = download_url.split('/download-csv/')[-1]
                    local_file_path = Path("./filled_templates") / actual_filename
                    
                    if local_file_path.exists():
                        await add_file_to_history(
                            filename=filled_filename,
                            file_path=str(local_file_path),
                            file_obj=None,
                            content_type='text/csv',
                            file_type="processed_csv",
                            gcs_folder="processed_csv"
                        )
                        logger.info(f"✅ Added filled CSV to file history: {filled_filename}")
                    else:
                        logger.warning(f"Local file not found for history upload: {local_file_path}")
                else:
                    logger.warning(f"Could not extract filename from download URL: {download_url}")
            except Exception as e:
                logger.warning(f"Could not add filled CSV to file history: {e}")
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Failed to process CSV: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to process CSV: {e}")

@router.get("/download-csv/{filename}")
async def download_filled_csv(filename: str, background_tasks: BackgroundTasks):
    """Download a filled CSV file and schedule cleanup after download"""
    try:
        file_path = Path("./filled_templates") / filename
        
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="CSV file not found")
        
        # Schedule file cleanup after download (30 seconds delay to ensure download completes)
        background_tasks.add_task(
            file_cleanup_service.schedule_file_cleanup,
            str(file_path),
            30  # 30 seconds delay
        )
        
        logger.info(f"📥 Downloading CSV: {filename} (cleanup scheduled)")
        
        return FileResponse(
            path=str(file_path),
            filename=filename,
            media_type='text/csv'
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Failed to download CSV: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to download CSV: {e}")

@router.post("/analyze-csv")
async def analyze_csv(
    device_id: str = Form(...),
    file: UploadFile = File(...)
):
    """Analyze a CSV file to show which cells can be filled - similar to document analysis"""
    try:
        # Verify device exists
        await get_device(device_id)
        
        # Validate file type
        if not file.filename.endswith('.csv'):
            raise HTTPException(
                status_code=400,
                detail="Only .csv files are supported"
            )
        
        # Read and analyze CSV
        csv_content = await file.read()
        
        # Parse CSV to identify empty cells
        df = csv_processor._parse_csv_content(csv_content)
        empty_cells = csv_processor._identify_empty_cells(df)
        
        logger.info(f"📊 Analyzing CSV: {len(df)} rows, {len(df.columns)} columns, {len(empty_cells)} empty cells")
        
        # Analyze each empty cell to estimate fillability - analyze more cells for better analysis
        fillable_cells = 0
        cell_analysis = {}
        column_analysis = {}
        
        # Group empty cells by column for better analysis
        cells_by_column = {}
        for cell_info in empty_cells:
            col_name = cell_info['column']
            if col_name not in cells_by_column:
                cells_by_column[col_name] = []
            cells_by_column[col_name].append(cell_info)
        
        # Analyze each column's fillability
        for col_name, column_cells in cells_by_column.items():
            column_fillable = 0
            column_confidence_scores = []
            column_sample_analysis = {}
            
            # Analyze up to 5 cells per column for detailed analysis
            sample_cells = column_cells[:5]
            
            for cell_info in sample_cells:
                row_idx = cell_info['row']
                
                # Generate search query for this cell
                context_info = csv_processor._extract_cell_context(df, row_idx, col_name)
                queries = await csv_processor._generate_cell_queries(context_info)
                
                # Test search for this cell
                if queries:
                    try:
                        query_embedding = await gemini_service.get_embedding(queries[0])
                        search_results = await pinecone_service.search_vectors(
                            query_vector=query_embedding,
                            device_id=device_id,
                            top_k=5
                        )
                        
                        can_fill = len(search_results) > 0 and search_results[0].score > 0.3
                        confidence = search_results[0].score if search_results else 0
                        
                        if can_fill:
                            column_fillable += 1
                            fillable_cells += 1
                        
                        column_confidence_scores.append(confidence)
                        
                        cell_key = f"Row {row_idx + 1}"
                        column_sample_analysis[cell_key] = {
                            "can_fill": can_fill,
                            "confidence": round(confidence, 3),
                            "sources": len(search_results),
                            "sample_query": queries[0] if queries else "No query generated"
                        }
                    except Exception as e:
                        logger.warning(f"Failed to analyze cell [{row_idx}, {col_name}]: {e}")
                        column_sample_analysis[f"Row {row_idx + 1}"] = {
                            "can_fill": False,
                            "confidence": 0,
                            "sources": 0,
                            "error": str(e)
                        }
            
            # Calculate column-level statistics
            avg_confidence = sum(column_confidence_scores) / len(column_confidence_scores) if column_confidence_scores else 0
            fill_rate = column_fillable / len(sample_cells) if sample_cells else 0
            
            column_analysis[col_name] = {
                "empty_cells_in_column": len(column_cells),
                "sample_cells_analyzed": len(sample_cells),
                "fillable_cells": column_fillable,
                "fill_rate": round(fill_rate, 2),
                "average_confidence": round(avg_confidence, 3),
                "sample_analysis": column_sample_analysis,
                "data_pattern": csv_processor._detect_column_pattern(df[col_name].dropna().tolist()[:5])
            }
        
        # Calculate overall statistics
        total_analyzed = sum(len(cells[:5]) for cells in cells_by_column.values())
        overall_fill_rate = fillable_cells / total_analyzed if total_analyzed > 0 else 0
        
        analysis_result = {
            "device_id": device_id,
            "csv_filename": file.filename,
            "analysis_type": "CSV Field Analysis",
            "summary": {
                "total_rows": len(df),
                "total_columns": len(df.columns),
                "total_empty_cells": len(empty_cells),
                "columns_with_empty_cells": len(cells_by_column),
                "sample_cells_analyzed": total_analyzed,
                "fillable_cells": fillable_cells,
                "overall_fill_rate": round(overall_fill_rate, 2),
                "analysis_status": "completed"
            },
            "columns": list(df.columns),
            "column_analysis": column_analysis,
            "recommendations": {
                "high_fill_rate_columns": [col for col, analysis in column_analysis.items() if analysis["fill_rate"] > 0.7],
                "low_fill_rate_columns": [col for col, analysis in column_analysis.items() if analysis["fill_rate"] < 0.3],
                "total_processable": len([col for col, analysis in column_analysis.items() if analysis["fill_rate"] > 0.0])
            }
        }
        
        logger.info(f"✅ CSV analysis completed: {fillable_cells}/{total_analyzed} cells can be filled")
        
        # Note: We don't add analysis results to file history since it's just analysis data, not a file
        # The analysis results are returned directly to the frontend for display
        
        return analysis_result
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Failed to analyze CSV: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to analyze CSV: {e}")
